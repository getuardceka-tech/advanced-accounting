"""
GETUARD AGENCY - Računovodstveni softver za Crnu Goru
Backend API
"""
from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import re
import uuid
import jwt
import bcrypt
import requests
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, Cm
import copy

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

TEMPLATES_DIR = ROOT_DIR / "templates"
GENERATED_DIR = ROOT_DIR / "generated"
GENERATED_DIR.mkdir(exist_ok=True)

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Auth config
JWT_SECRET = os.environ.get('JWT_SECRET', 'getuard-agency-secret-2026-change-me')
JWT_ALGORITHM = 'HS256'
JWT_EXPIRE_HOURS = 24 * 7  # 7 dana

# Default credentials (seed on startup)
DEFAULT_USERNAME = 'getuard'
DEFAULT_PASSWORD = 'Getuard1994.'

app = FastAPI(title="Getuard Agency API")
api_router = APIRouter(prefix="/api")
security = HTTPBearer()


# ============== MODELS ==============

class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    access_token: str
    user: dict

class Agency(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    naziv: str = "ADVANCED ACCOUNTING"
    adresa: str = "Brajša bb, Ulcinj"
    grad: str = "Ulcinj"
    pib: str = ""
    pdv_broj: str = ""
    ziro_racun: str = ""
    banka: str = ""
    telefon: str = "+382 69 172 204"
    email: str = "Advanced.acct@hotmail.com"
    direktor_ime: str = "GETUARD CEKOVIQ"
    direktor_jmbg: str = "0806994223008"
    djelatnost: str = "Računovodstvene usluge"
    logo_url: str = ""

class Company(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    naziv: str
    naziv_skraceni: str = ""
    pib: str
    maticni_broj: str = ""
    pdv_broj: str = ""
    adresa: str = ""
    grad: str = ""
    djelatnost: str = ""
    sifra_djelatnosti: str = ""
    oblik_organizovanja: str = ""  # iz IRMS-a: "Preduzetnik", "Društvo sa ograničenom odgovornošću", itd.
    direktor_ime: str = ""
    direktor_jmbg: str = ""
    direktor_adresa: str = ""
    ziro_racun: str = ""
    banka: str = ""
    telefon: str = ""
    email: str = ""
    # Flagovi
    pdv_obveznik: bool = False
    ioppd_obveznik: bool = False
    aktivna: bool = True
    napomena: str = ""
    # IRMS status (auto-popunjavanje iz Poreske uprave CG)
    irms_status: str = ""  # "Registrovan", "U likvidaciji", "U stečaju", "Mirovanje poslovanja", ...
    irms_checked_at: str = ""  # ISO datum poslednje provjere
    datum_registracije: str = ""  # iz IRMS-a, YYYY-MM-DD
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CompanyCreate(BaseModel):
    naziv: str
    pib: str
    naziv_skraceni: Optional[str] = ""
    maticni_broj: Optional[str] = ""
    pdv_broj: Optional[str] = ""
    adresa: Optional[str] = ""
    grad: Optional[str] = ""
    djelatnost: Optional[str] = ""
    sifra_djelatnosti: Optional[str] = ""
    oblik_organizovanja: Optional[str] = ""
    direktor_ime: Optional[str] = ""
    direktor_jmbg: Optional[str] = ""
    direktor_adresa: Optional[str] = ""
    ziro_racun: Optional[str] = ""
    banka: Optional[str] = ""
    telefon: Optional[str] = ""
    email: Optional[str] = ""
    pdv_obveznik: Optional[bool] = False
    ioppd_obveznik: Optional[bool] = False
    aktivna: Optional[bool] = True
    napomena: Optional[str] = ""
    irms_status: Optional[str] = ""
    irms_checked_at: Optional[str] = ""
    datum_registracije: Optional[str] = ""

class Employee(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    objekat_id: str = ""  # opciono — za koji objekat (poslovnicu) radi
    ime: str
    prezime: str
    jmbg: str = ""
    licna_karta: str = ""
    pasos: str = ""  # broj pasoša (za strance)
    is_stranac: bool = False  # da li je lice strani državljanin
    vrsta_isprave: str = "jmbg"  # jmbg / licna_karta / pasos — šta koristiti u pisanim ponudama
    adresa: str = ""
    grad: str = ""
    pozicija: str = ""
    strucna_sprema: str = ""
    plata_bruto: float = 0.0
    plata_neto: float = 0.0
    datum_pocetka: str = ""
    datum_kraja: str = ""
    datum_prestanka: str = ""  # datum prestanka radnog odnosa (odjava)
    vrsta_ugovora: str = "neodredjeno"  # odredjeno/neodredjeno
    radno_vrijeme: str = "puno"  # puno/skraceno
    sati_sedmicno: int = 40  # 40 = puno radno vrijeme, npr. 20 = pola
    telefon: str = ""
    email: str = ""
    aktivan: bool = True
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class EmployeeCreate(BaseModel):
    company_id: str
    objekat_id: Optional[str] = ""
    ime: str
    prezime: str
    jmbg: Optional[str] = ""
    licna_karta: Optional[str] = ""
    pasos: Optional[str] = ""
    is_stranac: Optional[bool] = False
    vrsta_isprave: Optional[str] = "jmbg"
    adresa: Optional[str] = ""
    grad: Optional[str] = ""
    pozicija: Optional[str] = ""
    strucna_sprema: Optional[str] = ""
    plata_bruto: Optional[float] = 0.0
    plata_neto: Optional[float] = 0.0
    datum_pocetka: Optional[str] = ""
    datum_kraja: Optional[str] = ""
    datum_prestanka: Optional[str] = ""
    vrsta_ugovora: Optional[str] = "neodredjeno"
    radno_vrijeme: Optional[str] = "puno"
    sati_sedmicno: Optional[int] = 40
    telefon: Optional[str] = ""
    email: Optional[str] = ""
    aktivan: Optional[bool] = True

class DocumentGenerateRequest(BaseModel):
    template_filename: str
    company_id: str
    employee_id: Optional[str] = None
    custom_fields: Dict[str, Any] = {}  # dodatna polja koja korisnik ručno popunjava (može sadržati liste i brojeve)


class AneksRequest(BaseModel):
    employee_id: str
    nova_vrsta_ugovora: str = "neodredjeno"  # neodredjeno/odredjeno
    novi_datum_kraja: Optional[str] = ""
    nova_plata_neto: Optional[float] = None
    nova_pozicija: Optional[str] = ""
    razlog: Optional[str] = ""
    update_employee: bool = True  # ažurirati i polja zaposlenog u bazi


class PunomoceRequest(BaseModel):
    """Podaci za generisanje Specijalnog punomoćja za osnivanje firme."""
    # Davalac punomoćja (osoba koja daje ovlaštenje)
    davaoc_ime_prezime: str  # "ARJANA CEKOVIQ"
    davaoc_is_stranac: bool = False
    davaoc_jmb: str = ""  # ako domaći
    davaoc_pasos: str = ""  # ako stranac
    davaoc_drzava: str = "Crne Gore"
    davaoc_adresa: str = ""
    
    # Firma koja se osniva
    firma_naziv: str = ""
    
    # Punomoćnik (npr. računovođa)
    punomocnik_ime_prezime: str = ""
    punomocnik_jmb: str = ""
    punomocnik_adresa: str = ""
    
    # Datum
    datum: str = ""  # default = today

# ============================================================================
# EVIDENCIJA RADA (Work Log) — integrisana evidencija svih radnih aktivnosti
# za svaku firmu po kategorijama.
# ============================================================================

WORK_KATEGORIJE = ["osnivanje", "pdv", "ioppd", "m4", "stvarni_vlasnici", "ostalo"]
WORK_STATUSI = ["u_toku", "poslato", "zavrseno"]


class WorkLogCreate(BaseModel):
    company_id: Optional[str] = None  # može biti "" za osnivanje firme prije nego što je u bazi
    company_naziv: str = ""  # ime firme (za prikaz, naročito za osnivanje)
    kategorija: str  # osnivanje / pdv / ioppd / m4 / stvarni_vlasnici / ostalo
    status: str = "u_toku"  # u_toku / poslato / zavrseno
    period: str = ""  # npr. "Maj 2026" ili "Q1 2026"
    napomena: str = ""
    iznos: Optional[float] = None  # opciono — npr. iznos PDV-a


class WorkLogUpdate(BaseModel):
    company_id: Optional[str] = None
    company_naziv: Optional[str] = None
    kategorija: Optional[str] = None
    status: Optional[str] = None
    period: Optional[str] = None
    napomena: Optional[str] = None
    iznos: Optional[float] = None


class FoundingRequest(BaseModel):
    """Podaci za generisanje 4 dokumenta osnivanja DOO firme."""
    # Osnivač
    osnivac_ime_prezime: str  # "ARJANA CEKOVIQ"
    osnivac_is_stranac: bool = False
    osnivac_jmb: str = ""  # ako domaći
    osnivac_pasos: str = ""  # ako stranac
    osnivac_drzava: str = "Crne Gore"  # "Crne Gore" ili "Albanija" itd.
    osnivac_adresa: str = ""  # "BRAJŠE BB.ULCINJ"
    osnivac_datum_rodjenja: str = ""  # za SAGLASNOST: "1985-12-20"
    osnivac_procenat: float = 100.0  # % udjela
    
    # Firma
    firma_naziv_pun: str  # "DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU ... \"ELA&ART\" ULCINJ"
    firma_naziv_skraceni: str  # "DOO \"ELA&ART\" ULCINJ"
    firma_naziv_pecat: str = ""  # "ELA&ART" (ide u pečat)
    firma_vrsta_djelatnosti_opis: str = "ZA PROIZVODNJU, PROMET I USLUGE"  # "ZA TRGOVINU" / "ZA UGOSTITELJSTVO" itd.
    firma_sjediste_adresa: str  # "BRAJŠE BB ULCINJ"
    firma_grad: str = "ULCINJ"
    firma_telefon: str = ""
    firma_email: str = ""
    firma_sifra_djelatnosti: str = "47.11"
    firma_naziv_djelatnosti: str = "Nespecijalizovana trgovina na malo pretežno hranom, pićima I duvanskim proizvodima"
    
    # Direktor (može biti isti kao osnivač)
    direktor_isti_kao_osnivac: bool = True
    direktor_ime_prezime: str = ""
    direktor_is_stranac: bool = False
    direktor_jmb: str = ""
    direktor_pasos: str = ""
    direktor_drzava: str = "Crne Gore"
    direktor_adresa: str = ""
    direktor_pol: str = "M"  # "M" = muško, "Z" = žensko (utiče na "saglasan/saglasna", "rodjen/rodjena")
    
    # Datumi
    datum_odluke: str = ""  # default = today
    osnovni_kapital: float = 1.00
    
    # Punomoćnik (Član 8.1 Odluke o osnivanju)
    podnosi_punomocnik: bool = False  # da li se podnosi preko punomoćnika
    punomocnik_ime_prezime: str = ""
    punomocnik_jmbg: str = ""
    punomocnik_adresa: str = ""


class PaymentRecord(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    year: int
    month: int  # 1-12
    iznos: float
    placeno: bool = False
    datum_placanja: str = ""
    napomena: str = ""

class PaymentUpdate(BaseModel):
    placeno: bool
    datum_placanja: Optional[str] = ""
    iznos: Optional[float] = None
    napomena: Optional[str] = ""

class PDVRecord(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    year: int
    month: int
    pdv_predato: bool = False
    pdv_datum: str = ""
    pdv_broj: str = ""
    pdv_status: str = "ceka"  # "ceka" | "u_toku" | "predato"
    ioppd_predato: bool = False
    ioppd_datum: str = ""
    ioppd_broj: str = ""
    ioppd_status: str = "ceka"

class PDVUpdate(BaseModel):
    pdv_predato: Optional[bool] = None
    pdv_datum: Optional[str] = None
    pdv_broj: Optional[str] = None
    pdv_status: Optional[str] = None
    ioppd_predato: Optional[bool] = None
    ioppd_datum: Optional[str] = None
    ioppd_broj: Optional[str] = None
    ioppd_status: Optional[str] = None


# ============== AUTH HELPERS ==============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
    except Exception:
        return False

def create_jwt(username: str) -> str:
    payload = {
        'sub': username,
        'exp': datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> str:
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        username = payload.get('sub')
        if not username:
            raise HTTPException(401, "Neispravan token")
        return username
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "Token istekao")
    except jwt.PyJWTError:
        raise HTTPException(401, "Neispravan token")


# ============== REMINDERS / PODSJETNICI ==============

@api_router.get("/reminders/expiring-contracts")
async def expiring_contracts(days: int = 30, username: str = Depends(get_current_user)):
    """Vraća sve zaposlene sa određenim ugovorom čiji datum_kraja ističe u narednih N dana."""
    today = datetime.now(timezone.utc).date()
    limit_date = today + timedelta(days=days)
    
    employees = await db.employees.find(
        {"vrsta_ugovora": "odredjeno", "datum_kraja": {"$ne": "", "$exists": True}, "aktivan": True},
        {"_id": 0}
    ).to_list(2000)
    
    expiring = []
    for emp in employees:
        try:
            end_str = emp.get("datum_kraja", "")
            if not end_str:
                continue
            end_dt = datetime.fromisoformat(end_str.replace('Z', '')).date()
            days_left = (end_dt - today).days
            if -7 <= days_left <= days:  # od 7 dana isteklog do N dana u budućnosti
                emp_copy = dict(emp)
                emp_copy["days_left"] = days_left
                emp_copy["end_date_formatted"] = end_dt.strftime("%d.%m.%Y")
                expiring.append(emp_copy)
        except Exception:
            continue
    
    # Enrich sa nazivima firmi
    company_ids = list(set(e.get("company_id") for e in expiring if e.get("company_id")))
    if company_ids:
        companies = await db.companies.find({"id": {"$in": company_ids}}, {"_id": 0}).to_list(1000)
        cmap = {c["id"]: c for c in companies}
        for e in expiring:
            c = cmap.get(e.get("company_id"))
            e["company_naziv"] = c.get("naziv", "—") if c else "—"
            e["company_pib"] = c.get("pib", "") if c else ""
    
    # Sort by days_left ascending (most urgent first)
    expiring.sort(key=lambda e: e.get("days_left", 999))
    return expiring


# ============== STARTUP - SEED USER + AGENCY ==============

@app.on_event("startup")
async def startup_seed():
    # Seed default user
    user = await db.users.find_one({"username": DEFAULT_USERNAME}, {"_id": 0})
    if not user:
        await db.users.insert_one({
            "username": DEFAULT_USERNAME,
            "password_hash": hash_password(DEFAULT_PASSWORD),
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        logging.info(f"Seeded default user: {DEFAULT_USERNAME}")
    
    # Seed default agency
    agency = await db.agency.find_one({}, {"_id": 0})
    if not agency:
        default_agency = Agency().model_dump()
        await db.agency.insert_one(default_agency)
        logging.info("Seeded default agency")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


# ============== AUTH ROUTES ==============

@api_router.post("/auth/login", response_model=LoginResponse)
async def login(data: LoginRequest):
    user = await db.users.find_one({"username": data.username}, {"_id": 0})
    if not user or not verify_password(data.password, user['password_hash']):
        raise HTTPException(401, "Pogrešno korisničko ime ili lozinka")
    token = create_jwt(data.username)
    return {"access_token": token, "user": {"username": data.username}}

@api_router.get("/auth/me")
async def me(username: str = Depends(get_current_user)):
    return {"username": username}


# ============== AGENCY ROUTES ==============

@api_router.get("/agency")
async def get_agency(username: str = Depends(get_current_user)):
    agency = await db.agency.find_one({}, {"_id": 0})
    if not agency:
        agency = Agency().model_dump()
        await db.agency.insert_one(agency)
    return agency

@api_router.put("/agency")
async def update_agency(data: Agency, username: str = Depends(get_current_user)):
    doc = data.model_dump()
    await db.agency.delete_many({})
    await db.agency.insert_one(dict(doc))
    return doc


# ============== IRMS LOOKUP ==============

@api_router.get("/companies/lookup-pib")
async def lookup_pib(pib: str, username: str = Depends(get_current_user)):
    """
    Automatski dohvat podataka firme iz IRMS portala Poreske uprave Crne Gore.
    Koristi javne API-je:
      1) GET /public/api/business-entities?identificationNumber={pib}  - search
      2) GET /public/api/business-entity/{taxpayerId}                  - detalji
      3) GET /public/api/business-entity/{taxpayerId}/ownership-roles  - direktori
    """
    pib = pib.strip()
    if not pib or not pib.isdigit() or len(pib) < 6:
        raise HTTPException(400, "PIB mora biti broj sa minimalno 6 cifara")
    
    result = {
        "success": False,
        "source": "IRMS",
        "pib": pib,
        "data": {},
        "message": ""
    }
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "application/json, text/plain, */*",
        "Accept-Language": "sr-Latn-ME,sr;q=0.9,en;q=0.8",
        "Origin": "https://irms.tax.gov.me",
        "Referer": "https://irms.tax.gov.me/public/search-register/business-entities",
    }
    
    try:
        # 1) Pretraga po PIB-u
        search_url = f"https://irms.tax.gov.me/public/api/business-entities?page=1&perPage=5&identificationNumber={pib}"
        search_resp = requests.get(search_url, headers=headers, timeout=10)
        if search_resp.status_code != 200:
            result["message"] = f"IRMS portal vratio status {search_resp.status_code}"
            return result
        
        search_data = search_resp.json()
        results_list = search_data.get("results", [])
        if not results_list:
            result["message"] = "Firma sa ovim PIB-om nije pronađena u IRMS registru"
            return result
        
        taxpayer = results_list[0]
        taxpayer_id = taxpayer.get("taxpayerId")
        
        if not taxpayer_id:
            result["message"] = "IRMS odgovor ne sadrži taxpayerId"
            return result
        
        # 2) Detalji
        detail_url = f"https://irms.tax.gov.me/public/api/business-entity/{taxpayer_id}"
        detail_resp = requests.get(detail_url, headers=headers, timeout=10)
        detail = detail_resp.json() if detail_resp.status_code == 200 else {}
        
        # 3) Ownership roles (za direktora)
        direktor_ime = ""
        try:
            roles_url = f"https://irms.tax.gov.me/public/api/business-entity/{taxpayer_id}/ownership-roles?id={taxpayer_id}&page=1&perPage=25"
            roles_resp = requests.get(roles_url, headers=headers, timeout=8)
            if roles_resp.status_code == 200:
                roles_data = roles_resp.json()
                for role in roles_data.get("results", []):
                    role_name = (role.get("role") or "").lower()
                    if "izvršni direktor" in role_name or "direktor" in role_name:
                        first = role.get("name") or ""
                        last = role.get("lastname") or ""
                        direktor_ime = f"{first} {last}".strip()
                        if direktor_ime:
                            break
                # Fallback: ako nije direktor, uzmi ovlašćenog zastupnika
                if not direktor_ime:
                    for role in roles_data.get("results", []):
                        first = role.get("name") or ""
                        last = role.get("lastname") or ""
                        direktor_ime = f"{first} {last}".strip()
                        if direktor_ime:
                            break
        except Exception as e:
            logging.warning(f"IRMS ownership-roles fetch failed: {e}")
        
        # Mapiranje na company schema
        full_name = detail.get("fullName") or taxpayer.get("fullName", "")
        short_name = detail.get("shortName") or ""
        registration_number = detail.get("registrationNumber") or taxpayer.get("registrationNumber", "")
        
        # Šifra djelatnosti: "5610, Djelatnosti restorana..." → ("5610", "Djelatnosti restorana...")
        main_activity = detail.get("mainActivity") or taxpayer.get("mainActivity", "")
        sifra_djelatnosti = ""
        djelatnost_naziv = main_activity
        if main_activity and "," in main_activity:
            parts = main_activity.split(",", 1)
            potential_code = parts[0].strip()
            if potential_code.isdigit():
                sifra_djelatnosti = potential_code
                djelatnost_naziv = parts[1].strip()
        
        # Grad: "Ulcinj, Crna Gora" → "Ulcinj"
        city_raw = detail.get("city") or ""
        grad = city_raw.split(",")[0].strip() if city_raw else ""
        
        # Datum registracije: "11/19/2025 00:00:00" → "2025-11-19"
        datum_registracije = ""
        reg_date_raw = detail.get("registrationDate") or ""
        if reg_date_raw:
            try:
                # Format: "11/19/2025 00:00:00" (MM/DD/YYYY)
                date_part = reg_date_raw.split()[0]
                m, d, y = date_part.split("/")
                datum_registracije = f"{y}-{int(m):02d}-{int(d):02d}"
            except Exception:
                pass
        
        result["success"] = True
        result["data"] = {
            "naziv": full_name,
            "naziv_skraceni": short_name,
            "maticni_broj": registration_number,
            "pib": detail.get("identificationNumber") or taxpayer.get("identificationNumber", pib),
            "adresa": detail.get("address") or "",
            "grad": grad,
            "email": detail.get("email") or "",
            "telefon": detail.get("phoneNumber") or "",
            "website": detail.get("website") or "",
            "djelatnost": djelatnost_naziv,
            "sifra_djelatnosti": sifra_djelatnosti,
            "direktor_ime": direktor_ime,
            "oblik_organizovanja": detail.get("legalStatus") or "",
            "datum_registracije": datum_registracije,
            "status": detail.get("taxpayerStatusDisplayName") or "",
        }
        result["message"] = "Podaci uspješno preuzeti sa IRMS portala"
        return result
    
    except requests.RequestException as e:
        logging.error(f"IRMS request failed: {e}")
        result["message"] = "IRMS portal trenutno nedostupan. Probajte ponovo kasnije."
        return result
    except Exception as e:
        logging.error(f"IRMS lookup error: {e}")
        result["message"] = f"Greška pri obradi IRMS odgovora: {str(e)[:120]}"
        return result


# ============== COMPANIES ROUTES ==============

@api_router.get("/companies")
@api_router.get("/companies")
async def list_companies(
    search: Optional[str] = None,
    pdv_only: bool = False,
    ioppd_only: bool = False,
    username: str = Depends(get_current_user)
):
    query: Dict[str, Any] = {}
    if pdv_only:
        query["pdv_obveznik"] = True
    if search:
        regex = re.escape(search)
        query["$or"] = [
            {"naziv": {"$regex": regex, "$options": "i"}},
            {"pib": {"$regex": regex, "$options": "i"}},
            {"direktor_ime": {"$regex": regex, "$options": "i"}},
        ]
    companies = await db.companies.find(query, {"_id": 0}).sort("naziv", 1).to_list(1000)
    
    # IOPPD filter: sve firme OSIM preduzetnika BEZ zaposlenih
    if ioppd_only:
        result = []
        for c in companies:
            oblik = (c.get("oblik_organizovanja") or "").lower()
            is_preduzetnik = "preduzetnik" in oblik
            if is_preduzetnik:
                # Provjeri da li ima zaposlenih
                emp_count = await db.employees.count_documents({"company_id": c["id"]})
                if emp_count == 0:
                    continue  # preskoči preduzetnika bez zaposlenih
            result.append(c)
        return result
    
    return companies

@api_router.get("/companies/irms-alerts")
async def irms_alerts(username: str = Depends(get_current_user)):
    """Vraća listu firmi čiji IRMS status NIJE 'Registrovan' (npr. u likvidaciji, stečaju, mirovanje)."""
    cursor = db.companies.find(
        {
            "irms_status": {"$exists": True, "$ne": "", "$nin": ["Registrovan", ""]}
        },
        {"_id": 0, "id": 1, "naziv": 1, "pib": 1, "irms_status": 1, "irms_checked_at": 1}
    )
    return [doc async for doc in cursor]


@api_router.get("/companies/{company_id}")
async def get_company(company_id: str, username: str = Depends(get_current_user)):
    c = await db.companies.find_one({"id": company_id}, {"_id": 0})
    if not c:
        raise HTTPException(404, "Firma nije pronađena")
    return c

@api_router.post("/companies")
async def create_company(data: CompanyCreate, username: str = Depends(get_current_user)):
    # Check duplicate PIB
    existing = await db.companies.find_one({"pib": data.pib}, {"_id": 0})
    if existing:
        raise HTTPException(400, f"Firma sa PIB-om {data.pib} već postoji")
    company = Company(**data.model_dump())
    await db.companies.insert_one(company.model_dump())
    return company.model_dump()

@api_router.put("/companies/{company_id}")
async def update_company(company_id: str, data: CompanyCreate, username: str = Depends(get_current_user)):
    existing = await db.companies.find_one({"id": company_id}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Firma nije pronađena")
    updated = {**existing, **data.model_dump()}
    await db.companies.update_one({"id": company_id}, {"$set": data.model_dump()})
    return updated

@api_router.delete("/companies/{company_id}")
async def delete_company(company_id: str, username: str = Depends(get_current_user)):
    await db.companies.delete_one({"id": company_id})
    await db.employees.delete_many({"company_id": company_id})
    return {"success": True}


# ============== EMPLOYEES ROUTES ==============

@api_router.get("/employees")
async def list_employees(
    company_id: Optional[str] = None,
    search: Optional[str] = None,
    username: str = Depends(get_current_user)
):
    query: Dict[str, Any] = {}
    if company_id:
        query["company_id"] = company_id
    if search:
        regex = re.escape(search)
        query["$or"] = [
            {"ime": {"$regex": regex, "$options": "i"}},
            {"prezime": {"$regex": regex, "$options": "i"}},
            {"jmbg": {"$regex": regex, "$options": "i"}},
        ]
    employees = await db.employees.find(query, {"_id": 0}).sort("prezime", 1).to_list(2000)
    # Pridruži naziv objekta
    obj_ids = list({(e.get("objekat_id") or "") for e in employees if e.get("objekat_id")})
    obj_map = {}
    if obj_ids:
        objs = await db.company_objekti.find({"id": {"$in": obj_ids}}, {"_id": 0, "id": 1, "naziv": 1}).to_list(500)
        obj_map = {o["id"]: o["naziv"] for o in objs}
    for e in employees:
        e["objekat_naziv"] = obj_map.get(e.get("objekat_id") or "", "")
    return employees

@api_router.get("/employees/{employee_id}")
async def get_employee(employee_id: str, username: str = Depends(get_current_user)):
    e = await db.employees.find_one({"id": employee_id}, {"_id": 0})
    if not e:
        raise HTTPException(404, "Zaposleni nije pronađen")
    return e

@api_router.post("/employees")
async def create_employee(data: EmployeeCreate, username: str = Depends(get_current_user)):
    # Validate company
    company = await db.companies.find_one({"id": data.company_id}, {"_id": 0})
    if not company:
        raise HTTPException(400, "Firma ne postoji")
    employee = Employee(**data.model_dump())
    await db.employees.insert_one(employee.model_dump())
    return employee.model_dump()

@api_router.put("/employees/{employee_id}")
async def update_employee(employee_id: str, data: EmployeeCreate, username: str = Depends(get_current_user)):
    existing = await db.employees.find_one({"id": employee_id}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Zaposleni nije pronađen")
    await db.employees.update_one({"id": employee_id}, {"$set": data.model_dump()})
    return {**existing, **data.model_dump()}

@api_router.delete("/employees/{employee_id}")
async def delete_employee(employee_id: str, username: str = Depends(get_current_user)):
    await db.employees.delete_one({"id": employee_id})
    return {"success": True}


# ============== PERSONS (centralna evidencija fizičkih lica) ==============

@api_router.get("/persons")
async def list_all_persons(
    search: Optional[str] = None,
    company_id: Optional[str] = None,
    username: str = Depends(get_current_user)
):
    """Vraća SVE zaposlene iz SVIH firmi sa pridruženim podacima o firmi."""
    query: Dict[str, Any] = {}
    if company_id:
        query["company_id"] = company_id
    if search:
        regex = re.escape(search)
        query["$or"] = [
            {"ime": {"$regex": regex, "$options": "i"}},
            {"prezime": {"$regex": regex, "$options": "i"}},
            {"jmbg": {"$regex": regex, "$options": "i"}},
            {"pozicija": {"$regex": regex, "$options": "i"}},
        ]
    
    employees = await db.employees.find(query, {"_id": 0}).sort("prezime", 1).to_list(5000)
    
    # Enrich sa nazivima firmi
    company_ids = list(set(e.get("company_id") for e in employees if e.get("company_id")))
    companies = await db.companies.find({"id": {"$in": company_ids}}, {"_id": 0}).to_list(1000)
    cmap = {c["id"]: c for c in companies}
    
    for e in employees:
        c = cmap.get(e.get("company_id"))
        e["company_naziv"] = c.get("naziv", "—") if c else "—"
        e["company_pib"] = c.get("pib", "") if c else ""
    
    return employees


# ============== TEMPLATES & DOCUMENT GENERATION ==============

@api_router.get("/templates")
async def list_templates(username: str = Depends(get_current_user)):
    """Vraća listu svih .docx, .doc, .pdf šablona u templates folderu."""
    from pdf_form_filler import is_pdf_form_template
    templates = []
    for f in sorted(TEMPLATES_DIR.iterdir()):
        if f.is_file() and f.suffix.lower() in ['.docx', '.doc', '.pdf', '.rtf']:
            # Friendly name
            name = f.stem.replace('_', ' ').strip()
            # Capitalize words
            name = ' '.join(w.capitalize() if not w.isupper() else w for w in name.split())
            category = _categorize_template(f.name)
            ext = f.suffix.lower()
            supports_gen = ext == '.docx' or (ext == '.pdf' and is_pdf_form_template(f.name))
            templates.append({
                "filename": f.name,
                "name": name,
                "extension": ext,
                "category": category,
                "supports_generation": supports_gen,
                "is_pdf_form": ext == '.pdf' and is_pdf_form_template(f.name),
            })
    return templates


def _categorize_template(filename: str) -> str:
    fl = filename.lower()
    if "ugovor" in fl:
        return "Ugovori"
    if "odluk" in fl:
        return "Odluke"
    if "obavjest" in fl or "obavest" in fl:
        return "Obavještenja"
    if "zahtjev" in fl or "zahtev" in fl:
        return "Zahtjevi"
    if "ovlasc" in fl or "punomoc" in fl or "ovlast" in fl:
        return "Ovlaštenja i punomoći"
    if "rjesenje" in fl or "rješenj" in fl:
        return "Rješenja"
    if "prijava" in fl:
        return "Prijave"
    if "izjava" in fl:
        return "Izjave"
    if "potvrda" in fl:
        return "Potvrde"
    if "saglasn" in fl:
        return "Saglasnosti"
    if "normativ" in fl:
        return "Normative"
    if "opoziv" in fl:
        return "Opozivi"
    if "obrazlozen" in fl:
        return "Obrazloženja"
    if "pisana ponuda" in fl:
        return "Pisane ponude"
    if "op obrazac" in fl:
        return "Obrasci"
    return "Ostalo"


def _format_date(date_str: str = "") -> str:
    """Formatira datum u DD.MM.YYYY"""
    if date_str:
        try:
            dt = datetime.fromisoformat(date_str.replace('Z', ''))
            return dt.strftime("%d.%m.%Y")
        except Exception:
            pass
    return datetime.now(timezone.utc).strftime("%d.%m.%Y")


# ============== SAMPLE VALUES MAPPING ==============
# Šabloni od korisnika sadrže konkretne podatke nekoliko "primjer" firmi.
# Kada se generiše dokument, sve te sample vrijednosti se zamijenjuju
# podacima izabrane firme/zaposlenog/agencije.

# Sample firme koje se pojavljuju u šablonima (sve se mapiraju na izabranu firmu):
SAMPLE_COMPANY_NAMES = [
    # Različite varijante kako se firma "CULT ULCINJ" pojavljuje u šablonima
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "CULT ULCINJ" ZA TRGOVINU, UGOSTITELJSTVO I USLUGE EXPORT-IMPORT- ULCINJ',
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "CULT ULCINJ" ZA TRGOVINU, UGOSTITELJSTVO I USLUGE "EXPORT-IMPORT" ULCINJ',
    'DOO "CULT ULCINJ"-ULCINJ',
    "DOO CULT ULCINJ-ULCINJ",
    "DOO CULT ULCINJ",
    "CULT ULCINJ",
    # MARINI GROUP
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "MARINI GROUP" ULCINJ',
    "MARINI GROUP",
    # SUMA FRESH MARKET
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "SUMA FRESH MARKET" - ULCINJ',
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU  "SUMA FRESH MARKET" ULCINJ',
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "SUMA FRESH MARKET" ULCINJ',
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "SUMA FRESH MARKET" -ULCINJ',
    "SUMA FRESH MARKET",
    # GONI COMPANY
    'DOO "GONI COMPANY" ZA PROIZVODNJU, PROMET I USLUGE, EXPORT - IMPORT ULCINJ',
    "GONI COMPANY",
    # SUR AFRODITA
    'PREDUZETNIK SINANI  ALIRAMI KOJI OBAVLJA  PRIVREDNU DJELATNOST  "SUR AFRODITA" ULCINJ',
    "SUR AFRODITA",
    # DARTI
    '"DARTI" D.O.O. ULCINJ',
    "DARTI",
    # FRIENDS CAFFE
    "DOO FRIENDS CAFFE ULCINJ",
    "FRIENDS CAFFE",
    # ELA&ART (OP OBRAZAC) - even though OP OBRAZAC is image-only, mapping is here for safety
    'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU ZA PROIZVODNJU, PROMET I USLUGE " ELA&ART " ULCINJ',
    # NOTE: NIKADA ne dodavaj split-ovane varijante poput "DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU ZA PROIZVODNJU,"
    # ili "PROMET I USLUGE..." kao posebne keys — prilikom zamjene sa pun nazivom (koji već sadrži ove dijelove)
    # nastaje duplikacija: "PUN_NAZIV PROMET I USLUGE ...".
    "ELA&ART",
]

# Sample PIBs koje treba zamijeniti (sa izabranom firmom)
SAMPLE_PIBS = ["03801969", "03796841", "03807851", "03663108", "03314367", "1906972223002"]

# Sample PDV brojevi
SAMPLE_PDV_NUMBERS = ["82/31-02356-8", "82/31-03288-7"]

# Sample telefoni
SAMPLE_PHONES = ["069832886", "069688102", "069628880"]

# Sample šifre djelatnosti
SAMPLE_DJELATNOST_CODES = ["4334", "5610", "4711"]

# Sample direktori (zamjenjuju se sa company.direktor_ime)
SAMPLE_DIRECTORS = [
    "JUSUF ELEZAGIĆ", "JUSUF ELEZAGIC",
    "ARIAN MARINI",
    "VESEL SUMA",
    "ALIRAMI SINANI", "SINANI ALIRAMI",
    "ZIJA DODIĆ", "ZIJA DODIC",
]

# Sample JMBG direktora
SAMPLE_DIRECTOR_JMBGS = [
    "0303987220166",
    "3105998220014",
    "3004985220014",
    "3004974220012",
]

# Sample adrese firmi
SAMPLE_COMPANY_ADDRESSES = [
    "Ulcinj, Ul.Vellezerit Frasheri bb.",
    "UL.VELLEZERIT FRASHERI BB ULCINJ",
    "UL.VLLEZERIT FRASHERI BB ULCINJ",
    "UL.VELLEZERIT FRASHERI BB",
    "VELLEZERIT FRASHERI BB",
    "VELIKA PLAZA BB - ULCINJ",
    "VLADIMIR BB. ULCINJ 85366",
    "VLADIMIR BB",
    "VLADIMIR   bb",
    "VLADIMIR BB ULCINJ",
    "DJERANE BB",
]

# Sample matični/registracijski brojevi
SAMPLE_REG_NUMBERS = ["5-1354657/001", "5-1344978"]

# Sample agency data (UVIJEK se zamjenjuju agencijskim podacima iz postavki)
# Agencija u šablonima može biti "Getuard Cekoviq" ili "Advanced Accounting" (D.O.O.)
SAMPLE_AGENCY_NAMES = [
    "D.O.O. ADVANCED ACCOUNTING- ULCINJ",
    "D.O.O. ADVANCED ACCOUNTING - ULCINJ",
    "DOO ADVANCED ACCOUNTING ULCINJ",
    "ADVANCED ACCOUNTING ULCINJ",
]
SAMPLE_AGENCY_PIBS = ["03719073"]
SAMPLE_AGENCY_DIRECTOR_NAMES = [
    "CEKOVIQ GETUARD", "GETUARD CEKOVIQ",
    "MIRSADA CEKOVIC", "MIRSADA CEKOVIQ",
]
SAMPLE_AGENCY_JMBGS = ["0806994223008", "2603972228013"]
SAMPLE_AGENCY_ADDRESSES = ["Ulcinja, Brajša bb.", "Brajša bb."]

# Sample zaposleni (zamjenjuju se izabranim zaposlenim)
SAMPLE_EMPLOYEE_NAMES = [
    "ALEKSANDER CUROVIĆ", "ALEKSANDER CUROVIC",
    "RENATO JAKU",
    "ZIJA DODIĆ", "ZIJA DODIC",
    "ALBERT OSMANOVIC", "ALBERT OSMANOVIĆ",
    "DAUT VELIC", "DAUT VELIĆ",
]
SAMPLE_EMPLOYEE_JMBGS = [
    "1411008223029", "039066621", "3004974220012",
    "0612986223008",
    "2602956223056",
]
SAMPLE_EMPLOYEE_LK = ["I3382349M"]  # broj lične karte
SAMPLE_EMPLOYEE_POSITIONS = [
    "KONOBAR", "Konobar", "konobar", "KONOBR",
    "KUVAR", "Kuvar", "kuvar",
    "pomoćni radnik u gradjevinu",
    "pomocni radnik u gradjevinu",
    "NK – nekvalifikovani radnik",
    "NK - nekvalifikovani radnik",
]

# Sample datumi početka rada koji se zamjenjuju sa emp.datum_pocetka
# (datum zasnivanja radnog odnosa, datum stupanja na rad, datum zaključenja ugovora)
SAMPLE_EMPLOYEE_START_DATES = [
    "09.02.2026",  # u UGOVOR O RADU Zaposlenih.docx
    "19.11.2025",  # u UGOVOR O RADU DIREKTOR.docx
    "01.02.2026",  # u UGOVOR O DOPUNSKOM RADU.docx (3x)
]

# Specifični datumi/periodi/dani koji se BRIŠU iz dokumenata
# (klijent ručno popuni u Wordu/PDF-u prema svojim potrebama)
SAMPLE_PERIODS_TO_BLANK = {
    # Periodi godišnjeg/sedmičnog odmora i pauze - briše se da klijent sam popuni
    "01.01.2026-31.12.2026": "________________________",
    "01.10.2026-31.10.2026": "________________________",
    "01.01.2026.godine": "_________ godine",
    # Specifični datumi koji NISU referenca na zakon
    "02.02.2026 god.": "______________ god.",
    # Specifična vremena pauze
    "10:00-10:30h": "____________",
    "10:00-10:30": "____________",
    # Specifični dan sedmičnog odmora - klijent bira po zaposlenom
    "NEDELJA": "____________",
    # Smjene u rasporedu (specifična satnica)
    "07:00  do 15:00  h": "______ do ______ h",
    "15:00  do_24:00  h": "______ do ______ h",
    "_07:00  do 15:00": "______ do ______",
    "_07:00": "______",
    "do_24:00": "do ______",
}


def _build_replacements(company: dict, employee: Optional[dict], agency: dict, custom: dict, template_filename: str = "") -> Dict[str, str]:
    """Gradi dictionary svih mogućih placeholdera za zamjenu."""
    today = datetime.now(timezone.utc)
    today_str = today.strftime("%d.%m.%Y")
    
    company_naziv = company.get("naziv", "")
    company_pib = company.get("pib", "")
    company_grad = company.get("grad", "") or "Ulcinj"
    company_adresa = company.get("adresa", "")
    full_adresa = f"{company_adresa}, {company_grad}" if company_adresa and company_grad else (company_adresa or company_grad)
    
    direktor_ime = company.get("direktor_ime", "") or "________________"
    direktor_jmbg = company.get("direktor_jmbg", "") or "________________"
    
    agency_director = agency.get("direktor_ime", "")
    agency_jmbg = agency.get("direktor_jmbg", "")
    agency_adresa = f"{agency.get('adresa','')}".strip() or "Brajša bb."
    
    emp_full = ""
    emp_jmbg = ""
    emp_pozicija = ""
    emp_grad = ""
    if employee:
        emp_full = f"{employee.get('ime','')} {employee.get('prezime','')}".strip()
        emp_jmbg = employee.get("jmbg", "")
        emp_pozicija = employee.get("pozicija", "")
        emp_grad = employee.get("grad", "") or "Ulcinju"
    
    repl = {}
    
    # ========== SMART REPLACEMENT - Sample → Real ==========
    # Sample company names → real company
    # Skraćeni naziv firme za kratke sample varijante
    naziv_skraceni = (company.get("naziv_skraceni") or "").strip() or company_naziv
    # Heuristika: kratke sample varijante (≤30 chars, npr. "DOO CULT ULCINJ" ili "CULT ULCINJ")
    # treba zamijeniti SKRAĆENIM nazivom firme — inače u tekstu poput "iz Ulcinja,donosi"
    # dobijamo cijeli pravni naziv unutar rečenice.
    for sample_name in SAMPLE_COMPANY_NAMES:
        if sample_name and company_naziv:
            if len(sample_name) <= 30 and naziv_skraceni:
                repl[sample_name] = naziv_skraceni
            else:
                repl[sample_name] = company_naziv
    
    # Sample PIBs → real PIB
    for sample_pib in SAMPLE_PIBS:
        if company_pib:
            repl[sample_pib] = company_pib
    
    # Sample directors → real director
    for sample_dir in SAMPLE_DIRECTORS:
        if direktor_ime and direktor_ime != "________________":
            repl[sample_dir] = direktor_ime
    
    # Razdvojeno ime/prezime direktora za templote koji ih traže odvojeno
    # (npr. Zahtjev iz kaznene evidencije fizičko lice)
    if direktor_ime and direktor_ime != "________________":
        parts = direktor_ime.strip().split(maxsplit=1)
        dir_first = parts[0] if parts else ""
        dir_last = parts[1] if len(parts) > 1 else ""
        if dir_first and dir_last:
            # "VESEL" (ime) i "SUMA" (prezime) standalone → koristi pravo ime/prezime
            # Pazi: SAMPLE_DIRECTORS sortira po dužini desc, "VESEL SUMA" se zamjenjuje PRIJE "VESEL"
            repl["VESEL"] = dir_first
            repl["SUMA"] = dir_last
    
    # Sample director JMBGs → real
    for sample_jmbg in SAMPLE_DIRECTOR_JMBGS:
        if direktor_jmbg and direktor_jmbg != "________________":
            repl[sample_jmbg] = direktor_jmbg
    
    # Sample addresses → real
    for sample_addr in SAMPLE_COMPANY_ADDRESSES:
        if company_adresa:
            repl[sample_addr] = full_adresa
    
    # Sample reg numbers + datum rješenja CRPS
    for sample_reg in SAMPLE_REG_NUMBERS:
        if company.get("maticni_broj"):
            repl[sample_reg] = company["maticni_broj"]
    # Datum rješenja CRPS - ako firma ima, koristi, inače blank
    repl["17.12.2025"] = company.get("datum_registracije") or "____________"
    # Datum CRPS rješenja iz Prijave zanatstva: "5-1344978  23.10.2025"
    repl["23.10.2025"] = company.get("datum_registracije") or "____________"
    
    # Šifra djelatnosti - sample "4711" iz Prijave trgovine + "4334", "5610" iz drugih
    if company.get("sifra_djelatnosti"):
        for sd_sample in SAMPLE_DJELATNOST_CODES:
            repl[sd_sample] = company["sifra_djelatnosti"]
    
    # Naziv djelatnosti
    if company.get("djelatnost"):
        repl["Bojenje i zastakljivanje"] = company["djelatnost"]
        repl["Trgovina na malo u nespecijalizovanim prodavnicama,pretežno hranom,pićem i duvanom"] = company["djelatnost"]
    
    # Žiro račun - sample iz Prijave trgovine i Prijave zanatstva
    if company.get("ziro_racun"):
        repl["535-26292-64"] = company["ziro_racun"]
        repl["530-797915-34"] = company["ziro_racun"]
    
    # PDV broj
    if company.get("pdv_broj"):
        for sample_pdv in SAMPLE_PDV_NUMBERS:
            repl[sample_pdv] = company["pdv_broj"]
    
    # Telefon - prvo company, ako nema onda agency
    company_tel = company.get("telefon") or agency.get("telefon", "")
    if company_tel:
        repl["+382 69 172 204"] = company_tel
        repl["069 172 204"] = company_tel
        # Sample telefoni iz različitih šablona
        for sample_phone in SAMPLE_PHONES:
            repl[sample_phone] = company_tel
    
    # Email - prvo company, ako nema onda agency
    company_email = company.get("email") or agency.get("email", "")
    if company_email:
        repl["Advanced.acct@hotmail.com"] = company_email
        repl["advanced.acct@hotmail.com"] = company_email
        repl["restauranculto@gmail.com"] = company_email
    
    # Datum rođenja direktora - ako firma ima, ostaje 30.04.1985 za korisnika da popuni ručno
    # (Nemamo to polje u modelu trenutno)
    repl["30.04.1985"] = "____________"
    
    # Ime oca/majke direktora - nema u DB, ostavi blank
    repl["QAMIL"] = "________________"
    repl["NADIRE"] = "________________"
    
    # Mjesto/opština rođenja direktora - nema u DB
    # "BAR" se javlja samo u kontekstu mjesta/opštine rođenja u Zahtjevu iz kaznene evidencije
    # ne diraj ostale upotrebe BAR. (Ostavi za korisnika.)
    
    # Naziv objekta (samo prvi varijant - tačan tekst iz šablona)
    if company.get("naziv_skraceni") or company.get("naziv"):
        obj_short = company.get("naziv_skraceni") or company.get("naziv", "")
        repl["LOUNGE BAR CULT"] = obj_short
    
    # Tipografski tipfeleri iz konverzije PDF→DOCX (UCLINJ → ULCINJ)
    if company_grad:
        repl["UCLINJ"] = company_grad
    
    # Agency name + PIB + director - always replace samples with current agency data
    agency_naziv = agency.get("naziv", "")
    agency_pib = agency.get("pib", "")
    for sample_ag_name in SAMPLE_AGENCY_NAMES:
        if agency_naziv:
            repl[sample_ag_name] = agency_naziv
    for sample_ag_pib in SAMPLE_AGENCY_PIBS:
        if agency_pib:
            repl[sample_ag_pib] = agency_pib
    for sample_ag in SAMPLE_AGENCY_DIRECTOR_NAMES:
        if agency_director:
            repl[sample_ag] = agency_director
    for sample_ag_jmbg in SAMPLE_AGENCY_JMBGS:
        if agency_jmbg:
            repl[sample_ag_jmbg] = agency_jmbg
    for sample_ag_addr in SAMPLE_AGENCY_ADDRESSES:
        if agency_adresa:
            repl[sample_ag_addr] = agency_adresa
    
    # Employee samples → real employee (only if selected)
    if employee and emp_full:
        for sample_emp in SAMPLE_EMPLOYEE_NAMES:
            repl[sample_emp] = emp_full
        for sample_emp_jmbg in SAMPLE_EMPLOYEE_JMBGS:
            if emp_jmbg:
                repl[sample_emp_jmbg] = emp_jmbg
        if emp_pozicija:
            for sample_pos in SAMPLE_EMPLOYEE_POSITIONS:
                # Don't replace DIREKTOR with employee position (DIREKTOR refers to company director)
                if sample_pos.upper() != "DIREKTOR":
                    repl[sample_pos] = emp_pozicija
        # Broj lične karte zaposlenog
        emp_lk = employee.get("licna_karta", "")
        if emp_lk:
            for sample_lk in SAMPLE_EMPLOYEE_LK:
                repl[sample_lk] = emp_lk
        # Datum zasnivanja radnog odnosa / stupanja na rad / zaključenja ugovora
        # → uzima se iz emp.datum_pocetka (kad je upisan u formi)
        emp_start = employee.get("datum_pocetka", "")
        emp_end = employee.get("datum_kraja", "")
        vrsta = (employee.get("vrsta_ugovora") or "neodredjeno").lower()
        
        formatted_start = ""
        formatted_end = ""
        if emp_start:
            try:
                dt = datetime.fromisoformat(emp_start.replace('Z', ''))
                formatted_start = dt.strftime("%d.%m.%Y")
            except Exception:
                formatted_start = emp_start
        if emp_end:
            try:
                dt = datetime.fromisoformat(emp_end.replace('Z', ''))
                formatted_end = dt.strftime("%d.%m.%Y")
            except Exception:
                formatted_end = emp_end
        
        # Konstruktor dinamičke fraze "određeno/neodređeno vrijeme rada i to od ..."
        if formatted_start:
            if vrsta == "odredjeno":
                end_part = formatted_end if formatted_end else "____________"
                contract_phrase = f"određeno vrijeme rada i to od  {formatted_start} - {end_part} godine"
            else:
                contract_phrase = f"neodređeno vrijeme rada i to od  {formatted_start} godine"
            
            # Zamjeni obje varijante iz šablona sa dinamičkom frazom
            repl["određeno vrijeme rada i to od  09.02.2026 -31.12.2026 godine"] = contract_phrase
            repl["neodređeno vrijeme rada i to od  19.11.2025 godine"] = contract_phrase
            
            # Dopunski rad: "period od 01.02.2026  i važi do 31.12.2026"
            if vrsta == "odredjeno":
                end_part2 = formatted_end if formatted_end else "____________"
                repl["period od 01.02.2026  i važi do 31.12.2026"] = f"period od {formatted_start}  i važi do {end_part2}"
            else:
                repl["period od 01.02.2026  i važi do 31.12.2026"] = f"period od {formatted_start} (na neodređeno vrijeme)"
        
        # Pojedinačni datumi zaposlenog (datum stupanja, datum zaključenja itd.)
        if formatted_start:
            for sample_date in SAMPLE_EMPLOYEE_START_DATES:
                repl[sample_date] = formatted_start
        # End-of-contract period blank/datum_kraja
        if formatted_end:
            repl["-31.12.2026 godine"] = f"- {formatted_end} godine"
            repl["važi do 31.12.2026"] = f"važi do {formatted_end}"
        else:
            repl["-31.12.2026 godine"] = "- ____________ godine"
            repl["važi do 31.12.2026"] = "važi do ____________"
        
        # Plata zaposlenog → uzima se iz forme (emp.plata_neto)
        emp_plata_neto = employee.get("plata_neto") or 0
        if emp_plata_neto > 0:
            plata_str = f"{float(emp_plata_neto):.2f}"
            # UGOVOR O RADU Zaposlenih/Direktor: "neto iznosu od 600.00 euro"
            repl["neto iznosu od 600.00 euro"] = f"neto iznosu od {plata_str} euro"
            repl["600.00 euro"] = f"{plata_str} euro"
            # UGOVOR O DOPUNSKOM RADU: "iznos od 300.00 eura"
            repl["iznos od 300.00 eura"] = f"iznos od {plata_str} eura"
            repl["300.00 eura"] = f"{plata_str} eura"
        
        # Datum prestanka radnog odnosa (odjava)
        emp_prestanak = employee.get("datum_prestanka", "")
        formatted_prestanak = ""
        if emp_prestanak:
            try:
                dt = datetime.fromisoformat(emp_prestanak.replace('Z', ''))
                formatted_prestanak = dt.strftime("%d.%m.%Y")
            except Exception:
                formatted_prestanak = emp_prestanak
            repl["[DATUM_PRESTANKA]"] = formatted_prestanak
            # Datum prestanka u rješenju o prestanku i obrazloženju za kašnjenje odjave
            # Sample datumi koji predstavljaju prestanak radnog odnosa
            repl["[DATUM_ODJAVE]"] = formatted_prestanak
        
        # POJEDINAČNO OBAVJESTENJE / ostali blank zaposlenog placeholderi
        for blank_zap in [
            "Zaposleni:________________________",
            "Zaposleni: ________________________",
            "Zaposlenom  __________________",
            "Zaposlenom __________________",
        ]:
            repl[blank_zap] = f"Zaposleni: {emp_full}"
        # Izjava o pravima/obavezama: "Ime i prezime: __________"
        for blank_imepr in [
            "Ime i prezime: ________________________",
            "Ime i prezime:________________________",
            "Ime i prezime: ____________________",
            "Ime i prezime:____________________",
        ]:
            repl[blank_imepr] = f"Ime i prezime: {emp_full}"
        
        if emp_jmbg:
            for blank_jmbg in [
                "JMB: _____________________",
                "JMB:_____________________",
                "JMBG: _____________________",
                "JMBG:  __________________",
                "JMBG: __________________",
            ]:
                repl[blank_jmbg] = f"JMBG: {emp_jmbg}"
        
        if emp_pozicija:
            for blank_pos in [
                "Radno mjesto: ____________________",
                "Radno mjesto:____________________",
                "radno mjesto:__________",
                "radno mjesto: __________",
                "Radno mjesto: ___________",
                "Radno mjesto:___________",
            ]:
                repl[blank_pos] = f"Radno mjesto: {emp_pozicija}"
            # Dodatne pozicije iz RJESENJE O PRESTANKU
            repl["pomocni radnik"] = emp_pozicija.lower() if emp_pozicija else "pomocni radnik"
        
        # Datum stupanja na rad - IZJAVA o pravima i obavezama
        if formatted_start:
            for blank_ds in [
                "Datum stupanja na rad: _______________",
                "Datum stupanja na rad:_______________",
            ]:
                repl[blank_ds] = f"Datum stupanja na rad: {formatted_start}"
        
        # Datum prestanka radnog odnosa - sample datumi u Rješenju o prestanku
        # ako je upisan datum_prestanka, koristi taj; inače današnji datum
        prestanak_date = formatted_prestanak if formatted_prestanak else today_str
        # Sample datumi prestanka u šablonima
        repl["31.03.2026"] = prestanak_date
        repl["28.02.2026"] = prestanak_date
    
    # Brisanje specifičnih datuma/periode/dana - klijent ručno popuni
    for sample_period, blank in SAMPLE_PERIODS_TO_BLANK.items():
        repl[sample_period] = blank
    
    # Header datumi u dokumentima ("Ulcinj, 01.01.2026 god.") → grad firme + današnji datum
    header_city = company_grad or "Ulcinj"
    repl["Ulcinj, 01.01.2026 god."] = f"{header_city}, {today_str} god."
    repl["Ulcinj, 02.02.2026 god."] = f"{header_city}, {today_str} god."
    repl["Ulcinj, 02.02.2026"] = f"{header_city}, {today_str}"
    repl["Ulcinj, 16.02.2026"] = f"{header_city}, {today_str}"
    repl["Ulcinj, 16.02.2026 godine"] = f"{header_city}, {today_str} godine"
    repl["Ulcinj, 02.04.2026 god."] = f"{header_city}, {today_str} god."
    repl["Ulcinj, 02.04.2026"] = f"{header_city}, {today_str}"
    repl["U ULCINJ, dana 02.06.2026 godine."] = f"U {header_city}, dana {today_str} godine."
    repl["U ULCINJ, dana 02.06.2026 godine"] = f"U {header_city}, dana {today_str} godine"
    repl["02.06.2026"] = today_str
    repl["U VLADIMIR, dana 01.03.2026  godine"] = f"U {header_city}, dana {today_str} godine"
    repl["U VLADIMIR, dana 01.03.2026 godine"] = f"U {header_city}, dana {today_str} godine"
    repl["dana 01.03.2026"] = f"dana {today_str}"
    # Standalone datumi koji su sample za današnji datum
    repl[" 01.03.2026 god."] = f" {today_str} god."
    repl[" 16.02.2026"] = f" {today_str}"
    repl["04.05.2026"] = today_str  # punomocje token
    
    # ========== MSG 292: Datumi koji se uvijek setuju na današnji (datum štampe) ==========
    # OP OBRAZAC: "PODGORICA, 18.05.2026 godine" → "PODGORICA, [danas] godine"
    repl["PODGORICA, 18.05.2026 godine"] = f"PODGORICA, {today_str} godine"
    # OP OBRAZAC: "UP I 02/02-057/2026-34002/2" → blank (broj rješenja)
    repl["UP I 02/02-057/2026-34002/2"] = "________________"
    # Prijava zanatstva: "31.10.2025" → današnji (datum prijave i ispunjenosti uslova)
    repl["31.10.2025"] = today_str
    # Prijava zanatstva: "Br./Nr.: 08- 306" → "Br./Nr.: 08- ___"
    repl["Br./Nr.: 08- 306"] = "Br./Nr.: 08- _____"
    # Odluka o blagajničkom maksimumu: "dana 02.02.2026 donosi" → "dana [danas] donosi"
    repl["dana 02.02.2026 donosi"] = f"dana {today_str} donosi"
    # Odluka za popust u prodavnicu: "dana 21.05.2026 god. donos" → "dana [danas] god. donosi"
    # Note: ovo je već pokriveno kroz "dana 01.03.2026" mapiranje gore, ali za svaki slučaj:
    repl["dana 01.01.2026 god."] = f"dana {today_str} god."
    repl["dana 02.02.2026 god."] = f"dana {today_str} god."
    repl["dana 16.02.2026 god."] = f"dana {today_str} god."
    
    # Broj broja dokumenta - ostavi blank za korisnika
    repl["BR:10/2026"] = "BR: ___/2026"
    
    # ========== BLANK PLACEHOLDERS (underscore lines) ==========
    # NAPOMENA: Ovi generički blank-line replacements se NE primenjuju na šablone
    # koji koriste eksplicitne (NAZIV_FIRME) placeholdere (npr. Zahtjev za uzorkovanje)
    # — jer bi short-key replacement pokrio dio long-underscore linije i ostavio remnant.
    _tn_lower_pre = (template_filename or "").lower()
    skip_legacy_blanks = "uzorkovanje" in _tn_lower_pre
    
    if company_naziv and not skip_legacy_blanks:
        for blank_naziv in [
            "NAZIV FIRME:____________________________",
            "NAZIV FIRME: ____________________________",
            "Naziv firme:_________________",
            "Naziv firme: _________________",
        ]:
            repl[blank_naziv] = f"NAZIV FIRME: {company_naziv}"
    
    if company_pib and not skip_legacy_blanks:
        for blank_pib in [
            "PIB: _______________________",
            "PIB:_______________________",
            "PIB: ______________________",
        ]:
            repl[blank_pib] = f"PIB: {company_pib}"
    
    if company_adresa and not skip_legacy_blanks:
        full_adr = f"{company_adresa}, {company_grad}" if company_adresa and company_grad else (company_adresa or company_grad)
        for blank_adr in [
            "Adresa: _____________________",
            "Adresa:_____________________",
            "Adresa: Ulcinj",
        ]:
            repl[blank_adr] = f"Adresa: {full_adr}"
    
    # Datum današnji za blank polje "Datum: ___"
    for blank_date in [
        "Datum: ________________",
        "Datum:________________",
        "Datum: _____________________",
        "Datum:_____________________",
    ]:
        repl[blank_date] = f"Datum: {today_str}"
    
    # Broj dokumenta - blank ostaje (korisnik popuni)
    # Ali sample iz odluke o popustu i sl. ostaje
    
    # Žiro račun u ODLUCI O PODIZANJU NOVCA
    if company.get("ziro_racun"):
        repl["sa žiro računa  društva  broj:_______________"] = f"sa žiro računa društva broj: {company['ziro_racun']}"
        repl["broj:_______________"] = f"broj: {company['ziro_racun']}"
    
    # ========== ALSO support [PLACEHOLDER] syntax for future templates ==========
    repl.update({
        # FIRMA - klijent
        "[NAZIV_FIRME]": company.get("naziv", ""),
        "[NAZIV_FIRME_SKRACENO]": company.get("naziv_skraceni") or company.get("naziv", ""),
        "[PIB_FIRME]": company.get("pib", ""),
        "[JIB_FIRME]": company.get("pib", ""),
        "[MATICNI_BROJ_FIRME]": company.get("maticni_broj", ""),
        "[PDV_BROJ_FIRME]": company.get("pdv_broj", ""),
        "[ADRESA_FIRME]": company.get("adresa", ""),
        "[GRAD_FIRME]": company.get("grad", ""),
        "[DJELATNOST_FIRME]": company.get("djelatnost", ""),
        "[SIFRA_DJELATNOSTI]": company.get("sifra_djelatnosti", ""),
        "[DIREKTOR_IME]": company.get("direktor_ime", ""),
        "[DIREKTOR_JMBG]": company.get("direktor_jmbg", ""),
        "[DIREKTOR_ADRESA]": company.get("direktor_adresa", "") or company.get("adresa", ""),
        "[ZIRO_RACUN_FIRME]": company.get("ziro_racun", ""),
        "[BANKA_FIRME]": company.get("banka", ""),
        "[TELEFON_FIRME]": company.get("telefon", ""),
        "[EMAIL_FIRME]": company.get("email", ""),
        
        # AGENCIJA
        "[NAZIV_AGENCIJE]": agency.get("naziv", ""),
        "[ADRESA_AGENCIJE]": agency.get("adresa", ""),
        "[GRAD_AGENCIJE]": agency.get("grad", ""),
        "[PIB_AGENCIJE]": agency.get("pib", ""),
        "[DIREKTOR_AGENCIJE]": agency.get("direktor_ime", ""),
        "[JMBG_AGENCIJE]": agency.get("direktor_jmbg", ""),
        
        # DATUMI
        "[DATUM_DANAS]": today.strftime("%d.%m.%Y"),
        "[DATUM]": today.strftime("%d.%m.%Y"),
        "[GODINA]": str(today.year),
        "[MJESEC]": str(today.month),
        "[DAN]": str(today.day),
    })
    
    if employee:
        repl.update({
            "[IME_RADNIKA]": employee.get("ime", ""),
            "[PREZIME_RADNIKA]": employee.get("prezime", ""),
            "[IME_PREZIME_RADNIKA]": f"{employee.get('ime','')} {employee.get('prezime','')}".strip(),
            "[JMBG_RADNIKA]": employee.get("jmbg", ""),
            "[LICNA_KARTA_RADNIKA]": employee.get("licna_karta", ""),
            "[ADRESA_RADNIKA]": employee.get("adresa", ""),
            "[GRAD_RADNIKA]": employee.get("grad", ""),
            "[POZICIJA_RADNIKA]": employee.get("pozicija", ""),
            "[STRUCNA_SPREMA]": employee.get("strucna_sprema", ""),
            "[PLATA_BRUTO]": f"{employee.get('plata_bruto', 0):.2f}",
            "[PLATA_NETO]": f"{employee.get('plata_neto', 0):.2f}",
            "[DATUM_POCETKA_RADA]": _format_date(employee.get("datum_pocetka", "")),
            "[VRSTA_UGOVORA]": employee.get("vrsta_ugovora", ""),
            "[RADNO_VRIJEME]": employee.get("radno_vrijeme", ""),
            "[TELEFON_RADNIKA]": employee.get("telefon", ""),
        })
        
        # UGOVOR O RADU — automatska zamjena broja sati i pun/nepuno
        # Šablon: "3. Zaposleni zasniva radni odnos sa __punim______ radnim vremenom u trajanju od  40  sati nedeljno."
        emp_sati = int(employee.get("sati_sedmicno") or 40)
        emp_vrste = (employee.get("radno_vrijeme") or "puno").lower()
        # Ako je <40, automatski je nepuno
        if emp_sati < 40 or emp_vrste in ("skraceno", "skraćeno", "nepuno"):
            radno_label = "nepunim"
        else:
            radno_label = "punim"
        # Originalni tekst iz šablona (sa underscore-ima)
        repl["sa __punim______ radnim vremenom u trajanju od  40  sati nedeljno."] = \
            f"sa {radno_label} radnim vremenom u trajanju od {emp_sati} sati nedeljno."
        # Varijante sa drugim brojem underscora
        repl["sa __punim_____ radnim vremenom u trajanju od  40  sati nedeljno."] = \
            f"sa {radno_label} radnim vremenom u trajanju od {emp_sati} sati nedeljno."
    
    # Custom fields od korisnika
    for k, v in custom.items():
        key = k if k.startswith("[") else f"[{k}]"
        repl[key] = str(v)
    
    # ========== MSG 292: Template-specifični prepisivi ==========
    tname_lower = (template_filename or "").lower()
    
    # "Rješenje o korišćenju godišnjeg odmora" → datum štampe ne treba (klijent popunjava ručno)
    if "rjesenje" in tname_lower and ("godisnj" in tname_lower or "godišnj" in tname_lower):
        # Resetuj sve današnje datume u headerima na blank
        repl[f"{header_city}, {today_str} god."] = f"{header_city}, ________________ god."
        repl[f"{header_city}, {today_str} godine"] = f"{header_city}, ________________ godine"
        repl[f"{header_city}, {today_str}"] = f"{header_city}, ________________"
        # "Datum: 21.05.2026" (već popunjeno iz blank field) → "Datum: ____"
        repl[f"Datum: {today_str}"] = "Datum: ________________"
        # Period korišćenja godišnjeg odmora — klijent popunjava
        repl["01.09.2026"] = "____________"
        repl["30.09.2026"] = "____________"
        # Broj radnih dana — klijent popunjava (sample je "2 2" sa space-om u templateu)
        repl["2 2 radn"] = "____ radn"
        repl["22 radn"] = "____ radn"
        # Broj rješenja "1/26" — klijent može ostaviti default ili popuniti
        # Ostavi kako jeste (može se kasnije prilagoditi)
        # Godina dvojno korišćenja "202 6" → tekuća godina
        repl["202 6 . godinu"] = f"{datetime.now().year} godinu"
        repl["2026. godinu"] = f"{datetime.now().year}. godinu"
    
    # "Rješenje o prestanku radnog odnosa kad ističe ugovor o radu" → datum štampe = today;
    # datum prestanka rada = employee.datum_prestanka (već mapirano kroz `prestanak_date`)
    # Ako je ovo termination rješenje a nema datum_prestanka, koristi datum_kraja
    if "prestan" in tname_lower and employee:
        emp_prestanak2 = employee.get("datum_prestanka") or employee.get("datum_kraja")
        if emp_prestanak2:
            try:
                dt = datetime.fromisoformat(emp_prestanak2.replace('Z',''))
                fp = dt.strftime("%d.%m.%Y")
                repl["31.03.2026"] = fp
                repl["28.02.2026"] = fp
            except Exception:
                pass
    
    # "Obrazloženje za poresku upravu kad kasnimo sa odjavama" → datum prestanka = employee.datum_prestanka
    if "obrazlozenje" in tname_lower and employee:
        emp_prestanak2 = employee.get("datum_prestanka") or employee.get("datum_kraja")
        if emp_prestanak2:
            try:
                dt = datetime.fromisoformat(emp_prestanak2.replace('Z',''))
                fp = dt.strftime("%d.%m.%Y")
                repl["31.03.2026"] = fp
                repl["28.02.2026"] = fp
            except Exception:
                pass
    
    # "Zahtjev iz kaznene evidencije za fizičko lice" → ako je employee_id naveden,
    # SUBJEKT je taj zaposleni (a NE direktor firme).
    if "kaznene evidencije fizicko" in tname_lower and employee and emp_full:
        emp_parts = emp_full.strip().split(maxsplit=1)
        emp_first = emp_parts[0] if emp_parts else ""
        emp_last = emp_parts[1] if len(emp_parts) > 1 else ""
        if emp_first and emp_last:
            # Override direktor mappings za ovaj specifičan šablon
            repl["VESEL"] = emp_first
            repl["SUMA"] = emp_last
            repl["VESEL SUMA"] = emp_full
        # Adresa zaposlenog umjesto direktora
        emp_adresa = employee.get("adresa", "")
        if emp_adresa:
            repl["VLADIMIR BB. ULCINJ 85366"] = emp_adresa
            repl["VLADIMIR BB"] = emp_adresa
        # JMBG zaposlenog
        if emp_jmbg:
            repl["3004985220014"] = emp_jmbg
    
    # "Prijava trgovine" — extras + datum logika
    if "prijava trgovine" in tname_lower or "prijava_trgovine" in tname_lower:
        # Tip prijave: pocetak / promjena
        tip = (custom.get("tip_prijave") or "pocetak").lower()
        # T0R0C1 sadrži "x" (čekirano "početak"). Ako promjena → ukloni X tu, dodaj kod promjene.
        # Pristup: zamijenimo postojeću "x" čekircu sa praznim, a u "promjena" liniji upišemo X.
        if tip == "promjena":
            # Ukloni X iz T0R0C1 ("-početak obavljanja trgovine 1)" → cell sa "x")
            # I dodaj X kod promjene; jednostavnije: dodaj '-promjena  podataka iz prijave2)' u replacement zajedno sa X-om
            # Pošto "x" se javlja i na drugim mjestima (kao 'X' velikim slovom je odvojeno), ovo radi targeted:
            repl["-početak obavljanja trgovine 1)"] = "-početak obavljanja trgovine 1)"  # ostavi
            repl["-promjena  podataka iz prijave2)"] = "-promjena  podataka iz prijave2)    X"
        
        # Sjedište + adresa objekta (T3R0C1, T3R0C3) — "Ulcinj" → user input, "VLADIMIR BB" → user input
        if custom.get("sjediste_objekta"):
            # Ovo je riskantno jer "Ulcinj" je svuda. Targetiramo cijelu ćeliju kombinaciju.
            pass  # ostavi za sad — Ulcinj se već zamjenjuje na druga mesta
        if custom.get("adresa_objekta"):
            repl["VLADIMIR BB"] = custom["adresa_objekta"]
        
        # Vrsta djelatnosti / "prodavnica" je već čekirana sa X u T4R0C2/C3 — ne treba mijenjati
        # Vrsta robe (T2R2C3 = "Trgovina na malo mješovitom robom")
        if custom.get("vrsta_djelatnosti"):
            repl["Trgovina na malo mješovitom robom"] = custom["vrsta_djelatnosti"]
        
        # Površina (T4R0C12 = "84")
        if custom.get("m2_poslovni") or custom.get("m2"):
            m2_val = str(custom.get("m2_poslovni") or custom.get("m2", ""))
            repl["84"] = m2_val
        
        # Datum početka rada (T5R0C2 = "01.02.2026")
        if custom.get("datum_pocetka_rada"):
            repl["01.02.2026"] = custom["datum_pocetka_rada"]
        # Datum podnošenja (T5R3C2 = "27.01.2026")
        repl["27.01.2026"] = today_str
        
        # Opis promjene
        if custom.get("opis_promjene"):
            # 6.Vrsta i opis promjene → ostavi placeholder
            pass
    
    # ============ CUSTOM FIELDS za druge odluke ============
    # 1) ODLUKA O POPUSTU — custom % popust
    if "popust" in tname_lower and custom.get("popust_procenat"):
        pct = str(custom["popust_procenat"]).replace(",", ".").strip()
        if not pct.endswith("%"):
            pct = f"{pct}%"
        repl["10%"] = pct
    
    # 2) OBAVJESTENJE O RADNOM VREMENU — user-input radno vrijeme
    if ("obavjestenje" in tname_lower or "obavještenje" in tname_lower) and "radn" in tname_lower:
        if custom.get("radno_vrijeme"):
            repl["08:00-12:00"] = custom["radno_vrijeme"]
    
    # 3) ODLUKA O RADNOM VREMENU - KOMUNALNA POLICIJA — radno vrijeme + dani
    if "komunaln" in tname_lower:
        rv = custom.get("radno_vrijeme", "")  # npr. "07:00 do 24:00"
        dani = custom.get("dani_rada", "")     # npr. "ponedeljak – nedelja"
        if rv:
            # Pojavljuje se 2 puta: "07:00 do 24:00 časova" + "od 07:00 do 24:00"
            # Format: razdvoji "od X do Y"
            repl["07:00 do 24:00"] = rv
        if dani:
            repl["ponedeljak – nedelja"] = dani
    
    # 4) ODLUKA ZA RAD TOKOM PRAZNIKA — custom lista praznika i godina
    if "praznika" in tname_lower or "prazni" in tname_lower:
        if custom.get("godina"):
            # "tokom 2026 god." → tokom {godina} god.
            repl["tokom 2026 god"] = f"tokom {custom['godina']} god"
            repl["2026"] = str(custom["godina"])
        # Note: lista praznika se popunjava preko paragraph editovanja u document, ne single replace
    

    # (Stari DOCX kod ispod je dead code — ostavljam za referencu ali se ne izvršava)
    if False and ("prijava_zanatstva" in tname_lower or "prijava zanatstva" in tname_lower):
        # 1.1. Naziv/ime/Emri _____...
        if company_naziv:
            repl["1.1. Naziv/ime/Emri ________________________________________________________________"] = \
                f"1.1. Naziv/ime/Emri  {company_naziv}"
        # Sjedište + adresa
        if company_adresa or company_grad:
            sj_adr = f"{company_grad or 'Ulcinj'}  adresa: {company_adresa}" if company_adresa else (company_grad or "")
            repl["a. Sjedište/Selia _____________________ adresa__________________________________________"] = \
                f"a. Sjedište/Selia  {company_grad}  adresa: {company_adresa or '____'}"
            repl["3.1. Sjedište/Selia________________________ adresa:______________________________________________"] = \
                f"3.1. Sjedište/Selia  {company_grad}  adresa: {company_adresa or '____'}"
        # 1.4. Šifra djelatnosti
        if company.get("sifra_djelatnosti"):
            repl["1.4. Šifra djelatnosti/Shifra e aktivitetit \t________________________________________"] = \
                f"1.4. Šifra djelatnosti/Shifra e aktivitetit \t{company['sifra_djelatnosti']}"
        # 1.5. Ime lica ovlašćenog za zastupanje (direktor)
        if direktor_ime and direktor_ime != "________________":
            repl["1.5.Ime lica ovlašćenog za zastupanje/Emri i personit të autorizuar për përfaqësim _____________________"] = \
                f"1.5.Ime lica ovlašćenog za zastupanje/Emri i personit të autorizuar për përfaqësim  {direktor_ime}"
        # 1.6. Žiro račun
        if company.get("ziro_racun"):
            repl["1.6.Žiro račun/i poslovna banka/Llogaria rrjedhëse dhe banka afariste ______________________________"] = \
                f"1.6.Žiro račun/i poslovna banka/Llogaria rrjedhëse dhe banka afariste  {company['ziro_racun']}"
        # 1.7. PIB
        if company_pib:
            repl["1.7. Poreski identifikacioni broj/Numri identifikues tatimor ______________________________________"] = \
                f"1.7. Poreski identifikacioni broj/Numri identifikues tatimor  {company_pib}"
        # 1.8. Telefon/email
        company_tel_val = company.get("telefon") or agency.get("telefon", "")
        if company_tel_val:
            repl["1.8.Telefon - i, fax - i, e-mail ______________________________________________________________"] = \
                f"1.8.Telefon - i, fax - i, e-mail  {company_tel_val}"
    
    # "Zahtjev za uzorkovanje i ispitivanje" (BRISEVA / HRANA / VODA ZA PIĆE) – novi šabloni sa eksplicitnim
    # placeholderima u zagradama (NAZIV_FIRME), (PIB_FIRME), (ADRESA_FIRME), itd.
    if "uzorkovanje" in tname_lower:
        company_tel_val2 = company.get("telefon") or agency.get("telefon", "")
        company_pdv_val = company.get("pdv_broj", "") or "____________"
        company_email_val = company.get("email") or agency.get("email", "") or "____________"
        company_sd_val = company.get("sifra_djelatnosti", "") or "____________"
        adresa_full = f"{company_adresa}, {company_grad}".strip(", ") if (company_adresa or company_grad) else "____________"
        naziv_objekta_val = company.get("naziv_skraceni") or company_naziv or "____________"
        
        # Ako je u custom poljima dat objekat, koristi ga umjesto sjedišta firme
        custom_objekat_naziv = (custom.get("naziv_objekta") or "").strip()
        custom_objekat_adresa = (custom.get("adresa_objekta") or "").strip()
        if custom_objekat_naziv:
            naziv_objekta_val = custom_objekat_naziv
        if custom_objekat_adresa:
            # Adresa firme se mijenja adresom objekta za ovaj dokument
            adresa_full = custom_objekat_adresa
        
        # Glavni placeholder mappings (case-sensitive, zagrade)
        repl["(NAZIV_FIRME)"] = company_naziv or "____________"
        repl["NAZIV_FIRME"] = company_naziv or "____________"  # bez zagrada (HRANA P8)
        repl["(ADRESA_FIRME)"] = adresa_full
        repl["(ADRESA FIRME)"] = adresa_full  # sa space-om (HRANA)
        repl["(PIB_FIRME)"] = company_pib or "____________"
        repl["(PDV_BROJ)"] = company_pdv_val
        repl["(BROJ_TELEFONA)"] = company_tel_val2 or "____________"
        repl["(IME_PREZIME_DIREKTORA)"] = direktor_ime if direktor_ime != "________________" else "____________"
        repl["(ime_prezime_direktora)"] = direktor_ime if direktor_ime != "________________" else "____________"
        repl["(DATUM_STAMPE)"] = today_str
        repl["(NAZIV_OBJEKTA)"] = naziv_objekta_val
        repl["(SIFRA_DJELATNOSTI)"] = company_sd_val
        repl["(email_adresa_firme)"] = company_email_val
        repl["(EMAIL_ADRESA_FIRME)"] = company_email_val
        repl["(ULCINJ)"] = company_grad or "Ulcinj"  # mjesto podnošenja
        
        # Čišćenje leftover hardcoded podataka iz originalnih PDF-ova:
        repl["DOO FRIENDS CAFFE ULCINJ "] = ""
        repl["DOO FRIENDS CAFFE ULCINJ"] = ""
        # HRANA P2 - "Podaci o objektu" placeholder je izgubljen pri konverziji
        if company_naziv:
            repl["Podaci o objektu (navesti tačan naziv i adresu)________________________________________ _____________________________________________________________________________"] = \
                f"Podaci o objektu: {company_naziv}, {adresa_full}".rstrip(", ")
    
    # UGOVOR O POZAJMICI — iznos pozajmice + datum sklapanja (klijent unosi) + žiro/banka
    if "pozajm" in tname_lower:
        # Datum sklapanja ugovora — iz custom polja (klijent ručno unosi).
        # Ako nije unesen, ostavlja crticu da klijent dopiše u Wordu.
        datum_raw = custom.get("datum_sklapanja") or ""
        datum_str = _format_date(str(datum_raw)) if datum_raw else "__________"
        repl["sklapaju dana 06.02.2026"] = f"sklapaju dana {datum_str}"
        repl["dana 06.02.2026"] = f"dana {datum_str}"
        repl["06.02.2026"] = datum_str
        
        # Iznos pozajmice iz custom polja zamjenjuje "__________" u Članu 1
        iznos_raw = custom.get("iznos_pozajmice")
        if iznos_raw not in (None, "", 0):
            try:
                iznos_num = float(str(iznos_raw).replace(",", ".").replace("€", "").strip())
                iznos_str = f"{iznos_num:,.2f}".replace(",", " ").replace(".", ",")
            except Exception:
                iznos_str = str(iznos_raw).strip()
            # Zamijeni dugu liniju "__________" sa iznosom (u Članu 1)
            repl["iznos od __________ €"] = f"iznos od {iznos_str} €"
            repl["iznos od __________"] = f"iznos od {iznos_str}"
            repl["__________ €"] = f"{iznos_str} €"
        
        # Član 7 — žiro-račun i banka iz baze firme.
        # Žiro: globalna zamjena "535-26292-64" → company.ziro_racun već radi (linija ~1012),
        # ali ako firma nema žiro/banku — ostavi crticu da klijent dopiše.
        company_ziro = (company.get("ziro_racun") or "").strip()
        company_banka = (company.get("banka") or "").strip()
        if not company_ziro:
            repl["žiro-račun 535-26292-64"] = "žiro-račun ____________________"
            repl["535-26292-64"] = "____________________"
        if company_banka:
            repl["kod  PRVA BANKA"] = f"kod {company_banka}"
            repl["kod PRVA BANKA"] = f"kod {company_banka}"
            repl["PRVA BANKA"] = company_banka
        else:
            repl["kod  PRVA BANKA"] = "kod ____________________"
            repl["kod PRVA BANKA"] = "kod ____________________"
            repl["PRVA BANKA"] = "____________________"
    
    # PISANA PONUDA ZA BORAVAK I RAD / PRODUZENJE
    if "pisana ponuda" in tname_lower:
        # Hardkodirana imena firmi i direktora u template-ima — zamijeni sa pravim
        # "Boravak i rad" template ima MARINI GROUP / ARIAN MARINI
        # "Produzenje" template ima UNICO COMPANY / EDONA BOLJEVIĆ
        sample_firme = [
            'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU "MARINI GROUP" ULCINJ',
            '"UNICO COMPANY" DOO ULCINJ ',
            '"UNICO COMPANY" DOO ULCINJ',
        ]
        sample_direktori = [
            ("ARIAN MARINI", "3105998220014"),
            ("EDONA BOLJEVIĆ ", "1402991228025"),
            ("EDONA BOLJEVIĆ", "1402991228025"),
        ]
        sample_pibovi = ["03796841", "03420205"]
        
        # Zamijeni naziv firme
        firma_target = company_naziv or "____________"
        for s in sample_firme:
            repl[s] = firma_target
        
        # Zamijeni PIB
        if company_pib:
            for p in sample_pibovi:
                repl[p] = company_pib
                repl[f"PIB:\t{p}"] = f"PIB:\t{company_pib}"
                repl[f"PIB:{p}"] = f"PIB:{company_pib}"
        
        # Zamijeni direktora
        if direktor_ime and direktor_ime != "________________":
            for d_name, d_jmb in sample_direktori:
                repl[d_name] = direktor_ime
            # JMB direktora — iz agency.direktor_jmbg ili prvi zaposlenik?
            ag_dir_jmb = agency.get("direktor_jmbg", "") or company.get("jmbg_direktora", "")
            if ag_dir_jmb:
                for d_name, d_jmb in sample_direktori:
                    repl[f"JMB {d_jmb}"] = f"JMB {ag_dir_jmb}"
                    repl[d_jmb] = ag_dir_jmb
        
        # 1) Broj zavedene ponude — generisan u generate_document funkciji i prosljeđen kroz custom
        broj_ponude = custom.get("broj_ponude") or ""
        if broj_ponude:
            repl["BROJ: 10/2026"] = f"BROJ: {broj_ponude}"
            repl["BROJ:11/2026"] = f"BROJ: {broj_ponude}"
            repl["BROJ: 11/2026"] = f"BROJ: {broj_ponude}"
            repl["BROJ:10/2026"] = f"BROJ: {broj_ponude}"
        
        # 2) Datum štampe ponude (gore u headeru i u P4 "je dana XX.XX.XXXX godine donio odluku")
        datum_ponude_raw = custom.get("datum_ponude") or ""
        if datum_ponude_raw:
            datum_ponude_str = _format_date(str(datum_ponude_raw))
        else:
            datum_ponude_str = today_str
        # Header datumi
        repl["Ulcinj, 04.03.2026godine"] = f"{header_city}, {datum_ponude_str} godine"
        repl["Ulcinj, 04.03.2026 godine"] = f"{header_city}, {datum_ponude_str} godine"
        repl["Ulcinj, 04.03.2026"] = f"{header_city}, {datum_ponude_str}"
        repl["Ulcinj, 13.05.2026"] = f"{header_city}, {datum_ponude_str}"
        # P4: "je dana 04.03.2026 godine, donio odluku"
        repl["je dana 04.03.2026 godine"] = f"je dana {datum_ponude_str} godine"
        repl["je dana 04.03.2026"] = f"je dana {datum_ponude_str}"
        repl["je dana 13.05.2026 ,godine"] = f"je dana {datum_ponude_str}, godine"
        repl["je dana 13.05.2026"] = f"je dana {datum_ponude_str}"
        repl["04.03.2026"] = datum_ponude_str
        repl["13.05.2026"] = datum_ponude_str
        
        # 3) Član 2 — radni odnos zasniva se / produžava se na određeno
        # "počev od 04.03.2026-03.03.2027.god."
        datum_od_raw = custom.get("datum_rad_od") or ""
        datum_do_raw = custom.get("datum_rad_do") or ""
        datum_od = _format_date(str(datum_od_raw)) if datum_od_raw else "__________"
        datum_do = _format_date(str(datum_do_raw)) if datum_do_raw else "__________"
        period_str = f"{datum_od}-{datum_do}.god."
        # Originalni stringovi u dva template-a
        repl["počev od  04.03.2026-03.03.2027.god."] = f"počev od {period_str}"
        repl["počev od 04.03.2026-03.03.2027.god."] = f"počev od {period_str}"
        repl["04.03.2026-03.03.2027.god."] = period_str
        repl["počev od 25.05.2026-24.05.2027.god."] = f"počev od {period_str}"
        repl["25.05.2026-24.05.2027.god."] = period_str
        
        # 4) Plata iz baze zaposlenog ako postoji
        if employee:
            emp_plata = employee.get("plata_neto") or 0
            if emp_plata > 0:
                plata_str = f"{float(emp_plata):.2f}"
                repl["mjesečna neto zarada u iznosu od 600.00e"] = f"mjesečna neto zarada u iznosu od {plata_str}e"
                repl["600.00e"] = f"{plata_str}e"
        
        # 5) JMBG / Broj isprave — uzima se na osnovu vrste_isprave i is_stranac
        # Stari hardkodovani brojevi koji se zamjenjuju
        if employee:
            vrsta_isprave = (employee.get("vrsta_isprave") or "jmbg").lower()
            is_stranac = bool(employee.get("is_stranac", False))
            
            # Logika izbora isprave
            if is_stranac:
                # Stranac: prioritet pasos > licna_karta > jmbg
                if vrsta_isprave == "pasos" and employee.get("pasos"):
                    broj_isprave = employee["pasos"]
                elif vrsta_isprave == "licna_karta" and employee.get("licna_karta"):
                    broj_isprave = employee["licna_karta"]
                elif employee.get("pasos"):
                    broj_isprave = employee["pasos"]
                elif employee.get("licna_karta"):
                    broj_isprave = employee["licna_karta"]
                else:
                    broj_isprave = employee.get("jmbg", "") or "____________"
            else:
                # Domaće lice: JMBG je standardno
                broj_isprave = employee.get("jmbg", "") or employee.get("licna_karta", "") or "____________"
            
            # Sample brojevi iz template-a (JMBG/BROJ ISPRAVE 039066621 ili 039437802)
            repl["JMBG/BROJ ISPRAVE 039066621"] = f"JMBG/BROJ ISPRAVE {broj_isprave}"
            repl["JMBG/BROJ ISPRAVE 039437802"] = f"JMBG/BROJ ISPRAVE {broj_isprave}"
            repl["039066621"] = broj_isprave
            repl["039437802"] = broj_isprave
            
            # Ime i prezime zaposlenog (RENATO JAKU / RENATO PALOKA)
            ime_prezime = f"{employee.get('ime','')} {employee.get('prezime','')}".strip().upper()
            if ime_prezime:
                repl["RENATO JAKU"] = ime_prezime
                repl["RENATO PALOKA"] = ime_prezime
            
            # Boravište
            boraviste = (employee.get("adresa") or "") + (", " + employee.get("grad", "") if employee.get("grad") else "")
            boraviste = boraviste.strip(", ") or "____________"
            repl["sa boravištem u ULCINJ ,u daljem"] = f"sa boravištem u {boraviste},u daljem"
            repl["sa boravištem u Ulcinju,u daljem"] = f"sa boravištem u {boraviste},u daljem"
            
            # Pozicija (radno mjesto)
            poz = employee.get("pozicija", "")
            if poz:
                repl["pomoćni radnik u gradjevinu"] = poz
                repl["pomocni radnik"] = poz
    
    return repl


def _fit_to_a4_one_page(doc: Document):
    """Postavlja A4 dimenzije, smanjuje margine i font da bi dokument stao u 1 stranicu.
    Koristi se za Obavještenje o knjizi prigovora (MSG 292).
    """
    from docx.shared import Cm, Pt
    # A4: 21.0 cm x 29.7 cm
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Smanji font veličinu za sve runove ako su veći od 11pt
    for para in doc.paragraphs:
        # Smanji prored
        try:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
        except Exception:
            pass
        for run in para.runs:
            try:
                cur = run.font.size
                if cur is None or cur > Pt(11):
                    run.font.size = Pt(10)
            except Exception:
                pass
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    try:
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(1)
                    except Exception:
                        pass
                    for run in para.runs:
                        try:
                            cur = run.font.size
                            if cur is None or cur > Pt(11):
                                run.font.size = Pt(10)
                        except Exception:
                            pass


def _populate_employees_table(doc, employees: list, tname_lower: str, custom: dict):
    """Popuni glavnu tabelu sa svim radnicima firme — za odluke koje imaju spisak."""
    if not doc.tables or not employees:
        return
    # Pretpostavka: prva tabela je spisak radnika (rb, ime, ...)
    tbl = doc.tables[0]
    rows = tbl.rows
    if len(rows) < 2:
        return
    # 1. red su zaglavlja, popuni od 2. reda
    headers = [c.text.strip().lower() for c in rows[0].cells]
    n_cols = len(rows[0].cells)
    
    # Defaults po template-u
    is_raspored = "rasporedu radnog" in tname_lower
    is_pauza = "pauze" in tname_lower
    is_sedmicni = "sedmic" in tname_lower or "sedmič" in tname_lower
    is_godisnji = "godisnji" in tname_lower or "godišnji" in tname_lower
    
    # Smjene/pauza/sedmični odmor — NE popunjavaj po defaultu.
    # Ostavi prazne ćelije da klijent sam popuni (po zahtjevu korisnika).
    smjena_default = custom.get("smjena_oznaka") or ""
    
    # Loop kroz radnike, ako tabela nema dovoljno redova — dodaj
    for i, emp in enumerate(employees):
        if i + 1 < len(rows):
            row = rows[i + 1]
        else:
            row = tbl.add_row()
        cells = row.cells
        ime_prezime = f"{emp.get('ime','')} {emp.get('prezime','')}".strip()
        pozicija = emp.get("pozicija", "")
        if n_cols >= 1:
            cells[0].text = f"{i+1}."
        if n_cols >= 2:
            cells[1].text = ime_prezime
        # 3. kolona — zavisi od template-a
        if n_cols >= 3:
            if is_raspored:
                cells[2].text = pozicija  # Radno mjesto
            elif is_pauza:
                # Klijent sam popunjava vrijeme pauze
                cells[2].text = custom.get("pauza_default") or ""
            elif is_sedmicni:
                # Klijent sam popunjava dan sedmičnog odmora
                cells[2].text = custom.get("sedmicni_default") or ""
            elif is_godisnji:
                cells[2].text = custom.get("godisnji_default") or ""
        # Za raspored radnog vremena — NE popunjavaj smjene po default-u.
        # Klijent sam popunjava ćelije po danima (I, II, X, ...).
        if is_raspored and n_cols >= 10:
            for d_idx in range(3, 10):
                cells[d_idx].text = smjena_default  # "" osim ako klijent unese override
        elif is_raspored and n_cols >= 4:
            # Manji broj kolona — staviti samo pozicija
            pass
    
    # Obriši višak praznih redova (osim ako ih je manje od potrebnog)
    extra_idx = len(employees) + 1
    while extra_idx < len(rows):
        row = rows[extra_idx]
        # Ako je red prazan (samo rb), ostavi praznim — ne brišemo
        extra_idx += 1
    
    # Postavi font 12pt za cijelu tabelu (zaglavlje + svi redovi)
    # po zahtjevu korisnika — za 4 Odluke (raspored / pauza / sedmični / godišnji).
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                runs = list(para.runs)
                if not runs:
                    # Prazna ćelija — dodaj prazan run sa 12pt da kad klijent kuca u Wordu, krene sa 12pt
                    r = para.add_run("")
                    try:
                        r.font.size = Pt(12)
                    except Exception:
                        pass
                else:
                    for run in runs:
                        try:
                            run.font.size = Pt(12)
                        except Exception:
                            pass


# Lista praznika Crne Gore (državni + vjerski)
CG_PRAZNICI = [
    {"datum": "01.01", "naziv": "Nova godina", "tip": "državni"},
    {"datum": "02.01", "naziv": "Nova godina (drugi dan)", "tip": "državni"},
    {"datum": "06.01", "naziv": "Badnji dan (pravoslavni)", "tip": "vjerski"},
    {"datum": "07.01", "naziv": "Božić (pravoslavni)", "tip": "vjerski"},
    {"datum": "08.01", "naziv": "Božić (pravoslavni, treći dan)", "tip": "vjerski"},
    {"datum": "01.05", "naziv": "Praznik rada", "tip": "državni"},
    {"datum": "02.05", "naziv": "Praznik rada (drugi dan)", "tip": "državni"},
    {"datum": "21.05", "naziv": "Dan nezavisnosti", "tip": "državni"},
    {"datum": "22.05", "naziv": "Dan nezavisnosti (drugi dan)", "tip": "državni"},
    {"datum": "13.07", "naziv": "Dan državnosti", "tip": "državni"},
    {"datum": "14.07", "naziv": "Dan državnosti (drugi dan)", "tip": "državni"},
    {"datum": "13.11", "naziv": "Njegošev dan", "tip": "državni"},
    {"datum": "14.11", "naziv": "Njegošev dan (drugi dan)", "tip": "državni"},
    {"datum": "25.12", "naziv": "Božić (katolički)", "tip": "vjerski"},
    {"datum": "26.12", "naziv": "Božić (katolički, drugi dan)", "tip": "vjerski"},
]

def _dan_nedjelje(d_str: str, year: int) -> str:
    """Vrati skraćeni naziv dana u sedmici (Pon, Uto, ...) za datum DD.MM godine."""
    try:
        dt = datetime.strptime(f"{d_str}.{year}", "%d.%m.%Y")
        dani = ["Pon", "Uto", "Sri", "Čet", "Pet", "Sub", "Ned"]
        return dani[dt.weekday()]
    except Exception:
        return ""


def _populate_praznici_lista(doc, praznici: list):
    """Zamijeni postojeću listu praznika u Odluci za rad tokom praznika sa custom listom.
    Praznici: lista dict-ova [{datum: '01.05.2026', naziv: 'Praznik rada', dan: 'Pet'}, ...]
    """
    # Pronađi paragrafe sa praznicima (sadrži "Praznik rada" ili "01. 05. 2026" pattern)
    target_paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Hard-coded references u sample template-u
        sample_indicators = ["Praznik rada", "Dan nezavisnosti", "Dan državnosti", "Njegošev dan", "Badnji dan", "Božić"]
        if any(ind in text for ind in sample_indicators) and len(text) < 100:
            target_paragraphs.append((i, p))
    
    # Obriši postojeće stare praznike (clear text)
    for idx, p in target_paragraphs:
        for run in p.runs:
            run.text = ""
    
    # Upiši nove praznike u prvi paragraf, ostatak postaje prazni
    if not target_paragraphs:
        return
    
    first_p = target_paragraphs[0][1]
    if not first_p.runs:
        return
    
    # Format: "Pet\t01. 05. 2026.\tPraznik rada"
    lines = []
    for pr in praznici:
        dan = pr.get("dan", "")
        datum = pr.get("datum", "")
        naziv = pr.get("naziv", "")
        if datum and naziv:
            lines.append(f"{dan}\t{datum}.\t{naziv}")
    
    if not lines:
        return
    
    # Prvi paragraf → prva linija, ostatak → preostalo, do limita
    for j, (idx, p) in enumerate(target_paragraphs):
        if j < len(lines):
            if p.runs:
                p.runs[0].text = lines[j]
        else:
            # Prazno — već je clear-ovano
            pass


def _docx_replace(doc: Document, replacements: Dict[str, str]):
    """Zamijenjuje sve placeholdere u dokumentu (paragrafima + tabelama)."""
    
    def replace_in_paragraph(paragraph):
        # Combine all runs text
        full_text = ''.join(run.text for run in paragraph.runs)
        new_text = full_text
        # Zamjene idu od najduže ka najkraćoj (da duži pattern ne bude pojeden kraćim).
        # KORISTIMO PRIVREMENI MARKER da bismo izbegli kaskadu kada zamjenska vrijednost
        # sadrži dijelove drugih ključeva (npr. "PROIZVODNJU," u novom nazivu firme).
        sorted_keys = sorted(replacements.keys(), key=len, reverse=True)
        markers: Dict[str, str] = {}
        for i, key in enumerate(sorted_keys):
            val = replacements[key]
            if key and key in new_text:
                marker = f"\x00REPL{i:04d}\x00"
                new_text = new_text.replace(key, marker)
                markers[marker] = val
        # Sada zamijeni markere sa stvarnim vrijednostima
        for marker, val in markers.items():
            new_text = new_text.replace(marker, val)
        
        if new_text != full_text and paragraph.runs:
            # Keep first run formatting, clear other runs
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
    
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            try:
                cells = list(row.cells)
            except Exception:
                # pdf2docx output može imati malformirane tc grid (vertical merge / missing tc).
                # Padaj na sirovo iteriranje raw <w:tc> elemenata.
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                cells = []
                from docx.table import _Cell
                for tc in row._tr.findall(f'{ns}tc'):
                    try:
                        cells.append(_Cell(tc, table))
                    except Exception:
                        pass
            for cell in cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    # Headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    replace_in_paragraph(paragraph)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    replace_in_paragraph(paragraph)


import subprocess

def _ensure_libreoffice_installed():
    """Self-healing: ako 'soffice' nije instaliran (može se desiti nakon container restart-a),
    pokreni apt-get install. Cache-uje rezultat da ne mlatimo apt na svaki poziv."""
    import shutil
    if shutil.which("soffice"):
        return True
    logging.warning("soffice not found — auto-installing LibreOffice…")
    try:
        subprocess.run(["apt-get", "install", "-y", "libreoffice-core", "libreoffice-writer"],
                       capture_output=True, timeout=240, check=False)
        return shutil.which("soffice") is not None
    except Exception as e:
        logging.error(f"LibreOffice auto-install failed: {e}")
        return False


def _convert_to_pdf(docx_path: Path) -> Optional[Path]:
    """Konvertuje docx u PDF koristeći LibreOffice headless. Vraća putanju do PDF-a ili None."""
    try:
        pdf_path = docx_path.with_suffix('.pdf')
        if pdf_path.exists():
            return pdf_path
        
        # Self-heal: ako LibreOffice fali, pokušaj da ga instaliraš
        if not _ensure_libreoffice_installed():
            logging.error("LibreOffice nije dostupan i auto-instalacija nije uspjela")
            return None
        
        # Run LibreOffice headless conversion
        result = subprocess.run(
            [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(docx_path.parent),
                str(docx_path)
            ],
            capture_output=True, timeout=60
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
        logging.warning(f"PDF conversion failed: {result.stderr.decode()[:200]}")
        return None
    except Exception as e:
        logging.error(f"PDF conversion error: {e}")
        return None


@api_router.post("/documents/generate")
async def generate_document(req: DocumentGenerateRequest, username: str = Depends(get_current_user)):
    template_path = TEMPLATES_DIR / req.template_filename
    if not template_path.exists():
        raise HTTPException(404, "Šablon nije pronađen")
    
    company = await db.companies.find_one({"id": req.company_id}, {"_id": 0})
    if not company:
        raise HTTPException(404, "Firma nije pronađena")
    
    employee = None
    if req.employee_id:
        employee = await db.employees.find_one({"id": req.employee_id}, {"_id": 0})
    
    agency = await db.agency.find_one({}, {"_id": 0}) or Agency().model_dump()
    
    # === PDF FORM OVERLAY (za nove BRISEVA/HRANA/VODA/BAZENI/Prijava zanatstva) ===
    from pdf_form_filler import is_pdf_form_template, fill_pdf_template
    if is_pdf_form_template(req.template_filename):
        output_filename = f"{uuid.uuid4().hex[:8]}_{template_path.stem}_{company.get('naziv_skraceni') or company.get('naziv','firma')[:20]}.pdf"
        output_filename = re.sub(r'[^\w\s.-]', '_', output_filename).replace(' ', '_')
        output_path = GENERATED_DIR / output_filename
        ok = fill_pdf_template(req.template_filename, output_path, company, agency, req.custom_fields)
        if not ok:
            raise HTTPException(500, "Greška pri popunjavanju PDF šablona")
        record = {
            "id": str(uuid.uuid4()),
            "filename": output_filename,
            "pdf_filename": output_filename,
            "template_filename": req.template_filename,
            "template": req.template_filename,
            "company_id": req.company_id,
            "company_naziv": company.get("naziv", ""),
            "employee_id": req.employee_id,
            "custom_fields": req.custom_fields or {},
            "created_by": username,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.generated_documents.insert_one(dict(record))
        return {"success": True, "filename": output_filename, "pdf_filename": output_filename, "record": {k: v for k, v in record.items() if k != "_id"}}
    
    if template_path.suffix.lower() != '.docx':
        raise HTTPException(400, "Trenutno se podržava samo .docx generisanje")
    
    # PISANA PONUDA — auto-inkrement brojača po firmi (osim ako je broj eksplicitno prosljeđen)
    custom_fields_in = dict(req.custom_fields or {})
    if "pisana ponuda" in req.template_filename.lower():
        if not custom_fields_in.get("broj_ponude"):
            # Atomski inkrement counter-a u companies kolekciji
            now_year = datetime.now(timezone.utc).year
            counter_field = f"ponuda_counter_{now_year}"
            updated = await db.companies.find_one_and_update(
                {"id": req.company_id},
                {"$inc": {counter_field: 1}},
                return_document=True,
                projection={"_id": 0, counter_field: 1},
            )
            new_num = (updated or {}).get(counter_field, 1)
            custom_fields_in["broj_ponude"] = f"{new_num:02d}/{now_year}"
    
    replacements = _build_replacements(company, employee, agency, custom_fields_in, req.template_filename)
    
    # Load template, replace, save
    doc = Document(str(template_path))
    _docx_replace(doc, replacements)
    
    # POST-PROCESSING: dodatne template-specific transformacije
    tname_lower = req.template_filename.lower()
    custom = req.custom_fields or {}
    
    # Tabela svih zaposlenih za odluke (rasporedu/sedmični/godišnji/pauza)
    is_table_template = any(k in tname_lower for k in [
        "rasporedu radnog", "sedmicnog odmora", "sedmičnog odmora",
        "godisnji odmor", "godišnji odmor", "koriscenje pauze", "korišćenje pauze",
    ])
    if is_table_template:
        emp_query: Dict[str, Any] = {"company_id": req.company_id}
        # Ako je u custom poljima poslata lista zaposlenih (npr. filtrirana po objektu), koristi samo njih
        table_ids = custom.get("table_employee_ids") or custom.get("bulk_employee_ids")
        if table_ids and isinstance(table_ids, list) and len(table_ids) > 0:
            emp_query["id"] = {"$in": table_ids}
        emp_list = await db.employees.find(emp_query, {"_id": 0}).to_list(length=500)
        # Sortiraj radnike po imenu
        emp_list.sort(key=lambda e: (e.get("ime", "") + " " + e.get("prezime", "")).strip().upper())
        _populate_employees_table(doc, emp_list, tname_lower, custom)
    
    # Lista praznika u Odluka za rad tokom praznika
    if "praznika" in tname_lower or "prazni" in tname_lower:
        praznici_lista = custom.get("praznici_lista") or []
        if praznici_lista:
            _populate_praznici_lista(doc, praznici_lista)
    
    # MSG 292: Formatiranje na 1 A4 stranicu za Obavještenje o knjizi prigovora
    if "knjige prigovora" in req.template_filename.lower() or "knjigu prigovora" in req.template_filename.lower() or "podnosenja prigovora" in req.template_filename.lower():
        _fit_to_a4_one_page(doc)
    
    # Save generated file
    output_filename = f"{uuid.uuid4().hex[:8]}_{template_path.stem}_{company.get('naziv_skraceni') or company.get('naziv','firma')[:20]}.docx"
    output_filename = re.sub(r'[^\w\s.-]', '_', output_filename).replace(' ', '_')
    output_path = GENERATED_DIR / output_filename
    doc.save(str(output_path))
    
    # Convert to PDF in background (it takes 2-3s)
    pdf_filename = output_filename.replace('.docx', '.pdf')
    _convert_to_pdf(output_path)  # Generates the PDF beside .docx
    
    # Save record
    record = {
        "id": str(uuid.uuid4()),
        "filename": output_filename,
        "pdf_filename": pdf_filename,
        "template": req.template_filename,
        "template_filename": req.template_filename,
        "company_id": req.company_id,
        "company_naziv": company.get("naziv", ""),
        "employee_id": req.employee_id,
        "employee_naziv": f"{employee.get('ime','')} {employee.get('prezime','')}".strip() if employee else "",
        "custom_fields": custom_fields_in or req.custom_fields or {},
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username
    }
    await db.generated_documents.insert_one(dict(record))
    
    return {
        "success": True,
        "filename": output_filename,
        "pdf_filename": pdf_filename,
        "download_url": f"/api/documents/download/{output_filename}",
        "preview_url": f"/api/documents/preview/{pdf_filename}",
        "record": {k: v for k, v in record.items() if k != "_id"}
    }


@api_router.get("/documents/history")
async def list_document_history(
    company_id: Optional[str] = None,
    template: Optional[str] = None,
    limit: int = 200,
    username: str = Depends(get_current_user)
):
    """Lista svih generisanih dokumenata, filter po firmi/šablonu."""
    query = {}
    if company_id:
        query["company_id"] = company_id
    if template:
        query["$or"] = [{"template": template}, {"template_filename": template}]
    
    docs = await db.generated_documents.find(query, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(length=limit)
    
    # Enrich sa firma nazivom ako fali u recordu (legacy zapisi)
    for d in docs:
        if not d.get("company_naziv") and d.get("company_id"):
            c = await db.companies.find_one({"id": d["company_id"]}, {"_id": 0, "naziv": 1, "naziv_skraceni": 1})
            if c:
                d["company_naziv"] = c.get("naziv_skraceni") or c.get("naziv", "")
        # Normalizuj template field
        if not d.get("template_filename"):
            d["template_filename"] = d.get("template", "")
    
    return docs


@api_router.get("/documents/history/{record_id}")
async def get_document_history_item(record_id: str, username: str = Depends(get_current_user)):
    """Dobij detalje (uključujući custom_fields) za jedan zapis."""
    rec = await db.generated_documents.find_one({"id": record_id}, {"_id": 0})
    if not rec:
        raise HTTPException(404, "Zapis nije pronađen")
    if not rec.get("template_filename"):
        rec["template_filename"] = rec.get("template", "")
    return rec


@api_router.delete("/documents/history/{record_id}")
async def delete_document_history_item(record_id: str, username: str = Depends(get_current_user)):
    """Obriši zapis (i pripadajući fajl ako postoji)."""
    rec = await db.generated_documents.find_one({"id": record_id}, {"_id": 0})
    if not rec:
        raise HTTPException(404, "Zapis nije pronađen")
    # Obriši fajlove
    for fn_key in ["filename", "pdf_filename"]:
        fn = rec.get(fn_key)
        if fn:
            fp = GENERATED_DIR / fn
            if fp.exists():
                try:
                    fp.unlink()
                except Exception:
                    pass
    await db.generated_documents.delete_one({"id": record_id})
    return {"success": True}


@api_router.get("/praznici/{year}")
async def list_praznici_godina(year: int, username: str = Depends(get_current_user)):
    """Vraća listu državnih i vjerskih praznika Crne Gore za datu godinu sa danom u sedmici."""
    result = []
    for p in CG_PRAZNICI:
        dan = _dan_nedjelje(p["datum"], year)
        result.append({
            "datum": f"{p['datum']}.{year}",
            "naziv": p["naziv"],
            "tip": p["tip"],
            "dan": dan,
        })
    return {"godina": year, "praznici": result}


@api_router.post("/documents/bulk-generate")
async def bulk_generate_documents(
    request: dict,
    username: str = Depends(get_current_user)
):
    """Generiše više dokumenata odjednom (jedan template × više zaposlenih) → ZIP fajl."""
    import zipfile
    
    template_filename = request.get("template_filename")
    company_id = request.get("company_id")
    employee_ids = request.get("employee_ids") or []
    custom_fields = request.get("custom_fields") or {}
    
    if not template_filename or not company_id or not employee_ids:
        raise HTTPException(400, "template_filename, company_id i employee_ids su obavezni")
    
    template_path = TEMPLATES_DIR / template_filename
    if not template_path.exists():
        raise HTTPException(404, "Šablon nije pronađen")
    
    company = await db.companies.find_one({"id": company_id}, {"_id": 0})
    if not company:
        raise HTTPException(404, "Firma nije pronađena")
    
    agency = await db.agency.find_one({}, {"_id": 0}) or Agency().model_dump()
    
    # Privremeni ZIP fajl
    zip_id = uuid.uuid4().hex[:8]
    safe_naziv = re.sub(r'[^\w-]', '_', (company.get("naziv_skraceni") or company.get("naziv","firma"))[:30])
    zip_filename = f"{zip_id}_{template_path.stem}_{safe_naziv}_BULK.zip"
    zip_filename = re.sub(r'[^\w.-]', '_', zip_filename)
    zip_path = GENERATED_DIR / zip_filename
    
    generated_count = 0
    failed = []
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for emp_id in employee_ids:
            employee = await db.employees.find_one({"id": emp_id}, {"_id": 0})
            if not employee:
                failed.append(f"emp_id {emp_id} not found")
                continue
            try:
                replacements = _build_replacements(company, employee, agency, custom_fields, template_filename)
                doc = Document(str(template_path))
                _docx_replace(doc, replacements)
                
                emp_safe = re.sub(r'[^\w-]', '_', f"{employee.get('ime','')}_{employee.get('prezime','')}")[:40]
                fname_in_zip = f"{template_path.stem}_{emp_safe}.docx"
                tmp_path = GENERATED_DIR / f"tmp_{uuid.uuid4().hex[:6]}.docx"
                doc.save(str(tmp_path))
                # Convert to PDF
                pdf_tmp = _convert_to_pdf(tmp_path)
                # Add both DOCX and PDF to ZIP
                zf.write(tmp_path, f"{template_path.stem}_{emp_safe}.docx")
                if pdf_tmp and Path(pdf_tmp).exists():
                    zf.write(pdf_tmp, f"{template_path.stem}_{emp_safe}.pdf")
                    Path(pdf_tmp).unlink(missing_ok=True)
                tmp_path.unlink(missing_ok=True)
                generated_count += 1
                
                # Zapiši u istoriju
                await db.generated_documents.insert_one({
                    "id": str(uuid.uuid4()),
                    "filename": fname_in_zip,
                    "pdf_filename": fname_in_zip.replace(".docx", ".pdf"),
                    "template": template_filename,
                    "template_filename": template_filename,
                    "company_id": company_id,
                    "company_naziv": company.get("naziv", ""),
                    "employee_id": emp_id,
                    "employee_naziv": f"{employee.get('ime','')} {employee.get('prezime','')}".strip(),
                    "custom_fields": custom_fields,
                    "bulk_zip": zip_filename,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                    "created_by": username,
                })
            except Exception as e:
                failed.append(f"emp {emp_id}: {str(e)[:80]}")
    
    return {
        "success": True,
        "generated_count": generated_count,
        "total_requested": len(employee_ids),
        "failed": failed,
        "zip_filename": zip_filename,
        "download_url": f"/api/documents/download/{zip_filename}",
    }


@api_router.get("/companies/{company_id}/objekti")
async def list_company_objects(company_id: str, username: str = Depends(get_current_user)):
    """Lista objekata (poslovnica/mjesta poslovanja) za firmu.
    Vraća prvo ručno-sačuvane objekte iz company_objekti kolekcije, zatim historijske
    nazive iz prijava (dedupe) kao dodatne suggestion-e."""
    # Ručno sačuvani objekti
    saved = await db.company_objekti.find({"company_id": company_id}, {"_id": 0}).sort("naziv", 1).to_list(200)
    
    seen = {}
    for o in saved:
        no = (o.get("naziv") or "").strip()
        if no:
            seen[no] = {
                "id": o.get("id"),
                "naziv_objekta": no,
                "adresa_objekta": (o.get("adresa") or "").strip(),
                "grad": o.get("grad", ""),
                "telefon": o.get("telefon", ""),
                "sifra_djelatnosti": o.get("sifra_djelatnosti", ""),
                "napomena": o.get("napomena", ""),
                "saved": True,
                "last_used": o.get("updated_at", ""),
            }
    
    # Historijski (iz prijava) — samo oni koji nisu već u listi
    docs = await db.generated_documents.find(
        {"company_id": company_id, "custom_fields.naziv_objekta": {"$exists": True, "$ne": ""}},
        {"_id": 0, "custom_fields": 1, "created_at": 1}
    ).sort("created_at", -1).limit(50).to_list(length=50)
    for d in docs:
        cf = d.get("custom_fields") or {}
        no = (cf.get("naziv_objekta") or "").strip()
        if no and no not in seen:
            seen[no] = {
                "naziv_objekta": no,
                "adresa_objekta": (cf.get("adresa_objekta") or "").strip(),
                "saved": False,
                "last_used": d.get("created_at", ""),
            }
    return list(seen.values())


class CompanyObjekat(BaseModel):
    naziv: str
    adresa: str = ""
    grad: str = ""
    telefon: str = ""
    sifra_djelatnosti: str = ""
    napomena: str = ""


@api_router.post("/companies/{company_id}/objekti")
async def create_company_objekat(company_id: str, req: CompanyObjekat, username: str = Depends(get_current_user)):
    # Provjeri da firma postoji
    co = await db.companies.find_one({"id": company_id}, {"_id": 0, "id": 1})
    if not co:
        raise HTTPException(404, "Firma nije pronađena.")
    if not req.naziv.strip():
        raise HTTPException(400, "Naziv objekta je obavezan.")
    rec = {
        "id": str(uuid.uuid4()),
        "company_id": company_id,
        "naziv": req.naziv.strip(),
        "adresa": req.adresa.strip(),
        "grad": req.grad.strip(),
        "telefon": req.telefon.strip(),
        "sifra_djelatnosti": req.sifra_djelatnosti.strip(),
        "napomena": req.napomena.strip(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.company_objekti.insert_one(dict(rec))
    return rec


@api_router.put("/companies/{company_id}/objekti/{objekat_id}")
async def update_company_objekat(company_id: str, objekat_id: str, req: CompanyObjekat, username: str = Depends(get_current_user)):
    updates = {
        "naziv": req.naziv.strip(),
        "adresa": req.adresa.strip(),
        "grad": req.grad.strip(),
        "telefon": req.telefon.strip(),
        "sifra_djelatnosti": req.sifra_djelatnosti.strip(),
        "napomena": req.napomena.strip(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    res = await db.company_objekti.update_one(
        {"id": objekat_id, "company_id": company_id},
        {"$set": updates},
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Objekat nije pronađen.")
    return {"id": objekat_id, **updates}


@api_router.delete("/companies/{company_id}/objekti/{objekat_id}")
async def delete_company_objekat(company_id: str, objekat_id: str, username: str = Depends(get_current_user)):
    res = await db.company_objekti.delete_one({"id": objekat_id, "company_id": company_id})
    if res.deleted_count == 0:
        raise HTTPException(404, "Objekat nije pronađen.")
    return {"success": True}


@api_router.post("/documents/generate-aneks")
async def generate_aneks(req: AneksRequest, username: str = Depends(get_current_user)):
    """Generiše aneks ugovora i opciono ažurira polja zaposlenog u bazi."""
    employee = await db.employees.find_one({"id": req.employee_id}, {"_id": 0})
    if not employee:
        raise HTTPException(404, "Zaposleni nije pronađen")
    
    company = await db.companies.find_one({"id": employee.get("company_id")}, {"_id": 0})
    if not company:
        raise HTTPException(404, "Firma nije pronađena")
    
    agency = await db.agency.find_one({}, {"_id": 0}) or Agency().model_dump()
    
    # Konstruiši opis izmjena u Članu 2
    izmjene_lines = []
    
    # Trajanje radnog odnosa
    today = datetime.now(timezone.utc).date()
    if req.nova_vrsta_ugovora == "neodredjeno":
        izmjene_lines.append(
            f"Radni odnos se mijenja u radni odnos NA NEODREĐENO VRIJEME, počev od {today.strftime('%d.%m.%Y')} godine."
        )
    elif req.nova_vrsta_ugovora == "odredjeno":
        end_str = "____________"
        if req.novi_datum_kraja:
            try:
                end_dt = datetime.fromisoformat(req.novi_datum_kraja.replace('Z',''))
                end_str = end_dt.strftime('%d.%m.%Y')
            except Exception:
                end_str = req.novi_datum_kraja
        izmjene_lines.append(
            f"Radni odnos se produžava NA ODREĐENO VRIJEME, počev od {today.strftime('%d.%m.%Y')} do {end_str} godine."
        )
    
    # Iznos zarade
    if req.nova_plata_neto is not None and req.nova_plata_neto > 0:
        stara = float(employee.get("plata_neto") or 0)
        izmjene_lines.append(
            f"Iznos neto zarade se mijenja sa {stara:.2f} eura na {float(req.nova_plata_neto):.2f} eura mjesečno."
        )
    
    # Radno mjesto
    if req.nova_pozicija:
        stara_poz = employee.get("pozicija", "")
        if stara_poz != req.nova_pozicija:
            izmjene_lines.append(
                f"Radno mjesto se mijenja iz \"{stara_poz}\" u \"{req.nova_pozicija}\"."
            )
    
    if req.razlog:
        izmjene_lines.append(f"Razlog izmjene: {req.razlog}")
    
    if not izmjene_lines:
        izmjene_lines.append("(Nije navedena izmjena — molimo popunite Aneks ručno u Wordu.)")
    
    izmjene_text = "\n".join(f"{i+1}. {line}" for i, line in enumerate(izmjene_lines))
    
    # Brojač aneksa
    aneks_count = await db.generated_documents.count_documents(
        {"template": "ANEKS UGOVORA O RADU.docx", "employee_id": req.employee_id}
    )
    aneks_broj = f"{aneks_count + 1}/{today.year}"
    
    # Build replacements
    replacements = _build_replacements(company, employee, agency, {}, "ANEKS UGOVORA O RADU.docx")
    replacements["[ANEKS_IZMJENE]"] = izmjene_text
    replacements["{ANEKS_BROJ}"] = aneks_broj
    
    # Adresa firme full
    adresa_full = f"{company.get('adresa','')}, {company.get('grad','')}".strip(", ")
    replacements["[ADRESA_FIRME_FULL]"] = adresa_full or company.get('adresa','') or company.get('grad','')
    
    # Generate
    template_path = TEMPLATES_DIR / "ANEKS UGOVORA O RADU.docx"
    doc = Document(str(template_path))
    _docx_replace(doc, replacements)
    
    output_filename = f"{uuid.uuid4().hex[:8]}_ANEKS_{employee.get('ime','')}_{employee.get('prezime','')}.docx"
    output_filename = re.sub(r'[^\w.-]', '_', output_filename)
    output_path = GENERATED_DIR / output_filename
    doc.save(str(output_path))
    
    pdf_filename = output_filename.replace('.docx', '.pdf')
    _convert_to_pdf(output_path)
    
    # Ažuriraj zaposlenog ako je traženo
    if req.update_employee:
        update = {"vrsta_ugovora": req.nova_vrsta_ugovora}
        if req.novi_datum_kraja is not None:
            update["datum_kraja"] = req.novi_datum_kraja if req.nova_vrsta_ugovora == "odredjeno" else ""
        if req.nova_plata_neto is not None and req.nova_plata_neto > 0:
            update["plata_neto"] = float(req.nova_plata_neto)
        if req.nova_pozicija:
            update["pozicija"] = req.nova_pozicija
        await db.employees.update_one({"id": req.employee_id}, {"$set": update})
    
    # Save record
    record = {
        "id": str(uuid.uuid4()),
        "filename": output_filename,
        "pdf_filename": pdf_filename,
        "template": "ANEKS UGOVORA O RADU.docx",
        "company_id": company["id"],
        "company_naziv": company.get("naziv", ""),
        "employee_id": req.employee_id,
        "employee_naziv": f"{employee.get('ime','')} {employee.get('prezime','')}".strip(),
        "aneks_broj": aneks_broj,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username,
    }
    await db.generated_documents.insert_one(dict(record))
    
    return {
        "success": True,
        "filename": output_filename,
        "pdf_filename": pdf_filename,
        "preview_url": f"/api/documents/preview/{pdf_filename}",
        "download_url": f"/api/documents/download/{output_filename}",
        "aneks_broj": aneks_broj,
        "record": record,
    }


# ============================================================================
# OSNIVANJE DOO — generisanje 4 dokumenta (Odluka o osnivanju, Imenovanje
# direktora, Saglasnost, Statut) na osnovu unetih podataka.
# ============================================================================

def _remove_yellow_highlights(doc: Document):
    """Ukloni žute highlight-ove sa svih runova (koristi se kao marker u template-ima)."""
    from docx.enum.text import WD_COLOR_INDEX
    for para in doc.paragraphs:
        for run in para.runs:
            try:
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    run.font.highlight_color = None
            except Exception:
                pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        try:
                            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                                run.font.highlight_color = None
                        except Exception:
                            pass


def _build_founding_replacements(req: 'FoundingRequest') -> Dict[str, str]:
    """Mapira sve hardkodirane vrijednosti iz 4 šablona na korisnikove podatke."""
    # Datum
    if req.datum_odluke:
        try:
            dt = datetime.fromisoformat(req.datum_odluke.replace('Z', ''))
            datum_str = dt.strftime("%d.%m.%Y")
        except Exception:
            datum_str = req.datum_odluke
    else:
        datum_str = datetime.now(timezone.utc).strftime("%d.%m.%Y")
    
    # Mjesec i godina (za statut headera)
    try:
        dt2 = datetime.fromisoformat((req.datum_odluke or datetime.now(timezone.utc).isoformat()).replace('Z', ''))
        mjeseci = ["JANUAR", "FEBRUAR", "MART", "APRIL", "MAJ", "JUN", "JUL", "AVGUST", "SEPTEMBAR", "OKTOBAR", "NOVEMBAR", "DECEMBAR"]
        mjesec_name = mjeseci[dt2.month - 1]
        godina_int = dt2.year
    except Exception:
        mjesec_name = "MAJ"
        godina_int = datetime.now(timezone.utc).year
    
    # Osnivač isprava
    if req.osnivac_is_stranac:
        osnivac_jmb_label = "Br. pasoša"
        osnivac_jmb_value = req.osnivac_pasos or "____________"
    else:
        osnivac_jmb_label = "JMB"
        osnivac_jmb_value = req.osnivac_jmb or "____________"
    
    # Direktor (može biti isti kao osnivač)
    if req.direktor_isti_kao_osnivac:
        direktor_ime = req.osnivac_ime_prezime
        direktor_jmb = osnivac_jmb_value
        direktor_drzava = req.osnivac_drzava
        direktor_adresa = req.osnivac_adresa
        direktor_pol = req.direktor_pol  # uvijek se uzima iz direktor_pol polja
    else:
        direktor_ime = req.direktor_ime_prezime
        if req.direktor_is_stranac:
            direktor_jmb = req.direktor_pasos or "____________"
        else:
            direktor_jmb = req.direktor_jmb or "____________"
        direktor_drzava = req.direktor_drzava
        direktor_adresa = req.direktor_adresa
        direktor_pol = req.direktor_pol
    
    # Gramatika za pol direktora
    is_zensko = (direktor_pol or "M").upper() == "Z"
    saglasan_label = "SAGLASNA" if is_zensko else "SAGLASAN"
    saglasan_lower = "saglasna" if is_zensko else "saglasan"
    rodjen_lower = "rodjena" if is_zensko else "rodjen"
    
    # Datum rođenja (za saglasnost)
    if req.osnivac_datum_rodjenja:
        try:
            dr = datetime.fromisoformat(req.osnivac_datum_rodjenja.replace('Z', ''))
            datum_rod_str = dr.strftime("%d.%m.%Y")
        except Exception:
            datum_rod_str = req.osnivac_datum_rodjenja
    else:
        datum_rod_str = "__.__.____"
    
    # Naziv pečata (default = skraćeni bez DOO)
    pecat = req.firma_naziv_pecat or req.firma_naziv_skraceni.replace("DOO", "").replace("\"", "").replace("ULCINJ", "").strip()
    
    # ===== Hardkodirane vrijednosti iz template-a koje se zamjenjuju =====
    # Sve sa "ELA&ART" / "ARJANA CEKOVIQ" itd. su sample, mapiramo ih na korisnikove vrijednosti.
    
    # Telefon - parsiraj na "+382" prefix i broj
    tel_raw = (req.firma_telefon or "").strip()
    if tel_raw:
        # Ukloni prefix "+382 " ili "+382"
        tel_no_prefix = tel_raw.replace('+382 ', '').replace('+382', '').strip()
        tel_full = tel_raw if tel_raw.startswith('+') else f"+382 {tel_no_prefix}"
    else:
        tel_no_prefix = "____________"
        tel_full = "+382 ____________"
    
    # === Adresa firme (sjedište) - puna adresa ===
    sjediste_full = req.firma_sjediste_adresa  # već je puna, npr. "SELITA BB ULCINJ"
    # === Adresa direktora (za Član 6 statuta) — pripada direktoru, ne firmi ===
    direktor_adresa_full = direktor_adresa or sjediste_full  # fallback na firmu ako prazno
    
    repl = {
        # === Vrsta djelatnosti opis (treba prvo da se zamijeni jer je dio fullname-a) ===
        'ZA PROIZVODNJU, PROMET I USLUGE': req.firma_vrsta_djelatnosti_opis,
        
        # === FIRMA - pun naziv (u Odluci, Imenovanju, Saglasnosti — gdje je u 1 runu) ===
        'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU ZA PROIZVODNJU, PROMET I USLUGE "ELA&ART " ULCINJ': req.firma_naziv_pun,
        'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU ZA PROIZVODNJU, PROMET I USLUGE "ELA&ART" ULCINJ': req.firma_naziv_pun,
        'DRUŠTVO SA OGRANIČENOM ODGOVORNOŠĆU ZA PROIZVODNJU, PROMET I USLUGE " ELA&ART " ULCINJ': req.firma_naziv_pun,
        
        # === Skraćeni naziv ===
        'DOO "ELA&ART " ULCINJ': req.firma_naziv_skraceni,
        'DOO "ELA&ART" ULCINJ': req.firma_naziv_skraceni,
        
        # === SJEDIŠTE FIRME (Statut P56 + razne pojavljivanja u Odluci/Imenovanju) ===
        # Specifične fraze prvo (najduže) — kontekstualno preciznije
        'Sjedište društva je u: BRAJŠE BB ULCINJ': f'Sjedište društva je u: {sjediste_full}',
        'Sjedište društva je u: BRAJŠE  BB.ULCINJ': f'Sjedište društva je u: {sjediste_full}',
        'Sjedište društva je u: BRAJŠE   BB.ULCINJ': f'Sjedište društva je u: {sjediste_full}',
        'Adresa sjedišta i adresa za prijem službene pošte je:': 'Adresa sjedišta i adresa za prijem službene pošte je:',
        'sjedište BRAJŠE BB ULCINJ': f'sjedište {sjediste_full}',
        'sjedište BRAJŠE BB.ULCINJ': f'sjedište {sjediste_full}',
        
        # === ADRESA OSNIVAČA u Odluci o osnivanju (P0 preambula i P63 potpis) ===
        # P0: "ARJANA CEKOVIQ , JMB: ... sa prebivalištem na adresi BRAJŠE BB. ULCINJ, kao osnivač"
        'prebivalištem na adresi BRAJŠE BB. ULCINJ': f'prebivalištem na adresi {req.osnivac_adresa}',
        'prebivalištem na adresi BRAJŠE  BB. ULCINJ': f'prebivalištem na adresi {req.osnivac_adresa}',
        'prebivalištem na adresi BRAJŠE BB.ULCINJ': f'prebivalištem na adresi {req.osnivac_adresa}',
        'sa prebivalištem na adresi BRAJŠE': f'sa prebivalištem na adresi {req.osnivac_adresa.split(",")[0] if req.osnivac_adresa else "BRAJŠE"}',
        
        # === ADRESA OSNIVAČA u Statutu - preambula (P33: "sa adresom BRAJŠE BB.ULCINJ") ===
        # Ovo se odnosi na osnivača, NE na sjedište firme!
        'sa adresom BRAJŠE BB.ULCINJ': f'sa adresom {req.osnivac_adresa}',
        'sa adresom BRAJŠE  BB.ULCINJ': f'sa adresom {req.osnivac_adresa}',
        'sa adresom BRAJŠE BB. ULCINJ': f'sa adresom {req.osnivac_adresa}',
        'sa adresom BRAJŠE BB ULCINJ': f'sa adresom {req.osnivac_adresa}',
        
        # === ADRESA PREBIVALIŠTA DIREKTORA (Statut Član 6 - P761, P765) ===
        'adresa prebivališta BRAJŠE  BB.ULCINJ': f'adresa prebivališta {direktor_adresa_full}',
        'adresa prebivališta BRAJŠE   BB.ULCINJ': f'adresa prebivališta {direktor_adresa_full}',
        'adresa prebivališta BRAJŠE BB.ULCINJ': f'adresa prebivališta {direktor_adresa_full}',
        'adresa prebivališta BRAJŠE BB ULCINJ': f'adresa prebivališta {direktor_adresa_full}',
        'sa adresom prebivališta BRAJŠE   BB.ULCINJ': f'sa adresom prebivališta {direktor_adresa_full}',
        'sa adresom prebivališta BRAJŠE  BB.ULCINJ': f'sa adresom prebivališta {direktor_adresa_full}',
        'sa adresom prebivališta BRAJŠE BB.ULCINJ': f'sa adresom prebivališta {direktor_adresa_full}',
        'sa adresom prebivališta BRAJŠE BB ULCINJ': f'sa adresom prebivališta {direktor_adresa_full}',
        
        # === PUNOMOĆNIK (Član 8.1 Odluke o osnivanju) ===
        # Originalni text: "može zastupati punomoćnik –ARJANA CEKOVIQ , JMBG 2012985225015"
        # Zamijena: ako podnosi_punomocnik=True → koristi unesene podatke
        #          inače → ostavi crtice (osnivač sam podnosi)
        'može zastupati punomoćnik –ARJANA CEKOVIQ , JMBG 2012985225015': (
            f'može zastupati punomoćnik – {req.punomocnik_ime_prezime}, JMBG {req.punomocnik_jmbg}'
            if req.podnosi_punomocnik and req.punomocnik_ime_prezime
            else 'može zastupati punomoćnik – ____________________, JMBG ____________________'
        ),
        'može zastupati punomoćnik – ARJANA CEKOVIQ , JMBG 2012985225015': (
            f'može zastupati punomoćnik – {req.punomocnik_ime_prezime}, JMBG {req.punomocnik_jmbg}'
            if req.podnosi_punomocnik and req.punomocnik_ime_prezime
            else 'može zastupati punomoćnik – ____________________, JMBG ____________________'
        ),
        
        # === Generic fallback za preostale 'BRAJŠE BB' (npr. samo "BRAJŠE BB ULCINJ" bez prefiksa) ===
        # Ovo ide ZADNJE jer su prethodne specifične zamjene preciznije
        'BRAJŠE BB ULCINJ': sjediste_full,
        'BRAJŠE BB.ULCINJ': sjediste_full,
        'BRAJŠE  BB.ULCINJ': sjediste_full,
        'BRAJŠE   BB.ULCINJ': sjediste_full,
        
        # === Osnivač ===
        'ARJANA CEKOVIQ': req.osnivac_ime_prezime,
        # === JMB osnivača ===
        '2012985225015': osnivac_jmb_value,
        'JMB: 2012985225015': f"{osnivac_jmb_label}: {osnivac_jmb_value}",
        'JMB:2012985225015': f"{osnivac_jmb_label}:{osnivac_jmb_value}",
        'JMBG 2012985225015': f"{osnivac_jmb_label} {osnivac_jmb_value}",
        # === Država OSNIVAČA - samo gdje se govori o državljanstvu/prebivalištu ===
        # Specifične fraze (NE diraju "propisi Crne Gore" — to ostaje uvijek)
        'iz  Crne Gore': f'iz {req.osnivac_drzava}',
        'iz Crne Gore': f'iz {req.osnivac_drzava}',
        # NE zamjenjuj generic "Crne Gore" — može pokvariti "propisi Crne Gore", "Vlade Crne Gore", "zakoni Crne Gore" itd.
        # === Datum rođenja (samo u SAGLASNOST) ===
        '20.12.1985': datum_rod_str,
        # === Pol direktora u Saglasnosti (rodjena/rodjen, saglasna/saglasan) ===
        'rodjena 20.12.1985': f'{rodjen_lower} {datum_rod_str}',
        ', rodjena ': f', {rodjen_lower} ',
        ' rodjena ': f' {rodjen_lower} ',
        'saglasna sam': f'{saglasan_lower} sam',
        ' saglasna ': f' {saglasan_lower} ',
        'SAGLASNA': saglasan_label,
        # Procenat udjela
        '100 %': f"{req.osnivac_procenat:g} %",
        '100%': f"{req.osnivac_procenat:g}%",
        
        # === Telefon (cijela poruka u statutu: "+382 69508359") ===
        '+382 69508359': tel_full,
        '69508359': tel_no_prefix,
        # Ne briši "+382" prefiks
        
        # === Email ===
        'advanced.acct@hotmail.com': req.firma_email or "____________",
        
        # === Djelatnost ===
        '47.11\tNespecijalizovana trgovina na malo pretežno hranom, pićima I duvanskim proizvodima': f"{req.firma_sifra_djelatnosti}\t{req.firma_naziv_djelatnosti}",
        '47.11': req.firma_sifra_djelatnosti,
        'Nespecijalizovana trgovina na malo pretežno hranom, pićima I duvanskim proizvodima': req.firma_naziv_djelatnosti,
        
        # === Datum statuta header ===
        'Ulcinj, MAJ  2026': f"Ulcinj, {mjesec_name} {godina_int}",
        'Ulcinj, MAJ 2026': f"Ulcinj, {mjesec_name} {godina_int}",
        'MAJ  2026': f"{mjesec_name} {godina_int}",
        'MAJ 2026': f"{mjesec_name} {godina_int}",
        'MAJ': mjesec_name,
        '2026': str(godina_int),
        
        # === Datum odluke ===
        '12.05.2026': datum_str,
        'dana 12.05.2026': f"dana {datum_str}",
        
        # === Osnovni kapital ===
        '1,00 EUR': f"{req.osnovni_kapital:.2f} EUR".replace('.', ','),
        
        # === Grad ===
        'U ULCINJ': f"U {req.firma_grad}",
        'U Ulcinju': f"U {req.firma_grad}",
        
        # === Naziv pečata/skraćeni (ELA&ART) — bitno zadnje da ne dira pun naziv ===
        '" ELA&ART "': f'" {pecat} "',
        '"ELA&ART"': f'"{pecat}"',
        ' ELA&ART ': f' {pecat} ',
        'ELA&ART': pecat,
    }
    
    # Ako se direktor razlikuje od osnivača, dodaj specifične zamjene
    # (Ovo nije potpuno: tek nakon prve generacije korisnik može testirati i mi proširimo)
    
    return repl


# Load activity codes once at module level
_SIFRE_DJELATNOSTI_CACHE = None

def _load_sifre_djelatnosti():
    global _SIFRE_DJELATNOSTI_CACHE
    if _SIFRE_DJELATNOSTI_CACHE is None:
        try:
            import json
            with open(ROOT_DIR / "data" / "sifre_djelatnosti.json", "r", encoding="utf-8") as f:
                _SIFRE_DJELATNOSTI_CACHE = json.load(f)
        except Exception:
            _SIFRE_DJELATNOSTI_CACHE = []
    return _SIFRE_DJELATNOSTI_CACHE


@api_router.get("/sifre-djelatnosti")
async def search_sifre(q: Optional[str] = None, limit: int = 50, username: str = Depends(get_current_user)):
    """Pretraga šifri djelatnosti — vraća listu (šifra, naziv) koja sadrži upit."""
    sifre = _load_sifre_djelatnosti()
    if not q:
        return sifre[:limit]
    q_lower = q.lower().strip()
    out = [s for s in sifre if q_lower in s["naziv"].lower() or q_lower in s["sifra"].lower()]
    return out[:limit]


# ============================================================================
# FINANSIJE — Cjenovnik, mjesečne uplate, dodatne usluge, troškovi
# ============================================================================

class CompanyPricing(BaseModel):
    company_id: str
    monthly_fee: float = 0.0  # standardna mjesečna naknada (EUR) - za firmu (bez objekata)
    napomena: str = ""


class ObjekatPricing(BaseModel):
    objekat_id: str
    company_id: str
    monthly_fee: float = 0.0  # cijena za konkretan objekat
    napomena: str = ""


class MonthlyPayment(BaseModel):
    company_id: str
    godina: int  # 2026
    mjesec: int  # 1-12
    iznos: float  # može biti drugačiji od default-a
    is_paid: bool = False
    datum_naplate: Optional[str] = ""  # ISO date
    napomena: Optional[str] = ""


class ExtraService(BaseModel):
    company_id: str
    naziv: str  # "Osnivanje DOO", "Izvještaj banci"...
    datum: str  # ISO date — kada je izvršeno
    iznos: float  # naplata
    is_paid: bool = False
    datum_naplate: Optional[str] = ""
    napomena: Optional[str] = ""


class Expense(BaseModel):
    naziv: str  # "Kancelarija najam", "Software Microsoft Office"...
    datum: str  # ISO date
    iznos: float
    kategorija: str = "opsti"  # "opsti" (opšti agencijski) / "usluga" (vezan za uslugu)
    extra_service_id: Optional[str] = ""  # ako je vezan za extra uslugu — povezuje sa ExtraService
    company_id: Optional[str] = ""  # ako je trošak vezan za firmu
    napomena: Optional[str] = ""


# === CJENOVNIK PO FIRMI ===

@api_router.get("/finance/pricing")
async def list_pricing(username: str = Depends(get_current_user)):
    """Cjenovnik za sve firme + svaki objekat ima zasebnu cijenu (ako je postavljena).
    monthly_fee na firmi = bazna cijena. Ukupno za firmu = bazna + zbir objekata."""
    pricings = await db.company_pricing.find({}, {"_id": 0}).to_list(500)
    pricing_by_cid = {p["company_id"]: p for p in pricings}
    
    obj_pricings = await db.objekat_pricing.find({}, {"_id": 0}).to_list(2000)
    obj_pricing_by_oid = {p["objekat_id"]: p for p in obj_pricings}
    
    objekti_all = await db.company_objekti.find({}, {"_id": 0}).to_list(2000)
    obj_by_cid: Dict[str, List[Dict]] = {}
    for o in objekti_all:
        obj_by_cid.setdefault(o["company_id"], []).append(o)
    
    companies = await db.companies.find({}, {"_id": 0, "id": 1, "naziv": 1, "naziv_skraceni": 1, "pib": 1}).to_list(500)
    
    out = []
    for c in companies:
        p = pricing_by_cid.get(c["id"], {})
        base_fee = float(p.get("monthly_fee", 0.0) or 0)
        # Objekti za ovu firmu
        objs = []
        objs_sum = 0.0
        for o in obj_by_cid.get(c["id"], []):
            op = obj_pricing_by_oid.get(o["id"], {})
            obj_fee = float(op.get("monthly_fee", 0.0) or 0)
            objs_sum += obj_fee
            objs.append({
                "objekat_id": o["id"],
                "naziv": o["naziv"],
                "adresa": o.get("adresa", ""),
                "monthly_fee": obj_fee,
                "napomena": op.get("napomena", ""),
            })
        out.append({
            "company_id": c["id"],
            "naziv": c["naziv"],
            "naziv_skraceni": c.get("naziv_skraceni", ""),
            "pib": c.get("pib", ""),
            "monthly_fee": base_fee,
            "napomena": p.get("napomena", ""),
            "objekti": objs,
            "total_fee": round(base_fee + objs_sum, 2),
        })
    return out


@api_router.put("/finance/pricing/{company_id}")
async def set_pricing(company_id: str, req: CompanyPricing, username: str = Depends(get_current_user)):
    rec = {
        "company_id": company_id,
        "monthly_fee": req.monthly_fee,
        "napomena": req.napomena or "",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.company_pricing.update_one(
        {"company_id": company_id},
        {"$set": rec},
        upsert=True,
    )
    return rec


@api_router.put("/finance/pricing/objekat/{objekat_id}")
async def set_objekat_pricing(objekat_id: str, req: ObjekatPricing, username: str = Depends(get_current_user)):
    rec = {
        "objekat_id": objekat_id,
        "company_id": req.company_id,
        "monthly_fee": req.monthly_fee,
        "napomena": req.napomena or "",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.objekat_pricing.update_one(
        {"objekat_id": objekat_id},
        {"$set": rec},
        upsert=True,
    )
    return rec


# === MJESEČNE UPLATE ===

@api_router.get("/finance/payments")
async def list_payments(
    godina: Optional[int] = None,
    mjesec: Optional[int] = None,
    company_id: Optional[str] = None,
    username: str = Depends(get_current_user),
):
    """Lista uplata. Ako su godina+mjesec dati — auto-generiše stavke za firme koje još nemaju zapis."""
    query: Dict = {}
    if godina:
        query["godina"] = godina
    if mjesec:
        query["mjesec"] = mjesec
    if company_id:
        query["company_id"] = company_id
    
    existing = await db.monthly_payments.find(query, {"_id": 0}).to_list(1000)
    
    # Auto-merge sa firmama da bismo prikazali sve firme bez obzira da li imaju uplatu
    if godina and mjesec:
        # Generiraj virtuelne stavke za firme koje nemaju zapis
        pricings = await db.company_pricing.find({}, {"_id": 0}).to_list(500)
        pricing_by_cid = {p["company_id"]: p for p in pricings}
        # Učitaj i cijene objekata (suma se dodaje na bazu)
        obj_pricings = await db.objekat_pricing.find({}, {"_id": 0}).to_list(2000)
        obj_sum_by_cid: Dict[str, float] = {}
        for op in obj_pricings:
            cid_local = op.get("company_id", "")
            obj_sum_by_cid[cid_local] = obj_sum_by_cid.get(cid_local, 0.0) + float(op.get("monthly_fee", 0) or 0)
        existing_cids = {p["company_id"] for p in existing}
        companies = await db.companies.find({}, {"_id": 0, "id": 1, "naziv": 1, "naziv_skraceni": 1}).to_list(500)
        
        out = list(existing)
        for c in companies:
            if c["id"] not in existing_cids:
                base_fee = pricing_by_cid.get(c["id"], {}).get("monthly_fee", 0.0) or 0
                obj_sum = obj_sum_by_cid.get(c["id"], 0.0)
                default_fee = round(float(base_fee) + obj_sum, 2)
                out.append({
                    "id": None,
                    "company_id": c["id"],
                    "company_naziv": c["naziv"],
                    "godina": godina,
                    "mjesec": mjesec,
                    "iznos": default_fee,
                    "is_paid": False,
                    "datum_naplate": "",
                    "napomena": "",
                    "_virtual": True,
                })
        # Dodaj naziv za postojeće
        co_by_cid = {c["id"]: c for c in companies}
        for p in out:
            if not p.get("company_naziv"):
                co = co_by_cid.get(p["company_id"])
                p["company_naziv"] = co["naziv"] if co else ""
        out.sort(key=lambda x: x.get("company_naziv", ""))
        return out
    
    # Dodaj naziv firme za istorijski pregled
    cids = list({p["company_id"] for p in existing})
    co_map = {}
    if cids:
        cos = await db.companies.find({"id": {"$in": cids}}, {"_id": 0, "id": 1, "naziv": 1}).to_list(500)
        co_map = {c["id"]: c["naziv"] for c in cos}
    for p in existing:
        p["company_naziv"] = co_map.get(p["company_id"], "")
    return existing


@api_router.post("/finance/payments")
async def upsert_payment(req: MonthlyPayment, username: str = Depends(get_current_user)):
    """Sačuvaj/ažuriraj uplatu za firmu × godina × mjesec."""
    query = {"company_id": req.company_id, "godina": req.godina, "mjesec": req.mjesec}
    existing = await db.monthly_payments.find_one(query, {"_id": 0})
    
    rec = {
        "company_id": req.company_id,
        "godina": req.godina,
        "mjesec": req.mjesec,
        "iznos": req.iznos,
        "is_paid": req.is_paid,
        "datum_naplate": req.datum_naplate or "",
        "napomena": req.napomena or "",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if existing:
        await db.monthly_payments.update_one(query, {"$set": rec})
        rec["id"] = existing.get("id")
    else:
        rec["id"] = str(uuid.uuid4())
        rec["created_at"] = datetime.now(timezone.utc).isoformat()
        await db.monthly_payments.insert_one(dict(rec))
    return rec


# === DODATNE USLUGE ===

@api_router.get("/finance/services")
async def list_services(
    company_id: Optional[str] = None,
    godina: Optional[int] = None,
    username: str = Depends(get_current_user),
):
    query: Dict = {}
    if company_id:
        query["company_id"] = company_id
    items = await db.extra_services.find(query, {"_id": 0}).sort("datum", -1).to_list(500)
    
    # Filter po godini
    if godina:
        items = [it for it in items if it.get("datum", "").startswith(str(godina))]
    
    # Dodaj naziv firme
    cids = list({i["company_id"] for i in items if i.get("company_id")})
    co_map = {}
    if cids:
        cos = await db.companies.find({"id": {"$in": cids}}, {"_id": 0, "id": 1, "naziv": 1}).to_list(500)
        co_map = {c["id"]: c["naziv"] for c in cos}
    for it in items:
        it["company_naziv"] = co_map.get(it.get("company_id", ""), "")
    return items


@api_router.post("/finance/services")
async def create_service(req: ExtraService, username: str = Depends(get_current_user)):
    rec = {
        "id": str(uuid.uuid4()),
        "company_id": req.company_id,
        "naziv": req.naziv,
        "datum": req.datum,
        "iznos": req.iznos,
        "is_paid": req.is_paid,
        "datum_naplate": req.datum_naplate or "",
        "napomena": req.napomena or "",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.extra_services.insert_one(dict(rec))
    return rec


@api_router.patch("/finance/services/{sid}")
async def update_service(sid: str, req: ExtraService, username: str = Depends(get_current_user)):
    updates = {
        "company_id": req.company_id,
        "naziv": req.naziv,
        "datum": req.datum,
        "iznos": req.iznos,
        "is_paid": req.is_paid,
        "datum_naplate": req.datum_naplate or "",
        "napomena": req.napomena or "",
    }
    await db.extra_services.update_one({"id": sid}, {"$set": updates})
    return updates


@api_router.delete("/finance/services/{sid}")
async def delete_service(sid: str, username: str = Depends(get_current_user)):
    await db.extra_services.delete_one({"id": sid})
    return {"success": True}


# === TROŠKOVI ===

@api_router.get("/finance/expenses")
async def list_expenses(
    kategorija: Optional[str] = None,
    godina: Optional[int] = None,
    company_id: Optional[str] = None,
    extra_service_id: Optional[str] = None,
    username: str = Depends(get_current_user),
):
    query: Dict = {}
    if kategorija:
        query["kategorija"] = kategorija
    if company_id:
        query["company_id"] = company_id
    if extra_service_id:
        query["extra_service_id"] = extra_service_id
    items = await db.expenses.find(query, {"_id": 0}).sort("datum", -1).to_list(1000)
    if godina:
        items = [it for it in items if it.get("datum", "").startswith(str(godina))]
    return items


@api_router.post("/finance/expenses")
async def create_expense(req: Expense, username: str = Depends(get_current_user)):
    rec = {
        "id": str(uuid.uuid4()),
        "naziv": req.naziv,
        "datum": req.datum,
        "iznos": req.iznos,
        "kategorija": req.kategorija,
        "extra_service_id": req.extra_service_id or "",
        "company_id": req.company_id or "",
        "napomena": req.napomena or "",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.expenses.insert_one(dict(rec))
    return rec


@api_router.patch("/finance/expenses/{eid}")
async def update_expense(eid: str, req: Expense, username: str = Depends(get_current_user)):
    updates = {
        "naziv": req.naziv,
        "datum": req.datum,
        "iznos": req.iznos,
        "kategorija": req.kategorija,
        "extra_service_id": req.extra_service_id or "",
        "company_id": req.company_id or "",
        "napomena": req.napomena or "",
    }
    await db.expenses.update_one({"id": eid}, {"$set": updates})
    return updates


@api_router.delete("/finance/expenses/{eid}")
async def delete_expense(eid: str, username: str = Depends(get_current_user)):
    await db.expenses.delete_one({"id": eid})
    return {"success": True}


# === PROFIT SUMMARY ===

@api_router.get("/finance/summary")
async def finance_summary(godina: int = None, username: str = Depends(get_current_user)):
    """Sažetak prihoda i rashoda za godinu."""
    if not godina:
        godina = datetime.now(timezone.utc).year
    
    # Prihodi mjesečnih naknada
    payments = await db.monthly_payments.find({"godina": godina, "is_paid": True}, {"_id": 0}).to_list(2000)
    income_monthly = sum(p.get("iznos", 0) for p in payments)
    income_monthly_pending = sum(p.get("iznos", 0) for p in await db.monthly_payments.find({"godina": godina, "is_paid": False}, {"_id": 0}).to_list(2000))
    
    # Prihodi extra usluga
    services = await db.extra_services.find({"is_paid": True}, {"_id": 0}).to_list(2000)
    services = [s for s in services if s.get("datum", "").startswith(str(godina))]
    income_extra = sum(s.get("iznos", 0) for s in services)
    
    services_pending = await db.extra_services.find({"is_paid": False}, {"_id": 0}).to_list(2000)
    services_pending = [s for s in services_pending if s.get("datum", "").startswith(str(godina))]
    income_extra_pending = sum(s.get("iznos", 0) for s in services_pending)
    
    # Troškovi
    expenses = await db.expenses.find({}, {"_id": 0}).to_list(2000)
    expenses_year = [e for e in expenses if e.get("datum", "").startswith(str(godina))]
    expense_opsti = sum(e.get("iznos", 0) for e in expenses_year if e.get("kategorija") == "opsti")
    expense_usluga = sum(e.get("iznos", 0) for e in expenses_year if e.get("kategorija") == "usluga")
    
    total_income = income_monthly + income_extra
    total_expense = expense_opsti + expense_usluga
    profit_net = total_income - total_expense
    profit_extra = income_extra - expense_usluga
    profit_monthly = income_monthly - expense_opsti
    
    # Po mjesecima
    monthly_breakdown = {}
    for m in range(1, 13):
        monthly_breakdown[m] = {
            "income_monthly": sum(p.get("iznos", 0) for p in payments if p.get("mjesec") == m),
            "income_extra": sum(s.get("iznos", 0) for s in services if s.get("datum", "").startswith(f"{godina}-{m:02d}")),
            "expense": sum(e.get("iznos", 0) for e in expenses_year if e.get("datum", "").startswith(f"{godina}-{m:02d}")),
        }
    
    return {
        "godina": godina,
        "income_monthly_paid": income_monthly,
        "income_monthly_pending": income_monthly_pending,
        "income_extra_paid": income_extra,
        "income_extra_pending": income_extra_pending,
        "total_income": total_income,
        "expense_opsti": expense_opsti,
        "expense_usluga": expense_usluga,
        "total_expense": total_expense,
        "profit_net": profit_net,
        "profit_monthly_services": profit_monthly,
        "profit_extra_services": profit_extra,
        "monthly_breakdown": monthly_breakdown,
    }


# === ALARMI ZA NEPLAĆENE RAČUNE ===

@api_router.get("/finance/overdue")
async def list_overdue(days: int = 30, username: str = Depends(get_current_user)):
    """Lista firmi sa neplaćenim mjesečnim naknadama starijim od X dana.
    Računa se: za svaki mjesec prošle/tekuće godine, ako je rok prošao i firma nije platila."""
    today = datetime.now(timezone.utc)
    # Cjenovnik (firma) + objekat pricing
    pricings = await db.company_pricing.find({}, {"_id": 0}).to_list(500)
    base_by_cid = {p["company_id"]: float(p.get("monthly_fee", 0) or 0) for p in pricings}
    obj_pricings = await db.objekat_pricing.find({}, {"_id": 0}).to_list(2000)
    obj_sum_by_cid: Dict[str, float] = {}
    for op in obj_pricings:
        cid_local = op.get("company_id", "")
        obj_sum_by_cid[cid_local] = obj_sum_by_cid.get(cid_local, 0.0) + float(op.get("monthly_fee", 0) or 0)
    # total fee = base + sum objekti
    all_cids = set(base_by_cid.keys()) | set(obj_sum_by_cid.keys())
    pricing_by_cid: Dict[str, Dict] = {}
    for cid in all_cids:
        total = round(base_by_cid.get(cid, 0.0) + obj_sum_by_cid.get(cid, 0.0), 2)
        if total > 0:
            pricing_by_cid[cid] = {"company_id": cid, "monthly_fee": total}
    if not pricing_by_cid:
        return []
    
    # Učitaj minimalnu godinu praćenja iz finance_settings
    settings = await db.finance_settings.find_one({"_id": "default"}, {"_id": 0}) or {}
    active_years = settings.get("active_years") or [datetime.now(timezone.utc).year, datetime.now(timezone.utc).year + 1]
    min_year = min(active_years)
    
    # Postojeće uplate
    payments = await db.monthly_payments.find({}, {"_id": 0}).to_list(5000)
    paid_keys = {(p["company_id"], p["godina"], p["mjesec"]) for p in payments if p.get("is_paid")}
    
    # Firme — naziv
    companies = await db.companies.find({"id": {"$in": list(pricing_by_cid.keys())}}, {"_id": 0, "id": 1, "naziv": 1}).to_list(500)
    co_by_cid = {c["id"]: c["naziv"] for c in companies}
    
    # Generiši listu neplaćenih mjeseci u zadnjih `days` (sve neplaćene mjesece od prije više od `days` dana)
    out: Dict[str, Dict] = {}
    # iteriraj zadnjih 24 mjeseca
    cur = today
    for _ in range(24):
        mjesec = cur.month
        godina = cur.year
        # Rok = kraj mjeseca + days dana grace
        # Smatra se overdue ako je prošlo > days od kraja mjeseca
        try:
            end_of_month = datetime(godina, mjesec + 1, 1, tzinfo=timezone.utc) if mjesec < 12 else datetime(godina + 1, 1, 1, tzinfo=timezone.utc)
        except Exception:
            end_of_month = today
        overdue_threshold = end_of_month + timedelta(days=days)
        if today < overdue_threshold:
            # mjesec još nije overdue
            cur = (cur.replace(day=1) - timedelta(days=1))
            continue
        if godina < min_year:
            # Skip stare godine ispod min_year iz settings-a
            break
        for cid, p in pricing_by_cid.items():
            if (cid, godina, mjesec) in paid_keys:
                continue
            entry = out.setdefault(cid, {
                "company_id": cid,
                "naziv": co_by_cid.get(cid, ""),
                "monthly_fee": p.get("monthly_fee", 0.0),
                "overdue_months": [],
                "total_owed": 0.0,
                "oldest_due": None,
            })
            entry["overdue_months"].append({"godina": godina, "mjesec": mjesec})
            entry["total_owed"] += p.get("monthly_fee", 0.0)
            due_str = f"{godina}-{mjesec:02d}"
            if not entry["oldest_due"] or due_str < entry["oldest_due"]:
                entry["oldest_due"] = due_str
        # prošli mjesec
        cur = (cur.replace(day=1) - timedelta(days=1))
    
    result = list(out.values())
    # Sortiraj po najstarijem dugovanju
    result.sort(key=lambda x: x["oldest_due"] or "9999")
    # Filtriraj firme bez dugovanja (može se desiti)
    result = [r for r in result if r["overdue_months"]]
    return result


# === FINANCE SETTINGS (aktivne godine) ===

@api_router.get("/finance/settings")
async def get_finance_settings(username: str = Depends(get_current_user)):
    s = await db.finance_settings.find_one({"_id": "default"}, {"_id": 0})
    if not s:
        cur_year = datetime.now(timezone.utc).year
        s = {"active_years": [cur_year, cur_year + 1]}
    s.setdefault("active_years", [])
    s["active_years"] = sorted(set(int(y) for y in s["active_years"]))
    return s


class FinanceSettings(BaseModel):
    active_years: List[int] = []


@api_router.put("/finance/settings")
async def update_finance_settings(req: FinanceSettings, username: str = Depends(get_current_user)):
    years = sorted(set(int(y) for y in req.active_years))
    await db.finance_settings.update_one(
        {"_id": "default"},
        {"$set": {"active_years": years, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"active_years": years}


@api_router.delete("/finance/clear-year/{year}")
async def clear_year_payments(year: int, username: str = Depends(get_current_user)):
    """Briše sve mjesečne uplate za datu godinu. Korisno kada su svi računi za prošlu godinu plaćeni
    i ne želiš da ti prikazuje overdue alarme za historijske godine."""
    r1 = await db.monthly_payments.delete_many({"godina": year})
    return {"deleted_payments": r1.deleted_count, "year": year}


@api_router.get("/finance/per-client")
async def finance_per_client(godina: int = None, username: str = Depends(get_current_user)):
    """Pojedinačni izvještaj profitabilnosti po klijentu za datu godinu.
    Direktni troškovi = expenses sa kategorija='usluga' koji su povezani sa firmom
    direktno (company_id) ili indirektno (preko extra_service_id → service.company_id)."""
    if not godina:
        godina = datetime.now(timezone.utc).year
    
    companies = await db.companies.find({}, {"_id": 0, "id": 1, "naziv": 1, "naziv_skraceni": 1, "pib": 1}).to_list(500)
    co_by_cid = {c["id"]: c for c in companies}
    
    # Cjenovnik (mjesečne naknade po firmi) — koristi se da bismo izračunali OČEKIVANE prihode
    pricings = await db.company_pricing.find({}, {"_id": 0}).to_list(500)
    base_by_cid = {p["company_id"]: float(p.get("monthly_fee", 0) or 0) for p in pricings}
    obj_pricings = await db.objekat_pricing.find({}, {"_id": 0}).to_list(2000)
    obj_sum_by_cid: Dict[str, float] = {}
    for op in obj_pricings:
        cid_local = op.get("company_id", "")
        obj_sum_by_cid[cid_local] = obj_sum_by_cid.get(cid_local, 0.0) + float(op.get("monthly_fee", 0) or 0)
    all_cids = set(base_by_cid.keys()) | set(obj_sum_by_cid.keys())
    pricing_by_cid: Dict[str, Dict] = {}
    for cid in all_cids:
        total = round(base_by_cid.get(cid, 0.0) + obj_sum_by_cid.get(cid, 0.0), 2)
        if total > 0:
            pricing_by_cid[cid] = {"company_id": cid, "monthly_fee": total}
    
    # Prihodi: mjesečne uplate (naplaćene)
    payments = await db.monthly_payments.find({"godina": godina}, {"_id": 0}).to_list(5000)
    payments_by_cid: Dict[str, List[Dict]] = {}
    for p in payments:
        payments_by_cid.setdefault(p.get("company_id", ""), []).append(p)
    
    # Prihodi: extra usluge u toj godini
    services_all = await db.extra_services.find({}, {"_id": 0}).to_list(2000)
    services = [s for s in services_all if s.get("datum", "").startswith(str(godina))]
    svc_by_id = {s["id"]: s for s in services_all}
    
    # Troškovi (direktni — kategorija usluga, u toj godini)
    expenses_all = await db.expenses.find({"kategorija": "usluga"}, {"_id": 0}).to_list(2000)
    expenses = [e for e in expenses_all if e.get("datum", "").startswith(str(godina))]
    
    # Određivanje koliko mjeseci je "prošlo" do danas za godinu (za očekivani prihod)
    today = datetime.now(timezone.utc)
    if godina < today.year:
        months_passed = 12
    elif godina == today.year:
        months_passed = today.month
    else:
        months_passed = 0  # buduća godina
    
    out: Dict[str, Dict] = {}
    
    def slot(cid: str) -> Dict:
        if cid not in out:
            co = co_by_cid.get(cid, {})
            out[cid] = {
                "company_id": cid,
                "naziv": co.get("naziv", "Nepoznata firma"),
                "naziv_skraceni": co.get("naziv_skraceni", ""),
                "pib": co.get("pib", ""),
                "monthly_fee": pricing_by_cid.get(cid, {}).get("monthly_fee", 0.0),
                "income_monthly_paid": 0.0,
                "income_monthly_pending": 0.0,
                "income_extra_paid": 0.0,
                "income_extra_pending": 0.0,
                "expense_direct": 0.0,
                "n_extra_services": 0,
                "n_paid_months": 0,
                "n_pending_months": 0,
            }
        return out[cid]
    
    # 1) Mjesečne uplate iz DB
    for p in payments:
        cid = p.get("company_id")
        if not cid:
            continue
        s = slot(cid)
        iznos = float(p.get("iznos") or 0)
        if p.get("is_paid"):
            s["income_monthly_paid"] += iznos
            s["n_paid_months"] += 1
        else:
            s["income_monthly_pending"] += iznos
            s["n_pending_months"] += 1
    
    # 2) Firme iz cjenovnika — za njih izračunaj OČEKIVANE pending mjesece koji još nisu u DB
    for cid, pricing in pricing_by_cid.items():
        s = slot(cid)
        fee = float(pricing.get("monthly_fee") or 0)
        # broj mjeseci za koje već imamo zapis (paid + pending)
        recorded_months = set()
        for p in payments_by_cid.get(cid, []):
            recorded_months.add(int(p.get("mjesec") or 0))
        # Mjeseci u godini koji su prošli ali nemaju zapis → očekivano dugovanje
        for m in range(1, months_passed + 1):
            if m not in recorded_months:
                s["income_monthly_pending"] += fee
                s["n_pending_months"] += 1
    
    # 3) Extra usluge
    for sv in services:
        cid = sv.get("company_id")
        if not cid:
            continue
        s = slot(cid)
        iznos = float(sv.get("iznos") or 0)
        if sv.get("is_paid"):
            s["income_extra_paid"] += iznos
        else:
            s["income_extra_pending"] += iznos
        s["n_extra_services"] += 1
    
    # 4) Direktni troškovi
    for e in expenses:
        cid = e.get("company_id") or ""
        if not cid:
            esid = e.get("extra_service_id") or ""
            sv = svc_by_id.get(esid)
            if sv:
                cid = sv.get("company_id", "")
        if not cid:
            continue
        s = slot(cid)
        s["expense_direct"] += float(e.get("iznos") or 0)
    
    # Profit
    for s in out.values():
        s["total_income_paid"] = round(s["income_monthly_paid"] + s["income_extra_paid"], 2)
        s["total_pending"] = round(s["income_monthly_pending"] + s["income_extra_pending"], 2)
        s["profit"] = round(s["total_income_paid"] - s["expense_direct"], 2)
        for k in ["income_monthly_paid", "income_monthly_pending", "income_extra_paid", "income_extra_pending", "expense_direct", "monthly_fee"]:
            s[k] = round(s[k], 2)
    
    result = list(out.values())
    result.sort(key=lambda x: x["profit"], reverse=True)
    return result


# === EXPORT (Excel + PDF) ===

def _month_name(m: int) -> str:
    return ["", "Januar", "Februar", "Mart", "April", "Maj", "Jun", "Jul", "Avgust", "Septembar", "Oktobar", "Novembar", "Decembar"][m]


@api_router.get("/finance/export/excel")
async def export_finance_excel(godina: int = None, username: str = Depends(get_current_user)):
    """Izvoz svih finansijskih podataka u Excel sa više listova."""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from fastapi.responses import StreamingResponse
    
    if not godina:
        godina = datetime.now(timezone.utc).year
    
    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1e40af")
    
    # === Sheet 1: Mjesečne uplate ===
    ws1 = wb.active
    ws1.title = "Mjesečne uplate"
    headers = ["Firma", "Godina", "Mjesec", "Iznos (€)", "Naplaćeno", "Datum naplate", "Napomena"]
    ws1.append(headers)
    for cell in ws1[1]:
        cell.font = header_font
        cell.fill = header_fill
    
    payments = await db.monthly_payments.find({"godina": godina}, {"_id": 0}).to_list(5000)
    cids = list({p["company_id"] for p in payments})
    cos = await db.companies.find({"id": {"$in": cids}}, {"_id": 0, "id": 1, "naziv": 1}).to_list(500)
    co_map = {c["id"]: c["naziv"] for c in cos}
    for p in sorted(payments, key=lambda x: (co_map.get(x["company_id"], ""), x["mjesec"])):
        ws1.append([
            co_map.get(p["company_id"], p["company_id"]),
            p.get("godina"),
            _month_name(p.get("mjesec", 0)),
            p.get("iznos", 0),
            "DA" if p.get("is_paid") else "NE",
            p.get("datum_naplate", ""),
            p.get("napomena", ""),
        ])
    for col in ws1.columns:
        ws1.column_dimensions[col[0].column_letter].width = 22
    
    # === Sheet 2: Dodatne usluge ===
    ws2 = wb.create_sheet("Dodatne usluge")
    ws2.append(["Datum", "Firma", "Naziv usluge", "Iznos (€)", "Naplaćeno", "Datum naplate", "Napomena"])
    for cell in ws2[1]:
        cell.font = header_font
        cell.fill = header_fill
    services = await db.extra_services.find({}, {"_id": 0}).to_list(2000)
    services = [s for s in services if s.get("datum", "").startswith(str(godina))]
    for s in sorted(services, key=lambda x: x.get("datum", "")):
        ws2.append([
            s.get("datum", ""),
            co_map.get(s.get("company_id", ""), s.get("company_id", "")),
            s.get("naziv", ""),
            s.get("iznos", 0),
            "DA" if s.get("is_paid") else "NE",
            s.get("datum_naplate", ""),
            s.get("napomena", ""),
        ])
    for col in ws2.columns:
        ws2.column_dimensions[col[0].column_letter].width = 22
    
    # === Sheet 3: Troškovi ===
    ws3 = wb.create_sheet("Troškovi")
    ws3.append(["Datum", "Naziv", "Kategorija", "Iznos (€)", "Napomena"])
    for cell in ws3[1]:
        cell.font = header_font
        cell.fill = header_fill
    expenses = await db.expenses.find({}, {"_id": 0}).to_list(2000)
    expenses = [e for e in expenses if e.get("datum", "").startswith(str(godina))]
    for e in sorted(expenses, key=lambda x: x.get("datum", "")):
        ws3.append([
            e.get("datum", ""),
            e.get("naziv", ""),
            "Opšti" if e.get("kategorija") == "opsti" else "Vezan za uslugu",
            e.get("iznos", 0),
            e.get("napomena", ""),
        ])
    for col in ws3.columns:
        ws3.column_dimensions[col[0].column_letter].width = 22
    
    # === Sheet 4: Sažetak ===
    ws4 = wb.create_sheet("Sažetak profita")
    ws4.append([f"FINANSIJSKI IZVJEŠTAJ — {godina}"])
    ws4["A1"].font = Font(bold=True, size=14)
    ws4.append([])
    
    income_monthly_paid = sum(p.get("iznos", 0) for p in payments if p.get("is_paid"))
    income_monthly_pending = sum(p.get("iznos", 0) for p in payments if not p.get("is_paid"))
    income_extra_paid = sum(s.get("iznos", 0) for s in services if s.get("is_paid"))
    income_extra_pending = sum(s.get("iznos", 0) for s in services if not s.get("is_paid"))
    expense_opsti = sum(e.get("iznos", 0) for e in expenses if e.get("kategorija") == "opsti")
    expense_usluga = sum(e.get("iznos", 0) for e in expenses if e.get("kategorija") == "usluga")
    profit_net = income_monthly_paid + income_extra_paid - expense_opsti - expense_usluga
    
    ws4.append(["Naplaćeni mjesečni prihodi", f"{income_monthly_paid:.2f} €"])
    ws4.append(["Čeka uplatu (mjesečno)", f"{income_monthly_pending:.2f} €"])
    ws4.append(["Naplaćene dodatne usluge", f"{income_extra_paid:.2f} €"])
    ws4.append(["Čeka uplatu (extra)", f"{income_extra_pending:.2f} €"])
    ws4.append(["Opšti troškovi", f"-{expense_opsti:.2f} €"])
    ws4.append(["Troškovi za usluge", f"-{expense_usluga:.2f} €"])
    ws4.append([])
    ws4.append(["ČISTI PROFIT", f"{profit_net:.2f} €"])
    ws4["A10"].font = Font(bold=True, size=12)
    ws4["B10"].font = Font(bold=True, size=12, color="10b981" if profit_net >= 0 else "ef4444")
    
    ws4.append([])
    ws4.append(["Mjesec", "Mjesečne (€)", "Extra (€)", "Troškovi (€)", "Profit (€)"])
    for cell in ws4[12]:
        cell.font = header_font
        cell.fill = header_fill
    for m in range(1, 13):
        income_m = sum(p.get("iznos", 0) for p in payments if p.get("is_paid") and p.get("mjesec") == m)
        extra_m = sum(s.get("iznos", 0) for s in services if s.get("is_paid") and s.get("datum", "").startswith(f"{godina}-{m:02d}"))
        exp_m = sum(e.get("iznos", 0) for e in expenses if e.get("datum", "").startswith(f"{godina}-{m:02d}"))
        ws4.append([_month_name(m), income_m, extra_m, -exp_m, income_m + extra_m - exp_m])
    
    for col in ws4.columns:
        ws4.column_dimensions[col[0].column_letter].width = 25
    
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"Finansije_{godina}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@api_router.get("/finance/export/pdf")
async def export_finance_pdf(godina: int = None, username: str = Depends(get_current_user)):
    """Izvoz finansijskog izvještaja u PDF (sažetak + mjesečni breakdown)."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from fastapi.responses import StreamingResponse
    
    if not godina:
        godina = datetime.now(timezone.utc).year
    
    # Učitaj podatke (slično kao u summary)
    payments = await db.monthly_payments.find({"godina": godina}, {"_id": 0}).to_list(5000)
    services = await db.extra_services.find({}, {"_id": 0}).to_list(2000)
    services = [s for s in services if s.get("datum", "").startswith(str(godina))]
    expenses = await db.expenses.find({}, {"_id": 0}).to_list(2000)
    expenses = [e for e in expenses if e.get("datum", "").startswith(str(godina))]
    
    income_monthly_paid = sum(p.get("iznos", 0) for p in payments if p.get("is_paid"))
    income_monthly_pending = sum(p.get("iznos", 0) for p in payments if not p.get("is_paid"))
    income_extra_paid = sum(s.get("iznos", 0) for s in services if s.get("is_paid"))
    income_extra_pending = sum(s.get("iznos", 0) for s in services if not s.get("is_paid"))
    expense_opsti = sum(e.get("iznos", 0) for e in expenses if e.get("kategorija") == "opsti")
    expense_usluga = sum(e.get("iznos", 0) for e in expenses if e.get("kategorija") == "usluga")
    profit_net = income_monthly_paid + income_extra_paid - expense_opsti - expense_usluga
    
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontSize=18, alignment=1, spaceAfter=6)
    sub_style = ParagraphStyle("sub", parent=styles["Normal"], fontSize=10, alignment=1, textColor=colors.grey, spaceAfter=18)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, spaceBefore=12, spaceAfter=8, textColor=colors.HexColor("#1e40af"))
    
    elements = []
    elements.append(Paragraph(f"Finansijski izvještaj — {godina}", title_style))
    elements.append(Paragraph(f"Advanced Accounting · Generisano: {datetime.now().strftime('%d.%m.%Y. %H:%M')}", sub_style))
    
    elements.append(Paragraph("Sažetak", h2))
    summary_data = [
        ["Stavka", "Iznos (EUR)"],
        ["Naplaćeni mjesečni prihodi", f"{income_monthly_paid:.2f}"],
        ["Čeka uplatu (mjesečno)", f"{income_monthly_pending:.2f}"],
        ["Naplaćene dodatne usluge", f"{income_extra_paid:.2f}"],
        ["Čeka uplatu (extra)", f"{income_extra_pending:.2f}"],
        ["Opšti troškovi", f"-{expense_opsti:.2f}"],
        ["Troškovi za usluge", f"-{expense_usluga:.2f}"],
        ["ČISTI PROFIT", f"{profit_net:.2f}"],
    ]
    t = Table(summary_data, colWidths=[10*cm, 5*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#dcfce7") if profit_net >= 0 else colors.HexColor("#fee2e2")),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f8fafc")]),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]))
    elements.append(t)
    
    elements.append(Paragraph("Mjesečni pregled", h2))
    month_data = [["Mjesec", "Mjesečne (€)", "Extra (€)", "Troškovi (€)", "Profit (€)"]]
    for m in range(1, 13):
        income_m = sum(p.get("iznos", 0) for p in payments if p.get("is_paid") and p.get("mjesec") == m)
        extra_m = sum(s.get("iznos", 0) for s in services if s.get("is_paid") and s.get("datum", "").startswith(f"{godina}-{m:02d}"))
        exp_m = sum(e.get("iznos", 0) for e in expenses if e.get("datum", "").startswith(f"{godina}-{m:02d}"))
        profit_m = income_m + extra_m - exp_m
        month_data.append([_month_name(m), f"{income_m:.2f}", f"{extra_m:.2f}", f"-{exp_m:.2f}", f"{profit_m:.2f}"])
    
    mt = Table(month_data, colWidths=[3*cm, 3*cm, 3*cm, 3*cm, 3*cm])
    mt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
    ]))
    elements.append(mt)
    
    doc.build(elements)
    buf.seek(0)
    fname = f"Finansijski_izvjestaj_{godina}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ============================================================================

@api_router.get("/work-logs")
async def list_work_logs(
    company_id: Optional[str] = None,
    kategorija: Optional[str] = None,
    status: Optional[str] = None,
    q: Optional[str] = None,
    limit: int = 200,
    username: str = Depends(get_current_user),
):
    """Lista evidencije rada — filteri: firma, kategorija, status, slobodna pretraga."""
    query: Dict = {}
    if company_id:
        query["company_id"] = company_id
    if kategorija:
        query["kategorija"] = kategorija
    if status:
        query["status"] = status
    if q:
        q_str = q.strip()
        query["$or"] = [
            {"company_naziv": {"$regex": q_str, "$options": "i"}},
            {"napomena": {"$regex": q_str, "$options": "i"}},
            {"period": {"$regex": q_str, "$options": "i"}},
        ]
    
    items = await db.work_logs.find(query, {"_id": 0}).sort("created_at", -1).to_list(limit)
    return items


@api_router.post("/work-logs")
async def create_work_log(req: WorkLogCreate, username: str = Depends(get_current_user)):
    """Kreiraj novu stavku evidencije."""
    if req.kategorija not in WORK_KATEGORIJE:
        raise HTTPException(400, f"Nepoznata kategorija. Dozvoljene: {', '.join(WORK_KATEGORIJE)}")
    if req.status not in WORK_STATUSI:
        raise HTTPException(400, f"Nepoznat status. Dozvoljene: {', '.join(WORK_STATUSI)}")
    
    # Ako je company_id dat ali bez naziva — auto-fetch
    naziv = req.company_naziv
    if req.company_id and not naziv:
        c = await db.companies.find_one({"id": req.company_id}, {"_id": 0, "naziv": 1})
        if c:
            naziv = c.get("naziv", "")
    
    log = {
        "id": str(uuid.uuid4()),
        "company_id": req.company_id or "",
        "company_naziv": naziv,
        "kategorija": req.kategorija,
        "status": req.status,
        "period": req.period or "",
        "napomena": req.napomena or "",
        "iznos": req.iznos,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username,
        "completed_at": datetime.now(timezone.utc).isoformat() if req.status == "zavrseno" else None,
    }
    await db.work_logs.insert_one(dict(log))
    return log


@api_router.patch("/work-logs/{log_id}")
async def update_work_log(log_id: str, req: WorkLogUpdate, username: str = Depends(get_current_user)):
    """Ažuriraj stavku — najčešće status (u_toku → poslato → zavrseno)."""
    existing = await db.work_logs.find_one({"id": log_id}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Evidencija nije pronađena")
    
    updates: Dict = {"updated_at": datetime.now(timezone.utc).isoformat()}
    for field in ["company_id", "company_naziv", "kategorija", "status", "period", "napomena", "iznos"]:
        v = getattr(req, field, None)
        if v is not None:
            updates[field] = v
    
    # Auto-postavi completed_at kad pređe u završeno
    if req.status == "zavrseno" and existing.get("status") != "zavrseno":
        updates["completed_at"] = datetime.now(timezone.utc).isoformat()
    elif req.status and req.status != "zavrseno":
        updates["completed_at"] = None
    
    await db.work_logs.update_one({"id": log_id}, {"$set": updates})
    out = await db.work_logs.find_one({"id": log_id}, {"_id": 0})
    return out


@api_router.delete("/work-logs/{log_id}")
async def delete_work_log(log_id: str, username: str = Depends(get_current_user)):
    r = await db.work_logs.delete_one({"id": log_id})
    if r.deleted_count == 0:
        raise HTTPException(404, "Evidencija nije pronađena")
    return {"success": True}


@api_router.get("/work-logs/stats")
async def work_logs_stats(username: str = Depends(get_current_user)):
    """Brojač po kategoriji i statusu za dashboard."""
    pipeline = [
        {"$group": {"_id": {"kategorija": "$kategorija", "status": "$status"}, "count": {"$sum": 1}}}
    ]
    raw = await db.work_logs.aggregate(pipeline).to_list(200)
    out: Dict = {}
    for item in raw:
        k = item["_id"]["kategorija"]
        s = item["_id"]["status"]
        out.setdefault(k, {"u_toku": 0, "poslato": 0, "zavrseno": 0, "total": 0})
        out[k][s] = item["count"]
        out[k]["total"] += item["count"]
    return out


# ============================================================================
# FOUNDING TEMPLATES — sačuvani šabloni osnivanja firmi (za ponovno korištenje)
# ============================================================================

@api_router.get("/founding/templates")
async def list_founding_templates(username: str = Depends(get_current_user)):
    """Lista sačuvanih šablona za osnivanje firmi."""
    items = await db.founding_templates.find({}, {"_id": 0}).sort("updated_at", -1).to_list(200)
    return items


@api_router.post("/founding/templates")
async def save_founding_template(payload: Dict[str, Any], username: str = Depends(get_current_user)):
    """Sačuvaj podatke o osnivanju kao šablon za ponovno korištenje."""
    name = payload.get("template_name", "").strip()
    if not name:
        raise HTTPException(400, "Naziv šablona je obavezan")
    
    data = payload.get("data", {})
    
    record = {
        "id": str(uuid.uuid4()),
        "template_name": name,
        "data": data,  # cijeli FoundingRequest payload
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username,
    }
    await db.founding_templates.insert_one(dict(record))
    return record


@api_router.delete("/founding/templates/{tpl_id}")
async def delete_founding_template(tpl_id: str, username: str = Depends(get_current_user)):
    r = await db.founding_templates.delete_one({"id": tpl_id})
    if r.deleted_count == 0:
        raise HTTPException(404, "Šablon nije pronađen")
    return {"success": True}


# ============================================================================
# VAULT — Lozinke za tokene i lične karte klijenata
# ============================================================================

class VaultCredential(BaseModel):
    naziv: str  # ime/prezime ILI naziv firme
    tip: str = "licna_karta"  # licna_karta / token / oba
    company_id: Optional[str] = None
    # Lična karta
    pin: Optional[str] = ""
    puk: Optional[str] = ""
    can: Optional[str] = ""
    broj_licne: Optional[str] = ""
    # Token
    token_password: Optional[str] = ""
    token_serial: Optional[str] = ""
    # Datumi
    datum_preuzimanja: Optional[str] = ""
    datum_isteka: Optional[str] = ""
    napomena: Optional[str] = ""


@api_router.get("/vault")
async def list_vault(
    q: Optional[str] = None,
    tip: Optional[str] = None,
    company_id: Optional[str] = None,
    sort_by_expiry: bool = True,
    limit: int = 500,
    username: str = Depends(get_current_user),
):
    """Lista lozinki — sortirano po datumu isteka (najbliži istek prvi)."""
    query: Dict = {}
    if tip:
        query["tip"] = tip
    if company_id:
        query["company_id"] = company_id
    if q:
        query["$or"] = [
            {"naziv": {"$regex": q, "$options": "i"}},
            {"broj_licne": {"$regex": q, "$options": "i"}},
            {"napomena": {"$regex": q, "$options": "i"}},
        ]
    
    items = await db.vault_credentials.find(query, {"_id": 0}).to_list(limit)
    
    # Sortiranje po datumu isteka (oni koji ističu prvi su gore)
    if sort_by_expiry:
        def sort_key(it):
            d = it.get("datum_isteka") or ""
            return d if d else "9999-99-99"
        items.sort(key=sort_key)
    
    # Dodaj polje is_expiring (2 mjeseca prije isteka)
    today = datetime.now(timezone.utc).date()
    for it in items:
        d_str = it.get("datum_isteka") or ""
        it["is_expired"] = False
        it["is_expiring"] = False
        it["days_to_expiry"] = None
        if d_str:
            try:
                d = datetime.fromisoformat(d_str.replace('Z', '')).date()
                delta = (d - today).days
                it["days_to_expiry"] = delta
                if delta < 0:
                    it["is_expired"] = True
                elif delta <= 60:
                    it["is_expiring"] = True
            except Exception:
                pass
    
    return items


@api_router.post("/vault")
async def create_vault(req: VaultCredential, username: str = Depends(get_current_user)):
    if not req.naziv.strip():
        raise HTTPException(400, "Naziv (ime/prezime ili firma) je obavezan")
    
    record = {
        "id": str(uuid.uuid4()),
        "naziv": req.naziv.strip(),
        "tip": req.tip,
        "company_id": req.company_id or "",
        "pin": req.pin or "",
        "puk": req.puk or "",
        "can": req.can or "",
        "broj_licne": req.broj_licne or "",
        "token_password": req.token_password or "",
        "token_serial": req.token_serial or "",
        "datum_preuzimanja": req.datum_preuzimanja or "",
        "datum_isteka": req.datum_isteka or "",
        "napomena": req.napomena or "",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username,
    }
    await db.vault_credentials.insert_one(dict(record))
    return record


@api_router.patch("/vault/{vid}")
async def update_vault(vid: str, req: VaultCredential, username: str = Depends(get_current_user)):
    existing = await db.vault_credentials.find_one({"id": vid}, {"_id": 0})
    if not existing:
        raise HTTPException(404, "Lozinka nije pronađena")
    
    updates = {
        "naziv": req.naziv.strip(),
        "tip": req.tip,
        "company_id": req.company_id or "",
        "pin": req.pin or "",
        "puk": req.puk or "",
        "can": req.can or "",
        "broj_licne": req.broj_licne or "",
        "token_password": req.token_password or "",
        "token_serial": req.token_serial or "",
        "datum_preuzimanja": req.datum_preuzimanja or "",
        "datum_isteka": req.datum_isteka or "",
        "napomena": req.napomena or "",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.vault_credentials.update_one({"id": vid}, {"$set": updates})
    return {**existing, **updates}


@api_router.delete("/vault/{vid}")
async def delete_vault(vid: str, username: str = Depends(get_current_user)):
    r = await db.vault_credentials.delete_one({"id": vid})
    if r.deleted_count == 0:
        raise HTTPException(404, "Lozinka nije pronađena")
    return {"success": True}


@api_router.post("/vault/bulk-import")
async def bulk_import_vault(payload: Dict[str, Any], username: str = Depends(get_current_user)):
    """Bulk import — prima listu objekata sa istim poljima kao VaultCredential."""
    items = payload.get("items") or []
    if not isinstance(items, list):
        raise HTTPException(400, "items mora biti lista")
    
    created = 0
    for it in items:
        if not (it.get("naziv") or "").strip():
            continue
        record = {
            "id": str(uuid.uuid4()),
            "naziv": (it.get("naziv") or "").strip(),
            "tip": it.get("tip") or "licna_karta",
            "company_id": it.get("company_id") or "",
            "pin": it.get("pin") or "",
            "puk": it.get("puk") or "",
            "can": it.get("can") or "",
            "broj_licne": it.get("broj_licne") or "",
            "token_password": it.get("token_password") or "",
            "token_serial": it.get("token_serial") or "",
            "datum_preuzimanja": it.get("datum_preuzimanja") or "",
            "datum_isteka": it.get("datum_isteka") or "",
            "napomena": it.get("napomena") or "",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "created_by": username,
        }
        await db.vault_credentials.insert_one(dict(record))
        created += 1
    
    return {"success": True, "created": created}



@api_router.post("/founding/generate")
async def generate_founding(req: FoundingRequest, username: str = Depends(get_current_user)):
    """Generiše 4 dokumenta za osnivanje DOO firme (Odluka, Imenovanje, Saglasnost, Statut)."""
    templates_map = [
        ("Odluka o osnivanju DOO.docx", "odluka_o_osnivanju"),
        ("ODLUKA O IMENOVANJE DIREKTORA.docx", "imenovanje_direktora"),
        ("SAGLASNOST SA IMENOVANJEM.docx", "saglasnost_imenovanjem"),
        ("statut 2026.docx", "statut"),
    ]
    
    replacements = _build_founding_replacements(req)
    short_name = re.sub(r'[^\w\s-]', '', req.firma_naziv_skraceni)[:25].strip().replace(' ', '_')
    output_files = []
    
    for tpl_filename, slug in templates_map:
        tpl_path = TEMPLATES_DIR / tpl_filename
        if not tpl_path.exists():
            continue
        
        doc = Document(str(tpl_path))
        _docx_replace(doc, replacements)
        _remove_yellow_highlights(doc)
        
        output_filename = f"{uuid.uuid4().hex[:8]}_{slug}_{short_name}.docx"
        output_filename = re.sub(r'[^\w\s.-]', '_', output_filename).replace(' ', '_')
        output_path = GENERATED_DIR / output_filename
        doc.save(str(output_path))
        
        # Konvertuj u PDF
        pdf_filename = output_filename.replace('.docx', '.pdf')
        _convert_to_pdf(output_path)
        
        # Snimi u history
        record = {
            "id": str(uuid.uuid4()),
            "filename": output_filename,
            "pdf_filename": pdf_filename,
            "template": tpl_filename,
            "template_filename": tpl_filename,
            "company_id": "osnivanje",  # virtual company ID for founding docs
            "company_naziv": req.firma_naziv_pun,
            "employee_id": None,
            "employee_naziv": req.osnivac_ime_prezime,
            "custom_fields": req.model_dump(),
            "created_at": datetime.now(timezone.utc).isoformat(),
            "created_by": username,
            "kategorija": "Osnivanje DOO",
        }
        await db.generated_documents.insert_one(dict(record))
        
        output_files.append({
            "label": tpl_filename.replace('.docx', ''),
            "slug": slug,
            "filename": output_filename,
            "pdf_filename": pdf_filename,
            "download_url": f"/api/documents/download/{output_filename}",
            "preview_url": f"/api/documents/preview/{pdf_filename}",
        })
    
    # Auto-kreiraj WorkLog entry za "Osnivanje DOO"
    work_log = {
        "id": str(uuid.uuid4()),
        "company_id": "",  # firma još nije osnovana
        "company_naziv": req.firma_naziv_pun,
        "kategorija": "osnivanje",
        "status": "u_toku",
        "period": "",
        "napomena": f"Osnivač: {req.osnivac_ime_prezime}, generisano 4 dokumenta",
        "iznos": None,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username,
        "completed_at": None,
    }
    await db.work_logs.insert_one(dict(work_log))
    
    return {
        "success": True,
        "firma_naziv": req.firma_naziv_pun,
        "files": output_files,
        "work_log_id": work_log["id"],
    }


# ============================================================================
# SPECIJALNO PUNOMOĆJE — generisanje dokumenta sa zamjenom žutih markera
# ============================================================================

def _build_punomoce_replacements(req: 'PunomoceRequest') -> Dict[str, str]:
    """Mapiranja za Specijalno punomoćje (template: SPECIJALNO PUNOMOCJE.docx).
    
    Hardkodirane sample vrijednosti u template-u koje se zamjenjuju:
      - Davalac: ARTA RESULBEGU-MURTEZA, JMB:1107981225018, iz Crne Gore
      - Punomoćnik: Getuard Cekoviq, JMB: 0806994223008
      - Firma: DOO "ULCINJ GRADNJA" Ulcinj  (sa " i „ navodnicima!)
      - Datum: 13.05. 2026
      - Grad: ULCINJ
    """
    # Datum
    if req.datum:
        try:
            dt = datetime.fromisoformat(req.datum.replace('Z', ''))
            datum_str = dt.strftime("%d.%m.%Y")
            datum_str_spaced = dt.strftime("%d.%m. %Y")  # template ima "13.05. 2026"
        except Exception:
            datum_str = req.datum
            datum_str_spaced = req.datum
    else:
        now = datetime.now(timezone.utc)
        datum_str = now.strftime("%d.%m.%Y")
        datum_str_spaced = now.strftime("%d.%m. %Y")
    
    # Davalac - isprava (JMBG ili pasoš)
    if req.davaoc_is_stranac:
        davaoc_label = "Br. pasoša"
        davaoc_id = req.davaoc_pasos or "____________"
    else:
        davaoc_label = "JMB"
        davaoc_id = req.davaoc_jmb or "____________"
    
    # Punomoćnik - obavezno
    pun_ime = req.punomocnik_ime_prezime or "____________"
    pun_jmb = req.punomocnik_jmb or "____________"
    
    # Firma
    firma = req.firma_naziv or "____________"
    
    # Grad iz adrese (uzima riječ poslije zadnje zapete ili poslednju riječ)
    davaoc_grad = "ULCINJ"
    if req.davaoc_adresa:
        adr = req.davaoc_adresa.strip()
        if ',' in adr:
            davaoc_grad = adr.split(',')[-1].strip().upper()
        else:
            # uzmi poslednju riječ
            parts = adr.split()
            if parts:
                davaoc_grad = parts[-1].upper()
    
    repl = {
        # === Davalac (Punomoćodavac) ===
        # Originalni text: "Ja, ARTA RESULBEGU-MURTEZA iz Crne Gore, JMB:1107981225018"
        'ARTA RESULBEGU-MURTEZA': req.davaoc_ime_prezime,
        # JMB davaoca
        'JMB:1107981225018': f"{davaoc_label}:{davaoc_id}",
        'JMB: 1107981225018': f"{davaoc_label}: {davaoc_id}",
        '1107981225018': davaoc_id,
        # Država davaoca
        'iz Crne Gore': f'iz {req.davaoc_drzava}',
        'iz  Crne Gore': f'iz {req.davaoc_drzava}',
        
        # === Punomoćnik ===
        # Originalni text: "ovlašćujem  Getuard Cekoviq JMB: 0806994223008"
        'Getuard Cekoviq': pun_ime,
        'GETUARD CEKOVIQ': pun_ime,
        'JMB: 0806994223008': f"JMB: {pun_jmb}",
        'JMB:0806994223008': f"JMB:{pun_jmb}",
        '0806994223008': pun_jmb,
        
        # === Naziv firme — tri pojavljivanja sa različitim navodnicima ===
        # P3, P4: DOO "ULCINJ GRADNJA" Ulcinj (sa " navodnicima)
        # P14: DOO „ULCINJ GRADNJA" Ulcinj (sa „ navodnikom - njemački otvarač)
        'DOO “ULCINJ GRADNJA“ Ulcinj': firma,
        'DOO „ULCINJ GRADNJA" Ulcinj': firma,
        'DOO "ULCINJ GRADNJA" Ulcinj': firma,
        'DOO „ULCINJ GRADNJA“ Ulcinj': firma,
        # Fallback bez "DOO" prefiksa
        '"ULCINJ GRADNJA"': f'"{firma}"' if not firma.startswith('"') and not firma.lower().startswith('doo') else firma,
        '„ULCINJ GRADNJA"': firma,
        '„ULCINJ GRADNJA“': firma,
        '“ULCINJ GRADNJA“': firma,
        'ULCINJ GRADNJA': firma.replace('DOO', '').replace('"', '').replace('„', '').replace('“', '').strip(),
        
        # === Datum ===
        '13.05. 2026': datum_str_spaced,
        '13.05.2026': datum_str,
        'dana 13.05. 2026': f'dana {datum_str_spaced}',
        'dana 13.05.2026': f'dana {datum_str}',
        
        # === Grad potpisa (P17: "U ULCINJ, dana...") ===
        'U ULCINJ, dana': f'U {davaoc_grad}, dana',
    }
    return repl


@api_router.post("/punomoce/generate")
async def generate_punomoce(req: PunomoceRequest, username: str = Depends(get_current_user)):
    """Generiše Specijalno punomoćje DOCX + PDF."""
    template_filename = "SPECIJALNO PUNOMOCJE.docx"
    tpl_path = TEMPLATES_DIR / template_filename
    if not tpl_path.exists():
        # Pokušaj alternativne nazive
        for alt in ["specijalno punomocje.docx", "Specijalno punomoćje.docx", "PUNOMOCJE.docx", "Punomocje.docx"]:
            ap = TEMPLATES_DIR / alt
            if ap.exists():
                tpl_path = ap
                template_filename = alt
                break
        else:
            raise HTTPException(
                404,
                "Šablon 'SPECIJALNO PUNOMOCJE.docx' nije pronađen. "
                "Molim te pošalji mi .docx fajl sa žutim oznakama gdje treba popuniti podatke."
            )
    
    replacements = _build_punomoce_replacements(req)
    
    doc = Document(str(tpl_path))
    _docx_replace(doc, replacements)
    _remove_yellow_highlights(doc)
    
    short_name = re.sub(r'[^\w\s-]', '', req.davaoc_ime_prezime or "Punomoce")[:25].strip().replace(' ', '_')
    output_filename = f"{uuid.uuid4().hex[:8]}_specijalno_punomocje_{short_name}.docx"
    output_filename = re.sub(r'[^\w\s.-]', '_', output_filename).replace(' ', '_')
    output_path = GENERATED_DIR / output_filename
    doc.save(str(output_path))
    
    # Konvertuj u PDF
    pdf_filename = output_filename.replace('.docx', '.pdf')
    _convert_to_pdf(output_path)
    
    # Snimi u history
    record = {
        "id": str(uuid.uuid4()),
        "filename": output_filename,
        "pdf_filename": pdf_filename,
        "template": template_filename,
        "template_filename": template_filename,
        "company_id": "punomocje",
        "company_naziv": req.firma_naziv or req.davaoc_ime_prezime,
        "employee_id": None,
        "employee_naziv": req.davaoc_ime_prezime,
        "custom_fields": req.model_dump(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": username,
        "kategorija": "Specijalno punomoćje",
    }
    await db.generated_documents.insert_one(dict(record))
    
    return {
        "success": True,
        "filename": output_filename,
        "pdf_filename": pdf_filename,
        "download_url": f"/api/documents/download/{output_filename}",
        "preview_url": f"/api/documents/preview/{pdf_filename}",
    }


@api_router.get("/documents/preview/{filename}")
async def preview_document(filename: str, token: Optional[str] = None):
    """Vraća PDF inline u browseru (za prikaz/štampu bez downloada).
    Podržava i generisane dokumente i izvorne PDF/DOCX šablone iz templates foldera.
    """
    if token:
        try:
            jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        except Exception:
            raise HTTPException(401, "Neispravan token")
    
    safe_name = os.path.basename(filename)
    
    # 1. Try GENERATED folder
    file_path = GENERATED_DIR / safe_name
    
    # 2. If PDF doesn't exist there, try to generate from sibling .docx (auto-create PDF)
    if not file_path.exists() and safe_name.endswith('.pdf'):
        docx_sibling = GENERATED_DIR / safe_name.replace('.pdf', '.docx')
        if docx_sibling.exists():
            _convert_to_pdf(docx_sibling)
    
    # 3. If still not found, try TEMPLATES folder (for source PDF/DOCX templates)
    if not file_path.exists():
        file_path = TEMPLATES_DIR / safe_name
        
        # If a .docx template, generate PDF copy in TEMPLATES folder (cached)
        if not file_path.exists() and safe_name.endswith('.pdf'):
            docx_in_templates = TEMPLATES_DIR / safe_name.replace('.pdf', '.docx')
            if docx_in_templates.exists():
                _convert_to_pdf(docx_in_templates)
                file_path = TEMPLATES_DIR / safe_name
    
    if not file_path.exists():
        # Provjeri da li je razlog što LibreOffice fali
        import shutil
        if safe_name.endswith('.pdf') and not shutil.which("soffice"):
            raise HTTPException(503, "LibreOffice nije dostupan na serveru — konverzija u PDF nije uspjela. Molim restartujte ili kontaktirajte podršku.")
        raise HTTPException(404, "Dokument nije pronađen")
    
    # ASCII-safe filename za HTTP header (RFC 5987 za UTF-8)
    from urllib.parse import quote
    ascii_name = safe_name.encode("ascii", "replace").decode("ascii").replace("?", "_")
    utf8_name = quote(safe_name)
    return FileResponse(
        path=str(file_path),
        filename=ascii_name,
        media_type='application/pdf',
        headers={
            "Content-Disposition": f"inline; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"
        }
    )

@api_router.get("/documents/download/{filename}")
async def download_document(filename: str, token: Optional[str] = None):
    # Allow token via query param for download
    if token:
        try:
            jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        except Exception:
            raise HTTPException(401, "Neispravan token")
    
    safe_name = os.path.basename(filename)
    file_path = GENERATED_DIR / safe_name
    if not file_path.exists():
        # Try templates folder
        file_path = TEMPLATES_DIR / safe_name
        if not file_path.exists():
            raise HTTPException(404, "Fajl nije pronađen")
    
    from urllib.parse import quote
    ascii_name = safe_name.encode("ascii", "replace").decode("ascii").replace("?", "_")
    utf8_name = quote(safe_name)
    return FileResponse(
        path=str(file_path),
        filename=ascii_name,
        media_type='application/octet-stream',
        headers={
            "Content-Disposition": f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"
        }
    )

@api_router.get("/documents")
async def list_generated_documents(
    company_id: Optional[str] = None,
    username: str = Depends(get_current_user)
):
    query: Dict[str, Any] = {}
    if company_id:
        query["company_id"] = company_id
    docs = await db.generated_documents.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    return docs


# ============== PDV / IOPPD TRACKING ==============

@api_router.get("/pdv-tracking")
async def get_pdv_tracking(year: int, month: int, username: str = Depends(get_current_user)):
    """Vraća listu svih PDV i IOPPD obveznika sa statusom predaje za dati mjesec."""
    companies = await db.companies.find(
        {"$or": [{"pdv_obveznik": True}, {"ioppd_obveznik": True}], "aktivna": True},
        {"_id": 0}
    ).sort("naziv", 1).to_list(1000)
    
    records = await db.pdv_records.find(
        {"year": year, "month": month}, {"_id": 0}
    ).to_list(1000)
    records_map = {r["company_id"]: r for r in records}
    
    result = []
    for c in companies:
        rec = records_map.get(c["id"], {})
        # Derive status — ako je predato=True, status="predato" (backward compat)
        pdv_status = rec.get("pdv_status") or ("predato" if rec.get("pdv_predato") else "ceka")
        ioppd_status = rec.get("ioppd_status") or ("predato" if rec.get("ioppd_predato") else "ceka")
        result.append({
            "company_id": c["id"],
            "company_naziv": c["naziv"],
            "pib": c["pib"],
            "pdv_obveznik": c.get("pdv_obveznik", False),
            "ioppd_obveznik": c.get("ioppd_obveznik", False),
            "pdv_predato": rec.get("pdv_predato", False),
            "pdv_datum": rec.get("pdv_datum", ""),
            "pdv_broj": rec.get("pdv_broj", ""),
            "pdv_status": pdv_status,
            "ioppd_predato": rec.get("ioppd_predato", False),
            "ioppd_datum": rec.get("ioppd_datum", ""),
            "ioppd_broj": rec.get("ioppd_broj", ""),
            "ioppd_status": ioppd_status,
        })
    return result

@api_router.put("/pdv-tracking/{company_id}")
async def update_pdv_tracking(
    company_id: str,
    year: int,
    month: int,
    data: PDVUpdate,
    username: str = Depends(get_current_user)
):
    existing = await db.pdv_records.find_one(
        {"company_id": company_id, "year": year, "month": month}, {"_id": 0}
    )
    
    update_dict = data.model_dump(exclude_none=True)
    
    # Auto-fill today's date when checking
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    if update_dict.get("pdv_predato") is True and not update_dict.get("pdv_datum"):
        if not existing or not existing.get("pdv_datum"):
            update_dict["pdv_datum"] = today
    if update_dict.get("ioppd_predato") is True and not update_dict.get("ioppd_datum"):
        if not existing or not existing.get("ioppd_datum"):
            update_dict["ioppd_datum"] = today
    
    # Sinhronizacija status ↔ predato
    if "pdv_status" in update_dict:
        update_dict["pdv_predato"] = (update_dict["pdv_status"] == "predato")
        if update_dict["pdv_predato"] and not update_dict.get("pdv_datum"):
            if not existing or not existing.get("pdv_datum"):
                update_dict["pdv_datum"] = today
    elif "pdv_predato" in update_dict:
        update_dict["pdv_status"] = "predato" if update_dict["pdv_predato"] else "ceka"
    
    if "ioppd_status" in update_dict:
        update_dict["ioppd_predato"] = (update_dict["ioppd_status"] == "predato")
        if update_dict["ioppd_predato"] and not update_dict.get("ioppd_datum"):
            if not existing or not existing.get("ioppd_datum"):
                update_dict["ioppd_datum"] = today
    elif "ioppd_predato" in update_dict:
        update_dict["ioppd_status"] = "predato" if update_dict["ioppd_predato"] else "ceka"
    
    if existing:
        await db.pdv_records.update_one(
            {"company_id": company_id, "year": year, "month": month},
            {"$set": update_dict}
        )
    else:
        record = PDVRecord(
            company_id=company_id, year=year, month=month, **update_dict
        ).model_dump()
        await db.pdv_records.insert_one(record)
    
    return {"success": True}


# ============== DASHBOARD STATS ==============

@api_router.get("/stats")
async def get_stats(username: str = Depends(get_current_user)):
    total_companies = await db.companies.count_documents({})
    active_companies = await db.companies.count_documents({"aktivna": True})
    pdv_count = await db.companies.count_documents({"pdv_obveznik": True, "aktivna": True})
    ioppd_count = await db.companies.count_documents({"ioppd_obveznik": True, "aktivna": True})
    total_employees = await db.employees.count_documents({})
    
    # Documents this month
    now = datetime.now(timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
    docs_this_month = await db.generated_documents.count_documents({"created_at": {"$gte": month_start}})
    
    return {
        "total_companies": total_companies,
        "active_companies": active_companies,
        "pdv_count": pdv_count,
        "ioppd_count": ioppd_count,
        "total_employees": total_employees,
        "docs_this_month": docs_this_month,
    }


@api_router.get("/")
async def root():
    return {"app": "Getuard Agency", "version": "1.0.0"}


@api_router.get("/health")
async def health_check():
    """Health endpoint za Railway / Render / uptime monitoring."""
    try:
        # Provjeri da li MongoDB radi
        await db.command("ping")
        return {"status": "ok", "db": "connected"}
    except Exception as e:
        return {"status": "degraded", "db": "error", "error": str(e)[:100]}


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
