"""Verify Statut template contains no 'izvršni direktor' variants."""
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document

TEMPLATE_PATH = Path("/app/backend/templates/statut 2026.docx")


@pytest.fixture(scope="module")
def template_full_xml():
    """Extract all XML content (document + headers + footers + tables) from docx."""
    assert TEMPLATE_PATH.exists(), f"Template not found: {TEMPLATE_PATH}"
    parts = []
    with zipfile.ZipFile(TEMPLATE_PATH, "r") as z:
        for name in z.namelist():
            if name.endswith(".xml"):
                parts.append(z.read(name).decode("utf-8", errors="ignore"))
    return "\n".join(parts)


@pytest.fixture(scope="module")
def doc_paragraphs():
    doc = Document(str(TEMPLATE_PATH))
    paragraphs = [p.text for p in doc.paragraphs]
    # Include table cell paragraphs too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paragraphs.append(p.text)
    return paragraphs


class TestStatutIzvrsniDirektor:
    def test_no_izvrsn_substring_in_xml(self, template_full_xml):
        """No occurrence of [Ii]zvr[šs]n anywhere in the docx XML."""
        pattern = re.compile(r"[Ii]zvr[šs]n")
        matches = pattern.findall(template_full_xml)
        assert not matches, f"Found {len(matches)} occurrences of 'izvršn' variants in template XML: {matches[:10]}"

    def test_no_izvrsn_in_paragraphs(self, doc_paragraphs):
        pattern = re.compile(r"[Ii]zvr[šs]n")
        offending = [(i, p) for i, p in enumerate(doc_paragraphs) if pattern.search(p)]
        assert not offending, f"Paragraphs still containing 'izvršn': {offending[:5]}"

    def test_direktora_drustva_phrase_present(self, doc_paragraphs):
        """Was 'Izvršnog direktora Društva' → should now contain 'direktora Društva'."""
        joined = "\n".join(doc_paragraphs)
        assert "direktora Društva" in joined, "Expected phrase 'direktora Društva' not found"

    def test_direktor_mora_osigurati_phrase_present(self, doc_paragraphs):
        """Was 'izvršni direktor mora osigurati' → should now contain 'direktor mora osigurati'."""
        joined = "\n".join(doc_paragraphs)
        assert "direktor mora osigurati" in joined, "Expected phrase 'direktor mora osigurati' not found"

    def test_direktor_word_still_present(self, doc_paragraphs):
        """Sanity: 'direktor' (in some form) still exists in the template."""
        joined = "\n".join(doc_paragraphs).lower()
        assert "direktor" in joined, "Template should still reference 'direktor'"
