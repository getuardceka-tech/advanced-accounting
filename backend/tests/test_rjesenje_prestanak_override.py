"""
Tests for RJESENJE O PRESTANKU RADNOG ODNOSA KAD ISTICE UGOVOR O RADU
- custom_fields.datum_prestanka_override priority
- employee.datum_prestanka fallback
- no dates -> today's date fallback
"""
import io
import os
import re
from datetime import datetime
from pathlib import Path

import pytest
import requests
from docx import Document
from dotenv import dotenv_values

frontend_env = dotenv_values("/app/frontend/.env")
base_url = os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")
if not base_url:
    raise RuntimeError("REACT_APP_BACKEND_URL missing")
BASE_URL = base_url.rstrip("/")

TEMPLATE = "RJESENJE O PRESTANKU RADNOG ODNOSA KAD ISTICE UGOVOR O RADU.docx"
COMPANY_ID = "3b295e48-3981-4edb-9d0c-5bb8aca93ebc"
EMPLOYEE_ID = "63283aaa-551b-4d67-9033-c946b336a610"
SAMPLE_DATES = ["31.03.2026", "28.02.2026"]


def _creds():
    content = Path("/app/memory/test_credentials.md").read_text(encoding="utf-8")
    user = re.search(r'(?im)^\s*(?:[-*]\s*)?(?:\*\*)?(?:Korisni[^:]*|username)(?:\*\*)?\s*:\s*`?([^`\s]+)', content)
    pwd = re.search(r'(?im)^\s*(?:[-*]\s*)?(?:\*\*)?(?:Lozinka|password)(?:\*\*)?\s*:\s*`?([^`\s]+)', content)
    if not user or not pwd:
        pytest.skip("credentials not found")
    return user.group(1), pwd.group(1)


@pytest.fixture(scope="module")
def client():
    s = requests.Session()
    user, pwd = _creds()
    r = s.post(f"{BASE_URL}/api/auth/login", json={"username": user, "password": pwd}, timeout=30)
    if r.status_code != 200:
        pytest.fail(f"Login failed {r.status_code}: {r.text[:300]}")
    token = r.json().get("access_token") or r.json().get("token")
    if not token:
        pytest.fail(f"No token in login response: {r.text[:300]}")
    s.headers.update({"Authorization": f"Bearer {token}", "Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def created_employee_ids():
    return []


@pytest.fixture(scope="module", autouse=True)
def cleanup(client, created_employee_ids):
    yield
    for eid in created_employee_ids:
        client.delete(f"{BASE_URL}/api/employees/{eid}", timeout=30)


def _company_id(client):
    r = client.get(f"{BASE_URL}/api/companies/{COMPANY_ID}", timeout=30)
    if r.status_code == 200:
        return COMPANY_ID
    r = client.get(f"{BASE_URL}/api/companies", timeout=30)
    assert r.status_code == 200, r.text[:300]
    comps = r.json()
    assert comps, "No companies available for testing"
    return comps[0]["id"]


def _generate_and_read(client, payload):
    r = client.post(f"{BASE_URL}/api/documents/generate", json=payload, timeout=90)
    assert r.status_code == 200, f"generate failed {r.status_code}: {r.text[:500]}"
    data = r.json()
    fn = data.get("filename")
    assert fn, f"No filename in response: {data}"
    d = client.get(f"{BASE_URL}/api/documents/download/{fn}", timeout=60)
    assert d.status_code == 200, f"download failed {d.status_code}"
    doc = Document(io.BytesIO(d.content))
    texts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return texts, doc


def _assert_no_sample_dates(texts):
    for sd in SAMPLE_DATES:
        hits = [t for t in texts if sd in t]
        assert not hits, f"Sample date {sd} still present: {hits}"


# --- Case 1: custom override wins ---
def test_override_from_modal(client):
    cid = _company_id(client)
    payload = {
        "template_filename": TEMPLATE,
        "company_id": cid,
        "employee_id": EMPLOYEE_ID,
        "custom_fields": {"datum_prestanka_override": "2026-12-31"},
    }
    r = client.get(f"{BASE_URL}/api/employees/{EMPLOYEE_ID}", timeout=30)
    if r.status_code != 200:
        # create fallback employee
        cr = client.post(f"{BASE_URL}/api/employees", json={
            "company_id": cid, "ime": "TEST", "prezime": "PRESTANAK",
            "jmbg": "1234567890123", "pozicija": "pomocni radnik",
            "vrsta_ugovora": "odredjeno", "datum_pocetka": "2026-01-01",
            "datum_kraja": "2026-06-30", "datum_prestanka": "2026-06-30",
        }, timeout=30)
        assert cr.status_code in (200, 201), cr.text[:300]
        payload["employee_id"] = cr.json()["id"]
    texts, _ = _generate_and_read(client, payload)
    _assert_no_sample_dates(texts)
    hits = [i for i, t in enumerate(texts) if "31.12.2026" in t]
    assert hits, f"Override date 31.12.2026 not found. Paragraphs: {texts[:12]}"
    assert 3 in hits, f"Paragraph 3 missing override date: {texts[3] if len(texts) > 3 else None}"
    assert 7 in hits, f"Paragraph 7 missing override date: {texts[7] if len(texts) > 7 else None}"


# --- Case 2: no override -> employee.datum_prestanka ---
def test_employee_datum_prestanka_used(client, created_employee_ids):
    cid = _company_id(client)
    cr = client.post(f"{BASE_URL}/api/employees", json={
        "company_id": cid, "ime": "TEST", "prezime": "PRESTANAKDVA",
        "jmbg": "1234567890124", "pozicija": "pomocni radnik",
        "vrsta_ugovora": "odredjeno", "datum_pocetka": "2026-01-01",
        "datum_kraja": "2026-09-30", "datum_prestanka": "2026-08-15",
    }, timeout=30)
    assert cr.status_code in (200, 201), cr.text[:300]
    eid = cr.json()["id"]
    created_employee_ids.append(eid)

    texts, _ = _generate_and_read(client, {
        "template_filename": TEMPLATE, "company_id": cid,
        "employee_id": eid, "custom_fields": {},
    })
    _assert_no_sample_dates(texts)
    assert any("15.08.2026" in t for t in texts), f"employee.datum_prestanka not used. p3={texts[3]!r} p7={texts[7]!r}"
    # datum_kraja should NOT override the prestanak placeholders
    assert not any("30.09.2026" in t and "zaključno sa danom" in t for t in texts)


# --- Case 3: no override, no datum_prestanka, no datum_kraja -> today ---
def test_fallback_to_today(client, created_employee_ids):
    cid = _company_id(client)
    cr = client.post(f"{BASE_URL}/api/employees", json={
        "company_id": cid, "ime": "TEST", "prezime": "PRESTANAKTRI",
        "jmbg": "1234567890125", "pozicija": "pomocni radnik",
        "vrsta_ugovora": "neodredjeno", "datum_pocetka": "2026-01-01",
        "datum_kraja": "", "datum_prestanka": "",
    }, timeout=30)
    assert cr.status_code in (200, 201), cr.text[:300]
    eid = cr.json()["id"]
    created_employee_ids.append(eid)

    texts, _ = _generate_and_read(client, {
        "template_filename": TEMPLATE, "company_id": cid,
        "employee_id": eid, "custom_fields": {},
    })
    today = datetime.now().strftime("%d.%m.%Y")
    _assert_no_sample_dates(texts)
    assert any(today in t for t in texts), f"Today's date {today} not found. p3={texts[3]!r} p7={texts[7]!r}"


# --- Case 4: empty override string must not break priority ---
def test_empty_override_falls_back_to_employee(client, created_employee_ids):
    cid = _company_id(client)
    cr = client.post(f"{BASE_URL}/api/employees", json={
        "company_id": cid, "ime": "TEST", "prezime": "PRESTANAKCETIRI",
        "jmbg": "1234567890126", "pozicija": "pomocni radnik",
        "vrsta_ugovora": "odredjeno", "datum_pocetka": "2026-01-01",
        "datum_kraja": "2026-10-31", "datum_prestanka": "2026-07-01",
    }, timeout=30)
    assert cr.status_code in (200, 201), cr.text[:300]
    eid = cr.json()["id"]
    created_employee_ids.append(eid)

    texts, _ = _generate_and_read(client, {
        "template_filename": TEMPLATE, "company_id": cid,
        "employee_id": eid, "custom_fields": {"datum_prestanka_override": "   "},
    })
    _assert_no_sample_dates(texts)
    assert any("01.07.2026" in t for t in texts), f"Fallback failed. p3={texts[3]!r} p7={texts[7]!r}"
