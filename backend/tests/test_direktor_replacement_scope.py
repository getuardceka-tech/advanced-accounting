"""
Regression tests for the scoped 'DIREKTOR' replacement fix (server.py ~1642-1661).

Rule under test (iteration_10):
    is_direktor_tpl = ("direktor" in filename) and ("imenovanj" in filename or "ugovor" in filename)
    repl["DIREKTOR"] = employee.pozicija  # only when is_direktor_tpl and employee is not a director

=> Templates that only mention 'imenovanj' (mobing, zastita od pozara, zastita na radu)
   must keep their 'DIREKTOR' signature label untouched.
=> 'IZVRŠNI DIREKTOR' in the Rješenje o prestanku signature block must stay intact.
"""
import io
import os
import re
from pathlib import Path

import pytest
import requests
from docx import Document
from dotenv import dotenv_values

frontend_env = dotenv_values("/app/frontend/.env")
base_url = os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")
if not base_url:
    raise RuntimeError("REACT_APP_BACKEND_URL missing")
BASE_URL = base_url.rstrip("/")

TPL_PRESTANAK = "RJESENJE O PRESTANKU RADNOG ODNOSA KAD ISTICE UGOVOR O RADU.docx"
TPL_IMENOVANJE_DIR = "ODLUKA O IMENOVANJE DIREKTORA.docx"
TPL_POZAR = "odluka o imenovanju odgovornog lica za zastitu od pozara.docx"
TPL_MOBING = "odluka o imenovanje lice za Mobing.docx"
TPL_ZNR = "odluka o imenovanju lica za zastitu na radu.docx"
TPL_UGOVOR_DIR = "UGOVOR O RADU DIREKTOR.docx"


def _creds():
    content = Path("/app/memory/test_credentials.md").read_text(encoding="utf-8")
    user = re.search(r'(?im)^\s*(?:[-*]\s*)?(?:\*\*)?(?:Korisni[^:]*|username)(?:\*\*)?\s*:\s*`?([^`\s]+)', content)
    pwd = re.search(r'(?im)^\s*(?:[-*]\s*)?(?:\*\*)?(?:Lozinka|password)(?:\*\*)?\s*:\s*`?([^`\s]+)', content)
    if not user or not pwd:
        pytest.skip("credentials not found")
    return user.group(1), pwd.group(1)


@pytest.fixture(scope="module")
def client():
    s = requests.Session()
    user, pwd = _creds()
    r = s.post(f"{BASE_URL}/api/auth/login", json={"username": user, "password": pwd}, timeout=30)
    if r.status_code != 200:
        pytest.fail(f"Login failed {r.status_code}: {r.text[:300]}")
    token = r.json().get("access_token") or r.json().get("token")
    if not token:
        pytest.fail(f"No token in login response: {r.text[:300]}")
    s.headers.update({"Authorization": f"Bearer {token}", "Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def company_id(client):
    r = client.get(f"{BASE_URL}/api/companies", timeout=30)
    assert r.status_code == 200, r.text[:300]
    comps = r.json()
    assert comps, "No companies available"
    return comps[0]["id"]


@pytest.fixture(scope="module")
def created_ids():
    return []


@pytest.fixture(scope="module", autouse=True)
def cleanup(client, created_ids):
    yield
    for eid in created_ids:
        client.delete(f"{BASE_URL}/api/employees/{eid}", timeout=30)


def _make_employee(client, company_id, created_ids, prezime, pozicija, jmbg):
    r = client.post(f"{BASE_URL}/api/employees", json={
        "company_id": company_id, "ime": "TEST", "prezime": prezime,
        "jmbg": jmbg, "pozicija": pozicija,
        "vrsta_ugovora": "odredjeno", "datum_pocetka": "2026-01-01",
        "datum_kraja": "2026-09-30", "datum_prestanka": "2026-08-15",
    }, timeout=30)
    assert r.status_code in (200, 201), r.text[:400]
    eid = r.json()["id"]
    created_ids.append(eid)
    return eid


def _generate_texts(client, payload):
    r = client.post(f"{BASE_URL}/api/documents/generate", json=payload, timeout=90)
    assert r.status_code == 200, f"generate failed {r.status_code}: {r.text[:500]}"
    fn = r.json().get("filename")
    assert fn, f"No filename: {r.json()}"
    d = client.get(f"{BASE_URL}/api/documents/download/{fn}", timeout=60)
    assert d.status_code == 200, f"download failed {d.status_code}"
    doc = Document(io.BytesIO(d.content))
    texts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return texts


# ---------------------------------------------------------------- prestanak
def test_prestanak_signature_keeps_izvrsni_direktor(client, company_id, created_ids):
    eid = _make_employee(client, company_id, created_ids, "PRODAVACJEDAN", "PRODAVAC", "1234567890131")
    texts = _generate_texts(client, {
        "template_filename": TPL_PRESTANAK, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    joined = "\n".join(texts)
    assert "IZVRŠNI PRODAVAC" not in joined.upper(), \
        f"Employee position leaked into signature label: {[t for t in texts if 'PRODAVAC' in t.upper()]}"
    assert "IZVRŠNI DIREKTOR" in joined.upper(), \
        f"Signature label 'IZVRŠNI DIREKTOR' missing. Tail paragraphs: {texts[-8:]}"


def test_prestanak_signature_other_position(client, company_id, created_ids):
    eid = _make_employee(client, company_id, created_ids, "KONOBARDVA", "KONOBAR", "1234567890132")
    texts = _generate_texts(client, {
        "template_filename": TPL_PRESTANAK, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    joined = "\n".join(texts).upper()
    assert "IZVRŠNI KONOBAR" not in joined
    assert "IZVRŠNI DIREKTOR" in joined


# ------------------------------------------- direktor + imenovanj (positive)
def test_odluka_imenovanje_direktora_with_director(client, company_id, created_ids):
    eid = _make_employee(client, company_id, created_ids, "DIREKTORTRI", "DIREKTOR", "1234567890133")
    texts = _generate_texts(client, {
        "template_filename": TPL_IMENOVANJE_DIR, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    joined = "\n".join(texts).upper()
    assert joined.strip(), "Generated document is empty"
    assert "ZA DIREKTORA DRUŠTVA IMENUJE SE" in joined, \
        f"Director appointment wording broken. Paragraphs: {texts[:12]}"
    assert "DIREKTOR ZASTUPA DRUŠTVO" in joined, \
        f"'Direktor zastupa Društvo' clause corrupted. Paragraphs: {texts[:15]}"


# --------------------------------------- imenovanj-only templates (negative)
def test_imenovanje_pozar_signature_label(client, company_id, created_ids):
    """'odluka o imenovanju odgovornog lica za zastitu od pozara.docx' has no
    'direktor' in the filename -> standalone 'DIREKTOR' signer label must survive."""
    eid = _make_employee(client, company_id, created_ids, "PRODAVACCETIRI", "PRODAVAC", "1234567890134")
    texts = _generate_texts(client, {
        "template_filename": TPL_POZAR, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    stripped = [t.strip().upper() for t in texts if t.strip()]
    assert "DIREKTOR" in stripped, \
        f"Standalone signature line 'DIREKTOR' was replaced by employee position. Tail: {stripped[-6:]}"
    assert "PRODAVAC" not in stripped, \
        f"Signature line replaced with employee position: {stripped[-6:]}"


def test_imenovanje_mobing_signature_label(client, company_id, created_ids):
    """'odluka o imenovanje lice za Mobing.docx' -> 'DIREKTOR' labels intact."""
    eid = _make_employee(client, company_id, created_ids, "PRODAVACPET", "PRODAVAC", "1234567890135")
    texts = _generate_texts(client, {
        "template_filename": TPL_MOBING, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    joined = "\n".join(texts).upper()
    assert joined.strip(), "Generated document is empty"
    direktor_lines = [t for t in texts if "DIREKTOR" in t.upper()]
    assert len(direktor_lines) >= 2, \
        f"Expected >=2 'DIREKTOR' occurrences (signature labels), got {direktor_lines}"
    assert "– PRODAVAC" not in joined and "- PRODAVAC" not in joined, \
        f"Signer label replaced with employee position: {[t for t in texts if 'PRODAVAC' in t.upper()]}"
    # the trailing signature label line must still read DIREKTOR
    assert any(t.strip().upper().endswith("DIREKTOR") for t in texts), \
        f"No trailing 'DIREKTOR' signature label. Tail: {texts[-6:]}"


def test_imenovanje_zastita_na_radu_signature_label(client, company_id, created_ids):
    """'odluka o imenovanju lica za zastitu na radu.docx' -> 'Direktor' label intact."""
    eid = _make_employee(client, company_id, created_ids, "PRODAVACSEST", "PRODAVAC", "1234567890136")
    texts = _generate_texts(client, {
        "template_filename": TPL_ZNR, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    stripped = [t.strip() for t in texts if t.strip()]
    assert stripped, "Generated document is empty"
    assert any(t.upper() == "DIREKTOR" for t in stripped), \
        f"Standalone 'Direktor' signature label missing/replaced. Tail: {stripped[-6:]}"
    assert not any(t.strip().upper() == "PRODAVAC" for t in stripped), \
        f"Signature label replaced with employee position. Tail: {stripped[-6:]}"


# --------------------------------------- direktor + ugovor (positive scope)
def test_ugovor_o_radu_direktor_generates(client, company_id, created_ids):
    """'UGOVOR O RADU DIREKTOR.docx' matches direktor+ugovor -> replacement applies
    to the 'Na radnom mjestu:DIREKTOR' field; document must not break."""
    eid = _make_employee(client, company_id, created_ids, "PRODAVACSEDAM", "PRODAVAC", "1234567890137")
    texts = _generate_texts(client, {
        "template_filename": TPL_UGOVOR_DIR, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    joined = "\n".join(texts)
    assert joined.strip(), "Generated document is empty"
    assert "Na radnom mjestu:" in joined, f"Position field missing. Paragraphs: {texts[:20]}"
    pos_line = next(t for t in texts if "Na radnom mjestu:" in t)
    assert pos_line.strip().upper().endswith("PRODAVAC"), \
        f"Employee position not mapped into 'Na radnom mjestu:' line -> {pos_line!r}"
    # NOTE (design): this template maps the person to the COMPANY DIRECTOR
    # (SAMPLE_DIRECTORS -> company.direktor_ime), not to the selected employee.
    # Sample director name must not leak.
    assert "JUSUF ELEZAGI" not in joined.upper(), "Sample director name leaked into document"
    assert "EKREM HOT" not in joined.upper(), "Sample employee name leaked"


def test_ugovor_o_radu_direktor_with_director_employee(client, company_id, created_ids):
    """A real DIREKTOR employee -> 'DIREKTOR' wording must stay as is."""
    eid = _make_employee(client, company_id, created_ids, "DIREKTOROSAM", "DIREKTOR", "1234567890138")
    texts = _generate_texts(client, {
        "template_filename": TPL_UGOVOR_DIR, "company_id": company_id,
        "employee_id": eid, "custom_fields": {},
    })
    joined = "\n".join(texts).upper()
    assert "NA RADNOM MJESTU:DIREKTOR" in joined.replace("  ", " "), \
        f"Director position wording broken: {[t for t in texts if 'RADNOM MJESTU' in t.upper()]}"
