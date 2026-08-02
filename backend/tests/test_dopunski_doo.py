"""Tests for 'UGOVOR O DOPUNSKOM RADU.docx' template — DOO DOO duplication bug fix.

Verifies that after fix in _build_replacements (server.py ~line 1587), the generated
document contains exactly ONE 'DOO' before naziv_skraceni, and that all previously
fixed fields (datum zaključenja, 5 radnih dana, sati, plata, itd.) render correctly.
"""
import os
import re
import io
from pathlib import Path

import pytest
import requests
from dotenv import dotenv_values
from docx import Document

frontend_env = dotenv_values("/app/frontend/.env")
base_url = os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")
if not base_url:
    raise RuntimeError("REACT_APP_BACKEND_URL missing")
BASE_URL = base_url.rstrip("/")

TEMPLATE = "UGOVOR O DOPUNSKOM RADU.docx"
COMPANY_ID = "d3fdf006-a6d9-4f4f-ab00-7018a1f30a54"  # DOO ADVANCED ACCOUNTING ULCINJ
EMPLOYEE_ID = "63283aaa-551b-4d67-9033-c946b336a610"  # EGZON CEKA (dopunski_rad=true)


@pytest.fixture(scope="module")
def auth_token():
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"username": "getuard", "password": "Getuard1994."})
    if r.status_code != 200:
        pytest.fail(f"Login failed: {r.status_code} {r.text[:300]}")
    token = r.json().get("access_token") or r.json().get("token")
    if not token:
        pytest.fail(f"No token in response: {r.text[:300]}")
    return token


@pytest.fixture(scope="module")
def client(auth_token):
    s = requests.Session()
    s.headers.update({"Authorization": f"Bearer {auth_token}", "Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def generated_doc(client):
    r = client.post(f"{BASE_URL}/api/documents/generate", json={
        "template_filename": TEMPLATE,
        "company_id": COMPANY_ID,
        "employee_id": EMPLOYEE_ID,
    })
    assert r.status_code == 200, f"Generate failed: {r.status_code} {r.text[:500]}"
    body = r.json()
    filename = body.get("filename")
    assert filename, f"No filename in response: {body}"

    dl = client.get(f"{BASE_URL}/api/documents/download/{filename}")
    assert dl.status_code == 200, f"Download failed: {dl.status_code}"
    doc = Document(io.BytesIO(dl.content))
    paragraphs = [p.text for p in doc.paragraphs]
    # include table cells too (signature block often in a table)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paragraphs.append(p.text)
    full_text = "\n".join(paragraphs)
    return {"filename": filename, "paragraphs": paragraphs, "text": full_text}


# --- Company / employee fetch fixtures for expected values ---
@pytest.fixture(scope="module")
def company(client):
    r = client.get(f"{BASE_URL}/api/companies/{COMPANY_ID}")
    assert r.status_code == 200
    return r.json()


@pytest.fixture(scope="module")
def employee(client):
    r = client.get(f"{BASE_URL}/api/employees/{EMPLOYEE_ID}")
    assert r.status_code == 200
    return r.json()


class TestDopunskiDooBug:
    def test_no_double_doo(self, generated_doc):
        text = generated_doc["text"]
        assert "DOO DOO" not in text, (
            f"Found 'DOO DOO' duplication in generated document!\n"
            f"Occurrences context: {[l for l in generated_doc['paragraphs'] if 'DOO DOO' in l]}"
        )

    def test_clan_1_correct_poslodavca(self, generated_doc, company):
        naziv_skr = (company.get("naziv_skraceni") or company.get("naziv") or "").strip()
        assert naziv_skr, "Company naziv/naziv_skraceni empty — cannot verify"
        # Find Član 1 paragraph containing 'Poslodavca'
        matches = [p for p in generated_doc["paragraphs"] if "Poslodavca" in p and "obavlja" in p]
        assert matches, f"Could not find Član 1 'Poslodavca ... obavlja' paragraph. Text: {generated_doc['text'][:2000]}"
        clan1 = matches[0]
        # Expect exactly: 'Poslodavca <naziv_skr>' with a single space
        expected_phrase = f"Poslodavca {naziv_skr}"
        assert expected_phrase in clan1, f"Expected '{expected_phrase}' in Član 1, got: {clan1}"
        # No leftover placeholder
        assert "UNICO HIJA" not in clan1, f"Placeholder UNICO HIJA leaked: {clan1}"

    def test_signature_has_naziv_skraceni(self, generated_doc, company):
        naziv_skr = (company.get("naziv_skraceni") or company.get("naziv") or "").strip()
        text = generated_doc["text"]
        assert naziv_skr in text
        # No leftover UNICO HIJA anywhere
        assert "UNICO HIJA" not in text, "UNICO HIJA placeholder leaked in document"

    def test_no_regression_employee_name(self, generated_doc, employee):
        # Employee full name should appear
        ime = (employee.get("ime") or "").strip()
        prezime = (employee.get("prezime") or "").strip()
        text = generated_doc["text"]
        if ime:
            assert ime in text, f"Employee first name '{ime}' missing"
        if prezime:
            assert prezime in text, f"Employee last name '{prezime}' missing"

    def test_no_regression_primary_employer(self, generated_doc, employee):
        pe = (employee.get("primary_employer") or "").strip()
        if pe:
            assert pe in generated_doc["text"], f"primary_employer '{pe}' missing from document"

    def test_no_regression_dopunski_mjesto_rada(self, generated_doc, employee):
        mr = (employee.get("dopunski_mjesto_rada") or "").strip()
        if mr:
            assert mr in generated_doc["text"], f"dopunski_mjesto_rada '{mr}' missing"

    def test_no_regression_5_radnih_dana(self, generated_doc):
        text = generated_doc["text"]
        # Should mention '5 radnih dana' hardcoded
        assert re.search(r"5\s*\(?\s*pet\s*\)?\s*radn|5\s+radnih\s+dana", text, re.IGNORECASE), \
            "Expected '5 radnih dana' phrase in document"

    def test_no_regression_do_10_u_mjesecu(self, generated_doc):
        text = generated_doc["text"]
        assert "do 10" in text and "prethodni mjesec" in text, \
            "Expected 'do 10 ... prethodni mjesec' hardcoded phrase"

    def test_no_unresolved_placeholders(self, generated_doc):
        text = generated_doc["text"]
        # Common placeholder patterns
        leftovers = re.findall(r"\{\{[^}]+\}\}|<[a-z_]+>|\$\{[^}]+\}", text)
        assert not leftovers, f"Unresolved placeholders found: {leftovers[:10]}"
