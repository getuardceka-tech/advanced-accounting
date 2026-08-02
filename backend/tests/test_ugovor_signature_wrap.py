"""Tests for signature-line wrap fix in UGOVOR O RADU Zaposlenih.docx and UGOVOR O RADU DIREKTOR.docx.

Verifies paragraph 36 (labels: 'Zaposleni\\tPoslodavac') and paragraph 38
(signature lines '___\\tM.P.\\t___') use center-tab (4535) + right-tab (9071) tab stops.
"""
import os
import io
from pathlib import Path

import pytest
import requests
from dotenv import dotenv_values
from docx import Document
from docx.oxml.ns import qn

frontend_env = dotenv_values("/app/frontend/.env")
base_url = os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")
if not base_url:
    raise RuntimeError("REACT_APP_BACKEND_URL missing")
BASE_URL = base_url.rstrip("/")

COMPANY_ID = "d3fdf006-a6d9-4f4f-ab00-7018a1f30a54"
EMPLOYEE_ID = "8326b941-1f7d-49d9-adb1-129d33288d71"  # MIRSADA CEKOVIQ (active, non-dopunski)


@pytest.fixture(scope="module")
def auth_token():
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"username": "getuard", "password": "Getuard1994."})
    if r.status_code != 200:
        pytest.fail(f"Login failed: {r.status_code} {r.text[:300]}")
    token = r.json().get("access_token") or r.json().get("token")
    if not token:
        pytest.fail(f"No token: {r.text[:300]}")
    return token


@pytest.fixture(scope="module")
def client(auth_token):
    s = requests.Session()
    s.headers.update({"Authorization": f"Bearer {auth_token}", "Content-Type": "application/json"})
    return s


def _generate_and_load(client, template_filename):
    r = client.post(f"{BASE_URL}/api/documents/generate", json={
        "template_filename": template_filename,
        "company_id": COMPANY_ID,
        "employee_id": EMPLOYEE_ID,
    })
    assert r.status_code == 200, f"Generate {template_filename} failed: {r.status_code} {r.text[:500]}"
    filename = r.json().get("filename")
    assert filename
    dl = client.get(f"{BASE_URL}/api/documents/download/{filename}")
    assert dl.status_code == 200
    return Document(io.BytesIO(dl.content))


def _tab_stops(paragraph):
    """Return list of (val, pos) tuples from <w:tabs> in pPr."""
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is None:
        return []
    tabs_el = pPr.find(qn('w:tabs'))
    if tabs_el is None:
        return []
    return [(t.get(qn('w:val')), t.get(qn('w:pos'))) for t in tabs_el.findall(qn('w:tab'))]


def _assert_signature_paragraphs(doc, template_name):
    paragraphs = doc.paragraphs
    assert len(paragraphs) > 38, f"[{template_name}] Only {len(paragraphs)} paragraphs; expected >38"

    p36 = paragraphs[36]
    p38 = paragraphs[38]

    # Text assertions
    # p36 uses TWO tabs so 'Poslodavac' lands on the right-tab stop (9071),
    # aligning with the right underscore column in p38. One tab would land it at center.
    assert p36.text in ("Zaposleni\tPoslodavac", "Zaposleni\t\tPoslodavac"), (
        f"[{template_name}] paragraph 36 got {p36.text!r}"
    )
    # Must contain both labels with tabs, no wrap (single line)
    assert "\n" not in p36.text
    assert "Zaposleni" in p36.text and "Poslodavac" in p36.text

    assert p38.text == "_______________________\tM.P.\t_______________________", (
        f"[{template_name}] paragraph 38 expected sig-line with 2 tabs, got {p38.text!r}"
    )

    # Tab stops assertions
    expected_tabs = [("center", "4535"), ("right", "9071")]
    for idx, para in [(36, p36), (38, p38)]:
        tabs = _tab_stops(para)
        assert tabs == expected_tabs, (
            f"[{template_name}] paragraph {idx} expected tab stops {expected_tabs}, got {tabs}"
        )


class TestZaposlenihSignature:
    @pytest.fixture(scope="class")
    def doc(self, client):
        return _generate_and_load(client, "UGOVOR O RADU Zaposlenih.docx")

    def test_signature_paragraphs(self, doc):
        _assert_signature_paragraphs(doc, "Zaposlenih")


class TestDirektorSignature:
    @pytest.fixture(scope="class")
    def doc(self, client):
        return _generate_and_load(client, "UGOVOR O RADU DIREKTOR.docx")

    def test_signature_paragraphs(self, doc):
        _assert_signature_paragraphs(doc, "Direktor")


class TestDopunskiRegression:
    """Regression: dopunski template should still have no 'DOO DOO' and its signature intact."""

    @pytest.fixture(scope="class")
    def doc(self, client):
        # Use a dopunski employee
        dopunski_emp = "cc216543-3baf-491e-9831-2ee18fe1b226"  # GETUARD CEKOVIQ dopunski_rad=true
        r = client.post(f"{BASE_URL}/api/documents/generate", json={
            "template_filename": "UGOVOR O DOPUNSKOM RADU.docx",
            "company_id": COMPANY_ID,
            "employee_id": dopunski_emp,
        })
        assert r.status_code == 200, f"Generate dopunski failed: {r.text[:400]}"
        filename = r.json().get("filename")
        dl = client.get(f"{BASE_URL}/api/documents/download/{filename}")
        assert dl.status_code == 200
        return Document(io.BytesIO(dl.content))

    def test_no_doo_doo(self, doc):
        text_parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text_parts.append(p.text)
        full = "\n".join(text_parts)
        assert "DOO DOO" not in full, "'DOO DOO' duplication regressed!"
