"""
GETUARD AGENCY - Document generation iteration 2 tests (MSG 292)
Tests:
- Templates count (58: 56 docx + 2 pdf)
- New mappings for various templates (datum behavior, employee fields)
- Newly-converted DOCX from PDF (prijava_zanatstva, brisevi, hrana, voda, kazneno, registracija)
- Obavještenje 1-page A4 fitting
- ANEKS generation regression
"""
import os
import re
from datetime import datetime, timezone, timedelta
from pathlib import Path

import pytest
import requests
from docx import Document

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')
assert BASE_URL, "REACT_APP_BACKEND_URL must be set"
API = f"{BASE_URL}/api"

USERNAME = "getuard"
PASSWORD = "Getuard1994."

GENERATED_DIR = Path("/app/backend/generated")

TODAY = datetime.now().strftime("%d.%m.%Y")


# ---------- Fixtures ----------
@pytest.fixture(scope="module")
def auth_headers():
    r = requests.post(f"{API}/auth/login", json={"username": USERNAME, "password": PASSWORD}, timeout=30)
    assert r.status_code == 200, f"Login failed: {r.text}"
    tok = r.json()["access_token"]
    return {"Authorization": f"Bearer {tok}", "Content-Type": "application/json"}


@pytest.fixture(scope="module")
def test_company(auth_headers):
    """Use an existing company (with ziro_racun) or create one."""
    companies = requests.get(f"{API}/companies", headers=auth_headers).json()
    # Prefer a company with a ziro_racun
    chosen = None
    for c in companies:
        if c.get("ziro_racun"):
            chosen = c
            break
    if not chosen and companies:
        chosen = companies[0]
    if chosen:
        # Patch missing fields
        upd = dict(chosen)
        changed = False
        if not upd.get("ziro_racun"):
            upd["ziro_racun"] = "510-00012345678-90"
            changed = True
        if not upd.get("adresa"):
            upd["adresa"] = "Glavna ulica 1"; changed = True
        if not upd.get("telefon"):
            upd["telefon"] = "+38269000000"; changed = True
        if not upd.get("email"):
            upd["email"] = "test@firma.me"; changed = True
        if not upd.get("sifra_djelatnosti"):
            upd["sifra_djelatnosti"] = "5610"; changed = True
        if changed:
            requests.put(f"{API}/companies/{chosen['id']}", json=upd, headers=auth_headers)
            chosen = requests.get(f"{API}/companies/{chosen['id']}", headers=auth_headers).json()
        return chosen
    # else create
    payload = {
        "naziv": "TEST_DOC_FIRMA_DOO",
        "pib": "98765432",
        "naziv_skraceni": "TEST DOC",
        "adresa": "Glavna ulica 1",
        "grad": "Ulcinj",
        "telefon": "+38269000000",
        "email": "test@firma.me",
        "ziro_racun": "510-00012345678-90",
        "sifra_djelatnosti": "5610",
        "direktor_ime": "Test Direktor",
        "direktor_jmbg": "1234567890123",
        "pdv_obveznik": True,
        "aktivna": True,
    }
    r = requests.post(f"{API}/companies", json=payload, headers=auth_headers)
    assert r.status_code == 200, r.text
    return r.json()


@pytest.fixture(scope="module")
def test_employee_with_datum_kraja(auth_headers, test_company):
    """Ensure there exists an employee with datum_kraja set on test_company."""
    emps = requests.get(f"{API}/employees?company_id={test_company['id']}", headers=auth_headers).json()
    for e in emps:
        if e.get("datum_kraja"):
            return e
    # else create
    payload = {
        "company_id": test_company["id"],
        "ime": "Vesel",
        "prezime": "Suma",
        "jmbg": "0101990123456",
        "pozicija": "Konobar",
        "datum_pocetka": "2025-01-01",
        "datum_kraja": "2026-12-31",
        "plata_neto": 500.0,
        "plata_bruto": 700.0,
        "vrsta_ugovora": "Na određeno",
        "radno_vrijeme": "Puno",
    }
    r = requests.post(f"{API}/employees", json=payload, headers=auth_headers)
    assert r.status_code == 200, r.text
    return r.json()


# ---------- Helpers ----------
def _generate(auth_headers, template_filename, company_id, employee_id=None, custom_fields=None):
    payload = {
        "template_filename": template_filename,
        "company_id": company_id,
    }
    if employee_id:
        payload["employee_id"] = employee_id
    if custom_fields:
        payload["custom_fields"] = custom_fields
    r = requests.post(f"{API}/documents/generate", json=payload, headers=auth_headers, timeout=120)
    assert r.status_code == 200, f"Generate failed for {template_filename}: {r.status_code} {r.text}"
    return r.json()


def _read_docx_text(filename):
    path = GENERATED_DIR / filename
    assert path.exists(), f"Generated docx missing: {path}"
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            try:
                cells = list(row.cells)
            except Exception:
                # malformed pdf2docx table - fall back to raw <w:tc> iteration
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                from docx.table import _Cell
                cells = []
                for tc in row._tr.findall(f'{ns}tc'):
                    try:
                        cells.append(_Cell(tc, tbl))
                    except Exception:
                        pass
            for cell in cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def _pdf_page_count(pdf_filename):
    pdf_path = GENERATED_DIR / pdf_filename
    if not pdf_path.exists() or fitz is None:
        return None
    with fitz.open(str(pdf_path)) as d:
        return d.page_count


# ---------- Tests ----------
class TestAuthAndTemplates:
    def test_login(self):
        r = requests.post(f"{API}/auth/login", json={"username": USERNAME, "password": PASSWORD})
        assert r.status_code == 200

    def test_templates_count_58(self, auth_headers):
        r = requests.get(f"{API}/templates", headers=auth_headers)
        assert r.status_code == 200
        templates = r.json()
        # Expect 58 total
        assert len(templates) == 58, f"Expected 58 templates, got {len(templates)}"
        # Verify 2 PDFs remain
        pdfs = [t for t in templates if t["extension"] == ".pdf"]
        pdf_names = {t["filename"] for t in pdfs}
        assert "OP OBRAZAC.pdf" in pdf_names
        assert "ZAHTJEV ZA ODOBRENJE ZA DUVAN.pdf" in pdf_names
        assert len(pdfs) == 2, f"Expected exactly 2 PDFs, got {len(pdfs)}: {pdf_names}"
        # 56 docx
        docx = [t for t in templates if t["extension"] == ".docx"]
        assert len(docx) == 56, f"Expected 56 .docx, got {len(docx)}"


class TestMsg292DatumBehavior:
    def test_obavjestenje_prigovora_one_page(self, auth_headers, test_company):
        out = _generate(auth_headers, "obavjestenje O NACINU PODNOSENJA PRIGOVORA za TRGOVINU.docx", test_company["id"])
        pdf_filename = out.get("pdf_filename")
        count = _pdf_page_count(pdf_filename)
        if count is None:
            pytest.skip("PDF generation not available; cannot verify page count")
        assert count == 1, f"Expected 1 page PDF, got {count} for {pdf_filename}"

    def test_rjesenje_godisnjeg_odmora_blank_datum(self, auth_headers, test_company):
        out = _generate(auth_headers, "rjesenje o koriscenje godisnjeg odmora.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        # 'Datum:' must be blank (no today's date right after)
        # Look for any literal today's date occurrences that would indicate failure
        # We accept underscores/blank line after "Datum:"
        # Heuristic: ensure TODAY does not appear right after a "Datum:" label
        m = re.search(r"Datum:\s*([^\n]*)", text)
        if m:
            after = m.group(1)
            assert TODAY not in after, f"Datum line should be blank but contains today: '{after}'"

    def test_odluka_blagajnicki_maksimum_today(self, auth_headers, test_company):
        out = _generate(auth_headers, "odluka o blagajnickom maksimumu.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert TODAY in text, f"Expected today's date {TODAY} in generated text"

    def test_odluka_popust_prodavnicu_today(self, auth_headers, test_company):
        out = _generate(auth_headers, "ODLUKA ZA POPUST U PRODAVNICU.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert TODAY in text, f"Expected today's date {TODAY} in generated text"

    def test_odluka_podizanje_novca_ziro_racun(self, auth_headers, test_company):
        out = _generate(auth_headers, "ODLUKA O PODIZANJE NOVCA SA ZIRO RACUNA.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert TODAY in text, f"Expected today's date in document"
        ziro = test_company.get("ziro_racun", "")
        if ziro:
            assert ziro in text, f"Expected ziro_racun '{ziro}' in document text"

    def test_rjesenje_prestanak_uses_datum_kraja(self, auth_headers, test_company, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        out = _generate(
            auth_headers,
            "RJESENJE O PRESTANKU RADNOG ODNOSA KAD ISTICE UGOVOR O RADU.docx",
            test_company["id"],
            employee_id=emp["id"],
        )
        text = _read_docx_text(out["filename"])
        # Format datum_kraja to dd.mm.yyyy
        dk = emp.get("datum_kraja", "")
        try:
            dk_fmt = datetime.fromisoformat(dk.replace("Z", "")).strftime("%d.%m.%Y")
        except Exception:
            dk_fmt = dk
        assert dk_fmt in text, f"Expected employee datum_kraja '{dk_fmt}' in text"
        # Original template placeholders 31.03.2026 / 28.02.2026 must NOT remain literally
        assert "31.03.2026" not in text or dk_fmt == "31.03.2026", "Template literal 31.03.2026 not replaced"

    def test_obrazlozenje_poreska(self, auth_headers, test_company, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        out = _generate(
            auth_headers,
            "OBRAZLOZENJE ZA PORESKU UPRAVU KAD KASNIMO SA ODJAVAMA.docx",
            test_company["id"],
            employee_id=emp["id"],
        )
        text = _read_docx_text(out["filename"])
        # Header today
        assert TODAY in text, "Expected today's date in header"
        # Employee datum prestanka
        dk = emp.get("datum_kraja", "")
        try:
            dk_fmt = datetime.fromisoformat(dk.replace("Z", "")).strftime("%d.%m.%Y")
        except Exception:
            dk_fmt = dk
        assert dk_fmt in text, f"Expected datum_kraja {dk_fmt} in text"

    def test_pojedinacno_obavjestenje_mobing(self, auth_headers, test_company, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        out = _generate(
            auth_headers,
            "POJEDINACNO OBAVJESTENJE ZAPOSLENIH ZA MOBING.docx",
            test_company["id"],
            employee_id=emp["id"],
        )
        text = _read_docx_text(out["filename"])
        full_name = f"{emp['ime']} {emp['prezime']}"
        assert full_name in text, f"Expected '{full_name}' in text"
        assert emp.get("pozicija", "") in text, f"Expected pozicija '{emp.get('pozicija')}' in text"

    def test_izjava_zaposlenog_upoznavanje(self, auth_headers, test_company, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        out = _generate(
            auth_headers,
            "izjava zaposlenog o upoznavanje sa pravima obavezama i odgovornostima .docx",
            test_company["id"],
            employee_id=emp["id"],
        )
        text = _read_docx_text(out["filename"])
        full_name = f"{emp['ime']} {emp['prezime']}"
        assert full_name in text, f"Expected employee name in text"
        assert emp.get("jmbg", "") in text, f"Expected JMBG in text"
        assert emp.get("pozicija", "") in text, f"Expected pozicija in text"


class TestConvertedFromPdf:
    def _naziv_present(self, text, company):
        n = (company.get("naziv") or "").upper()
        ns = (company.get("naziv_skraceni") or "").upper()
        t = text.upper()
        return n and (n in t or (ns and ns in t) or n.split()[0] in t)

    def test_prijava_zanatstva(self, auth_headers, test_company):
        out = _generate(auth_headers, "prijava_zanatstva_.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert self._naziv_present(text, test_company), f"Expected company naziv in prijava_zanatstva"
        if test_company.get("pib"):
            assert test_company["pib"] in text

    def test_brisevi(self, auth_headers, test_company):
        out = _generate(auth_headers, "Zahtjev za uzorkovanje i ispitivanje - BRISEVA.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert self._naziv_present(text, test_company)
        if test_company.get("pib"):
            assert test_company["pib"] in text

    def test_hrana(self, auth_headers, test_company):
        out = _generate(auth_headers, "Zahtjev za uzorkovanje i ispitivanje - HRANA.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert self._naziv_present(text, test_company)
        if test_company.get("pib"):
            assert test_company["pib"] in text

    def test_voda(self, auth_headers, test_company):
        out = _generate(auth_headers, "Zahtjev za uzorkovanje i ispitivanje - VODA ZA PICE.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert self._naziv_present(text, test_company)

    def test_kazneno_fizicko_employee(self, auth_headers, test_company, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        out = _generate(
            auth_headers,
            "ZAHTJEV IZ KAZNENE EVIDENCIJE FIZICKO LICE.docx",
            test_company["id"],
            employee_id=emp["id"],
        )
        text = _read_docx_text(out["filename"]).upper()
        ime = emp["ime"].upper()
        prezime = emp["prezime"].upper()
        assert ime in text, f"Expected employee IME '{ime}' in text"
        assert prezime in text, f"Expected employee PREZIME '{prezime}' in text"

    def test_registracija_bezbjednost_hrane(self, auth_headers, test_company):
        out = _generate(auth_headers, "zahtjev za registraciju objekta za bezbjednost hrane.docx", test_company["id"])
        text = _read_docx_text(out["filename"])
        assert self._naziv_present(text, test_company)


class TestRegression:
    def test_ugovor_o_radu_zaposlenih(self, auth_headers, test_company, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        out = _generate(
            auth_headers,
            "UGOVOR O RADU Zaposlenih.docx",
            test_company["id"],
            employee_id=emp["id"],
        )
        text = _read_docx_text(out["filename"])
        # Verify company name + PIB + director + employee + JMBG present
        n = test_company.get("naziv", "")
        ns = test_company.get("naziv_skraceni", "")
        assert (n and n in text) or (ns and ns in text) or (n and n.split()[0] in text), "Company naziv missing"
        if test_company.get("pib"):
            assert test_company["pib"] in text, "PIB missing"
        full_name = f"{emp['ime']} {emp['prezime']}"
        assert full_name in text or emp["ime"] in text, "Employee name missing"
        assert emp.get("jmbg", "") in text, "JMBG missing"

    def test_aneks_generation(self, auth_headers, test_employee_with_datum_kraja):
        emp = test_employee_with_datum_kraja
        payload = {
            "employee_id": emp["id"],
            "nova_vrsta_ugovora": "Na neodređeno",
            "nova_plata_neto": 600.0,
        }
        r = requests.post(f"{API}/documents/generate-aneks", json=payload, headers=auth_headers, timeout=120)
        assert r.status_code == 200, f"Aneks failed: {r.text}"
        d = r.json()
        assert d.get("success") is True
        assert d.get("filename", "").endswith(".docx")
