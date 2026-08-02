"""Tests for 'UGOVOR O DOPUNSKOM RADU.docx' — signature wrap fix (paragraphs 36 & 37).

Verifies that the signature line uses a TAB (\u0009) with a right-aligned tab stop
at 9071 twips instead of a run of \u2006 (SIX-PER-EM-SPACE) characters, so that a
long naziv_skraceni like 'DOO ADVANCED ACCOUNTING ULCINJ' does not wrap to a new
line under the employee name.
"""
import io
import os
from pathlib import Path

import pytest
import requests
from dotenv import dotenv_values
from docx import Document
from docx.oxml.ns import qn

frontend_env = dotenv_values("/app/frontend/.env")
base_url = os.environ.get("REACT_APP_BACKEND_URL") or frontend_env.get("REACT_APP_BACKEND_URL")
if not base_url:
    raise RuntimeError("REACT_APP_BACKEND_URL missing")
BASE_URL = base_url.rstrip("/")

TEMPLATE = "UGOVOR O DOPUNSKOM RADU.docx"
COMPANY_ID = "d3fdf006-a6d9-4f4f-ab00-7018a1f30a54"  # DOO ADVANCED ACCOUNTING ULCINJ
EMPLOYEE_ID = "63283aaa-551b-4d67-9033-c946b336a610"  # EGZON CEKA


@pytest.fixture(scope="module")
def auth_token():
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"username": "getuard", "password": "Getuard1994."})
    if r.status_code != 200:
        pytest.fail(f"Login failed: {r.status_code} {r.text[:300]}")
    token = r.json().get("access_token") or r.json().get("token")
    if not token:
        pytest.fail(f"No token in response: {r.text[:300]}")
    return token


@pytest.fixture(scope="module")
def client(auth_token):
    s = requests.Session()
    s.headers.update({"Authorization": f"Bearer {auth_token}", "Content-Type": "application/json"})
    return s


@pytest.fixture(scope="module")
def company(client):
    r = client.get(f"{BASE_URL}/api/companies/{COMPANY_ID}")
    assert r.status_code == 200, f"Company fetch failed: {r.status_code} {r.text[:300]}"
    return r.json()


@pytest.fixture(scope="module")
def employee(client):
    r = client.get(f"{BASE_URL}/api/employees/{EMPLOYEE_ID}")
    assert r.status_code == 200, f"Employee fetch failed: {r.status_code} {r.text[:300]}"
    return r.json()


@pytest.fixture(scope="module")
def generated_doc(client):
    r = client.post(f"{BASE_URL}/api/documents/generate", json={
        "template_filename": TEMPLATE,
        "company_id": COMPANY_ID,
        "employee_id": EMPLOYEE_ID,
    })
    assert r.status_code == 200, f"Generate failed: {r.status_code} {r.text[:500]}"
    body = r.json()
    filename = body.get("filename")
    assert filename, f"No filename in response: {body}"

    dl = client.get(f"{BASE_URL}/api/documents/download/{filename}")
    assert dl.status_code == 200, f"Download failed: {dl.status_code}"
    doc = Document(io.BytesIO(dl.content))
    return doc


def _find_signature_paragraphs(doc):
    """Locate the two paragraphs: 'Izvršilac posla ... Poslodavac' label row + the name row directly after."""
    paras = doc.paragraphs
    label_idx = None
    # The signature label paragraph is short and starts with 'Izvršilac posla' and
    # ends with 'Poslodavac' — use tab-separated form, not the long 'Po isteku...' one.
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("Izvršilac posla") and t.endswith("Poslodavac") and len(t) < 40:
            label_idx = i
            break
    return label_idx, paras


class TestSignatureWrapFix:
    def test_label_paragraph_uses_tab(self, generated_doc):
        label_idx, paras = _find_signature_paragraphs(generated_doc)
        assert label_idx is not None, "Could not find 'Izvršilac posla ... Poslodavac' paragraph"
        label_text = paras[label_idx].text
        # No SIX-PER-EM-SPACE runs (\u2006) should remain
        assert "\u2006" not in label_text, (
            f"Label paragraph still contains \\u2006 chars: {label_text!r}"
        )
        # Must be a tab-separated label
        assert label_text == "Izvršilac posla\tPoslodavac", (
            f"Expected 'Izvršilac posla\\tPoslodavac', got: {label_text!r}"
        )

    def test_name_paragraph_uses_tab(self, generated_doc, company, employee):
        label_idx, paras = _find_signature_paragraphs(generated_doc)
        assert label_idx is not None
        name_para = paras[label_idx + 1]
        name_text = name_para.text

        naziv_skr = (company.get("naziv_skraceni") or company.get("naziv") or "").strip()
        ime = (employee.get("ime") or "").strip()
        prezime = (employee.get("prezime") or "").strip()
        full_name = f"{ime} {prezime}".strip()

        assert "\u2006" not in name_text, (
            f"Name paragraph still contains \\u2006 chars: {name_text!r}"
        )
        # Must contain exactly ONE tab separating the two names
        assert name_text.count("\t") == 1, (
            f"Expected exactly one TAB in name paragraph, got: {name_text!r}"
        )
        left, right = name_text.split("\t", 1)
        assert left.strip() == full_name, f"Left side mismatch: {left!r} vs {full_name!r}"
        # Right side must contain the naziv_skraceni contiguously (template may prepend 'DOO ')
        if naziv_skr:
            assert naziv_skr in right, f"naziv_skraceni {naziv_skr!r} not in right side {right!r}"
        assert "DOO ADVANCED ACCOUNTING ULCINJ" in right, (
            f"Expected 'DOO ADVANCED ACCOUNTING ULCINJ' contiguously on right side, got: {right!r}"
        )

    def test_no_placeholder_and_naziv_contiguous(self, generated_doc):
        label_idx, paras = _find_signature_paragraphs(generated_doc)
        name_text = paras[label_idx + 1].text
        assert "UNICO HIJA" not in name_text, f"Placeholder leaked: {name_text!r}"
        # The naziv_skraceni must be contiguous — no whitespace/newline between ACCOUNTING and ULCINJ
        assert "DOO ADVANCED ACCOUNTING ULCINJ" in name_text, (
            f"Company name not contiguous / missing: {name_text!r}"
        )

    def test_no_doo_doo_regression(self, generated_doc):
        full = "\n".join(p.text for p in generated_doc.paragraphs)
        for t in generated_doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full += "\n" + p.text
        assert "DOO DOO" not in full, "Regression: DOO DOO duplication found"

    def test_right_tab_stop_9071_in_name_paragraph_pPr(self, generated_doc):
        label_idx, paras = _find_signature_paragraphs(generated_doc)
        name_para = paras[label_idx + 1]
        pPr = name_para._p.find(qn("w:pPr"))
        assert pPr is not None, "Name paragraph has no pPr"
        tabs = pPr.find(qn("w:tabs"))
        assert tabs is not None, (
            f"Name paragraph pPr has no <w:tabs> element. XML: {name_para._p.xml[:800]}"
        )
        found = False
        for tab in tabs.findall(qn("w:tab")):
            val = tab.get(qn("w:val"))
            pos = tab.get(qn("w:pos"))
            if val == "right" and pos == "9071":
                found = True
                break
        assert found, (
            f"Expected <w:tab w:val='right' w:pos='9071'/> in pPr. "
            f"Got tabs XML: {tabs.xml if tabs is not None else None}"
        )

    def test_right_tab_stop_in_label_paragraph_pPr(self, generated_doc):
        label_idx, paras = _find_signature_paragraphs(generated_doc)
        label_para = paras[label_idx]
        pPr = label_para._p.find(qn("w:pPr"))
        assert pPr is not None, "Label paragraph has no pPr"
        tabs = pPr.find(qn("w:tabs"))
        assert tabs is not None, "Label paragraph pPr has no <w:tabs>"
        found = any(
            tab.get(qn("w:val")) == "right" and tab.get(qn("w:pos")) == "9071"
            for tab in tabs.findall(qn("w:tab"))
        )
        assert found, f"Label paragraph missing right-tab@9071. XML: {tabs.xml}"
