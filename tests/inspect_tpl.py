from docx import Document
import sys, glob

for f in [
    "odluka o imenovanju odgovornog lica za zastitu od pozara.docx",
    "odluka o imenovanje lice za Mobing.docx",
    "odluka o imenovanju lica za zastitu na radu.docx",
    "UGOVOR O RADU DIREKTOR.docx",
]:
    p = "/app/backend/templates/" + f
    print("=" * 20, f)
    d = Document(p)
    texts = [x.text.strip() for x in d.paragraphs if x.text.strip()]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.extend(pp.text.strip() for pp in cell.paragraphs if pp.text.strip())
    hits = [t for t in texts if "DIREKTOR" in t.upper()]
    print("DIREKTOR lines:", hits)
    print("tail:", texts[-8:])
